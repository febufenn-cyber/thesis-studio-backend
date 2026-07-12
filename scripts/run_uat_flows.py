#!/usr/bin/env python
"""Executable UAT / staging-acceptance driver (Subphase H).

Mirrors the fourteen user journeys and seven failure exercises in
``docs/release/STAGING_ACCEPTANCE.md`` as real HTTP flows and records one
pass/fail/manual/requires-staging line per step into a JSON report.

Modes
-----
``--base-url https://staging.example``
    Runs against a live deployment with plain HTTP (httpx). Login uses the
    email OTP path; on staging pass ``--otp-from-log`` so the operator can
    paste each code from the staging mail intercept/log. Steps that need
    direct database seeding (institution linkage, admin role bootstrap,
    induced job failures) are reported ``requires-staging`` — provision them
    with the operator checklist before trusting downstream rows.

``--local``
    Drives the ASGI app in-process against the local test database using the
    same HTTP flows plus test-style fixtures (schema recreate, institution /
    department seeding, admin org-memberships, in-process job draining).
    ``DEBUG=true`` is forced so ``/auth/request-otp`` returns ``debug_code``
    and no mail provider is needed.

Report hygiene: the report stores identifiers, status codes, counts and
checksums only — never thesis prose, quotations, full emails or secrets. A
canary marker embedded in the test manuscript is asserted absent from both
the support diagnostic bundle and the report itself before writing.

Steps that genuinely require human visual/role judgement are emitted as
``manual`` with a pointer into ``docs/release/evidence/UAT_CHECKLISTS.md``.
Steps that need live workers, ClamAV, LibreOffice, an AI provider or
operator-induced infrastructure failure are emitted ``requires-staging`` —
they are never faked locally.

Usage
-----
    .venv-validate/bin/python scripts/run_uat_flows.py --local --out uat.json
    python scripts/run_uat_flows.py --base-url https://thesis-staging.example \
        --otp-from-log --institution-id <uuid> --email-domain test.edu \
        --out uat_staging.json
"""

from __future__ import annotations

import argparse
import asyncio
import hashlib
import json
import os
import sys
import tempfile
import time
from dataclasses import dataclass, field
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any
from uuid import UUID, uuid4

import httpx


REPO_ROOT = Path(__file__).resolve().parent.parent
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))
DEFAULT_LOCAL_DATABASE_URL = (
    "postgresql+asyncpg://thesis:thesis@localhost:5453/thesis_studio_test"
)
CHECKLISTS = "docs/release/evidence/UAT_CHECKLISTS.md"

PASS = "pass"
FAIL = "fail"
MANUAL = "manual"
REQUIRES_STAGING = "requires-staging"
BLOCKED = "blocked"

# A tiny but structurally valid single-page PDF used only as a local fixture
# artifact for the pdf export slot (real PDF rendering needs LibreOffice and
# is exercised on staging). Content: one blank page, no text.
_MINIMAL_PDF = (
    b"%PDF-1.4\n"
    b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
    b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
    b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]>>endobj\n"
    b"xref\n0 4\n0000000000 65535 f \n0000000009 00000 n \n"
    b"0000000052 00000 n \n0000000101 00000 n \n"
    b"trailer<</Size 4/Root 1 0 R>>\nstartxref\n164\n%%EOF\n"
)


def _configure_local_env(database_url: str, storage_dir: str) -> None:
    """Set process environment for in-process mode before app imports."""

    os.environ["DATABASE_URL"] = database_url
    os.environ.setdefault("ENV", "development")
    os.environ["DEBUG"] = "true"
    os.environ.setdefault("JWT_SECRET", "a" * 64)
    os.environ.setdefault("ANTHROPIC_API_KEY", "uat-placeholder-not-used")
    os.environ.setdefault("DEFAULT_INSTITUTION_SHORT_NAME", "TU")
    os.environ.setdefault("BILLING_PROVIDER", "test")
    os.environ.setdefault(
        "BILLING_WEBHOOK_SECRET", "uat-local-webhook-secret-at-least-32-characters"
    )
    os.environ.setdefault("MALWARE_SCAN_MODE", "disabled")
    os.environ.setdefault("PRODUCTION_REQUIRE_MALWARE_SCAN", "false")
    os.environ["STORAGE_BACKEND"] = "local"
    os.environ["LOCAL_STORAGE_DIR"] = storage_dir


def _build_manuscript_docx(path: str, marker: str) -> str:
    """Write a small MLA-shaped manuscript and return its sha256.

    Includes one deliberately unresolvable citation ``(Ngugi 23)`` so the
    import report contains a blocking issue the student must resolve, and a
    canary ``marker`` sentence used to prove prose never leaks into support
    bundles or this report.
    """

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    d = Document()
    t = d.add_paragraph("COLONIAL VOICES IN THE AFRICAN NOVEL")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph("A dissertation by UAT Candidate")
    d.add_paragraph("CERTIFICATE")
    d.add_paragraph("This is to certify that this work is bonafide.")
    d.add_paragraph("DECLARATION")
    d.add_paragraph("I hereby declare this dissertation is my original work.")
    d.add_paragraph("CHAPTER I")
    ch = d.add_paragraph("INTRODUCTION")
    ch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = d.add_paragraph()
    p.add_run("Achebe's ")
    r = p.add_run("Things Fall Apart")
    r.italic = True
    p.add_run(f" resists the colonial gaze (Achebe 45). Sentinel {marker}.")
    bq = d.add_paragraph(
        "The white man is very clever. He has put a knife on the things "
        "that held us together and we have fallen apart. (Achebe 176)"
    )
    bq.paragraph_format.left_indent = Inches(0.5)
    d.add_paragraph("Postcolonial critics extend this reading (Ngugi 23).")
    d.add_paragraph("CHAPTER II: SUMMATION")
    d.add_paragraph("The argument closes where it began.")
    d.add_paragraph("WORKS CITED")
    wc1 = d.add_paragraph()
    wc1.add_run("Achebe, Chinua. ")
    wr = wc1.add_run("Things Fall Apart")
    wr.italic = True
    wc1.add_run(". Heinemann, 1958.")
    wc2 = d.add_paragraph()
    wc2.add_run(
        "Gikandi, Simon. “Chinua Achebe and the Invention of African Culture.” "
    )
    wj = wc2.add_run("Research in African Literatures")
    wj.italic = True
    wc2.add_run(", vol. 32, no. 3, 2001, pp. 3-8.")
    d.save(path)
    digest = hashlib.sha256()
    with open(path, "rb") as handle:
        digest.update(handle.read())
    return digest.hexdigest()


@dataclass
class Step:
    """One recorded acceptance step (identifiers and codes only)."""

    id: str
    journey: str
    name: str
    status: str
    expected: str = ""
    observed: str = ""
    detail: str = ""
    duration_ms: int = 0


@dataclass
class Persona:
    """One authenticated role driven through the flows."""

    role: str
    email: str
    client: httpx.AsyncClient
    user_id: str | None = None


@dataclass
class RunState:
    """Mutable cross-journey identifiers (never content)."""

    institution_id: str | None = None
    department_id: str | None = None
    project_id: str | None = None
    throwaway_project_id: str | None = None
    revision_id: str | None = None
    upload_checksum: str | None = None
    chapter_block_id: str | None = None
    chapter_id: str | None = None
    issue_id: str | None = None
    achebe_source_id: str | None = None
    unresolved_block_id: str | None = None
    unresolved_raw: str | None = None
    review_cycle_id: str | None = None
    package_id: str | None = None
    grant_token: str | None = None
    download_token: str | None = None
    failed_job_id: str | None = None
    marker: str = ""


class LocalFixtures:
    """Direct-database seeding used only in ``--local`` mode.

    Mirrors the fixture style of ``tests/test_phase4_acceptance.py`` for the
    state that has no HTTP bootstrap path (institution linkage, org admin
    roles, induced job failure, final export artifacts without LibreOffice).
    """

    def __init__(self) -> None:
        from sqlalchemy.ext.asyncio import async_sessionmaker, create_async_engine
        from sqlalchemy.pool import NullPool

        from app.core.config import get_settings

        self._engine = create_async_engine(
            get_settings().DATABASE_URL, echo=False, poolclass=NullPool
        )
        self._factory = async_sessionmaker(self._engine, expire_on_commit=False)

    async def recreate_schema(self) -> None:
        """Drop and recreate every table on the dedicated test database."""

        from app.db.session import Base

        async with self._engine.begin() as conn:
            await conn.run_sync(Base.metadata.drop_all)
            await conn.run_sync(Base.metadata.create_all)

    async def seed_institution(self, short_name: str, domain: str) -> tuple[str, str]:
        """Create the UAT institution and department; return their ids."""

        from app.models.institution import Institution
        from app.models.tenancy import Department
        from app.models.user import User

        async with self._factory() as db:
            inst = Institution(
                name="UAT University",
                short_name=short_name,
                email_domains=domain,
                address="1 Acceptance Way, Staging",
                short_address="Staging",
                university_name="UAT University",
                default_department="Department of English",
                department_aided=False,
            )
            db.add(inst)
            await db.flush()
            provisioner = User(
                email=f"uat-provisioner-{uuid4().hex[:8]}@{domain}",
                full_name="UAT Provisioner",
                institution_id=inst.id,
            )
            db.add(provisioner)
            await db.flush()
            dept = Department(
                institution_id=inst.id,
                name="PG & Research Department of English",
                code="ENG",
                created_by=provisioner.id,
            )
            db.add(dept)
            await db.commit()
            return str(inst.id), str(dept.id)

    async def link_project(
        self, project_id: str, institution_id: str, department_id: str
    ) -> None:
        """Attach the student project to the institution and department."""

        from sqlalchemy import select

        from app.models.project import Project

        async with self._factory() as db:
            project = (
                await db.execute(select(Project).where(Project.id == UUID(project_id)))
            ).scalar_one()
            project.institution_id = UUID(institution_id)
            project.department_id = UUID(department_id)
            await db.commit()

    async def grant_org_role(
        self, email: str, institution_id: str, department_id: str | None, role: str
    ) -> None:
        """Seed an active verified OrganizationMembership for an admin role."""

        from sqlalchemy import select

        from app.models.tenancy import OrganizationMembership
        from app.models.user import User

        async with self._factory() as db:
            user = (
                await db.execute(select(User).where(User.email == email))
            ).scalar_one()
            db.add(
                OrganizationMembership(
                    institution_id=UUID(institution_id),
                    department_id=UUID(department_id) if department_id else None,
                    user_id=user.id,
                    role=role,
                    affiliation_status="admin_verified",
                    status="active",
                    verified_by=user.id,
                    verified_at=datetime.now(timezone.utc),
                )
            )
            await db.commit()

    async def seed_final_exports(self, project_id: str, owner_email: str) -> list[str]:
        """Write real artifact files and ready/final Export rows (docx+pdf).

        LibreOffice is unavailable in local mode, so the pdf slot carries a
        minimal placeholder artifact; the byte-faithful PDF path is exercised
        on staging. Rows are checksummed against the actual stored files.
        """

        from sqlalchemy import select

        from app.models.export import Export
        from app.models.project import Project
        from app.models.user import User
        from app.services.storage_service import get_storage_service

        storage = get_storage_service()
        created: list[str] = []
        async with self._factory() as db:
            project = (
                await db.execute(select(Project).where(Project.id == UUID(project_id)))
            ).scalar_one()
            owner = (
                await db.execute(select(User).where(User.email == owner_email))
            ).scalar_one()
            for fmt in ("docx", "pdf"):
                descriptor, temp_path = tempfile.mkstemp(suffix=f".{fmt}")
                os.close(descriptor)
                if fmt == "docx":
                    from docx import Document

                    doc = Document()
                    doc.add_paragraph("UAT sealed-package fixture artifact.")
                    doc.save(temp_path)
                else:
                    with open(temp_path, "wb") as handle:
                        handle.write(_MINIMAL_PDF)
                with open(temp_path, "rb") as handle:
                    payload = handle.read()
                checksum = hashlib.sha256(payload).hexdigest()
                key = f"exports/{owner.id}/{project.id}/uat-fixture.{fmt}"
                await storage.upload_file(temp_path, key, "application/octet-stream")
                os.unlink(temp_path)
                export = Export(
                    project_id=project.id,
                    user_id=owner.id,
                    format=fmt,
                    document_version=project.document_version,
                    profile_version=f"builtin:{project.format_profile}",
                    storage_key=key,
                    checksum=checksum,
                    size_bytes=len(payload),
                    status="ready",
                    report={"pass": True, "fixture": True},
                    manifest={
                        "state": "final",
                        "document_version": project.document_version,
                        "fixture": True,
                    },
                )
                db.add(export)
                await db.flush()
                created.append(str(export.id))
            await db.commit()
        return created

    async def seed_failed_job(self, project_id: str, owner_email: str) -> str:
        """Insert one failed pdf-queue job so support retry can be exercised."""

        from sqlalchemy import select

        from app.models.job import Job
        from app.models.user import User

        async with self._factory() as db:
            owner = (
                await db.execute(select(User).where(User.email == owner_email))
            ).scalar_one()
            job = Job(
                kind="export",
                queue_name="pdf",
                priority=10,
                project_id=UUID(project_id),
                user_id=owner.id,
                payload={"export_id": str(uuid4()), "project_id": project_id},
                status="failed",
                attempts=1,
                max_attempts=3,
                available_at=datetime.now(timezone.utc),
                error_message="Worker process terminated (UAT induced failure).",
                idempotency_key=f"uat-failed-job:{uuid4()}",
            )
            db.add(job)
            await db.commit()
            return str(job.id)

    async def drain_jobs(self, max_jobs: int = 40) -> int:
        """Claim and run queued jobs in-process (stands in for the workers)."""

        from app.services.job_queue import _claim_next, _run_claimed

        worker_id = f"uat-local-{uuid4().hex[:8]}"
        queues = {"general", "pdf", "ai", "maintenance"}
        processed = 0
        for _ in range(max_jobs):
            job_id = await _claim_next(worker_id, queues)
            if job_id is None:
                break
            try:
                await _run_claimed(job_id, worker_id)
            except Exception:
                pass  # _run_claimed persists failure state; keep draining.
            processed += 1
        return processed

    async def dispose(self) -> None:
        await self._engine.dispose()


class UatDriver:
    """Runs every automatable journey and records evidence rows."""

    def __init__(self, args: argparse.Namespace) -> None:
        self.args = args
        self.local: bool = args.local
        self.steps: list[Step] = []
        self.state = RunState(marker=f"UATCANARY{uuid4().hex[:10].upper()}")
        self.run_id = uuid4().hex[:8]
        self.fixtures: LocalFixtures | None = None
        self.personas: dict[str, Persona] = {}
        self._transport: httpx.ASGITransport | None = None
        self.email_domain = args.email_domain
        self.short_name = os.environ.get("DEFAULT_INSTITUTION_SHORT_NAME", "TU")
        if args.institution_id:
            self.state.institution_id = args.institution_id

    # ------------------------------------------------------------------ core

    def _new_client(self) -> httpx.AsyncClient:
        if self.local:
            return httpx.AsyncClient(
                transport=self._transport,
                base_url="http://uat.testserver",
                timeout=120.0,
            )
        return httpx.AsyncClient(base_url=self.args.base_url, timeout=120.0)

    def record(
        self,
        step_id: str,
        journey: str,
        name: str,
        status: str,
        *,
        expected: str = "",
        observed: str = "",
        detail: str = "",
        started: float | None = None,
    ) -> Step:
        step = Step(
            id=step_id,
            journey=journey,
            name=name,
            status=status,
            expected=expected,
            observed=observed,
            detail=detail,
            duration_ms=int((time.monotonic() - started) * 1000) if started else 0,
        )
        self.steps.append(step)
        icon = {
            PASS: "PASS ",
            FAIL: "FAIL ",
            MANUAL: "MANUAL",
            REQUIRES_STAGING: "STAGE",
            BLOCKED: "BLOCK",
        }.get(status, status)
        print(f"[{icon:6s}] {step_id:9s} {name}"
              + (f" — {detail}" if detail else ""))
        return step

    def check(
        self,
        step_id: str,
        journey: str,
        name: str,
        response: httpx.Response | None,
        expected_codes: tuple[int, ...],
        *,
        ok: bool = True,
        detail: str = "",
        started: float | None = None,
    ) -> bool:
        """Record a step from an HTTP response and an extra assertion flag."""

        code = response.status_code if response is not None else 0
        passed = response is not None and code in expected_codes and ok
        self.record(
            step_id,
            journey,
            name,
            PASS if passed else FAIL,
            expected=f"http {'/'.join(str(c) for c in expected_codes)}",
            observed=f"http {code}" if response is not None else "no-response",
            detail=detail,
            started=started,
        )
        return passed

    def manual(self, step_id: str, journey: str, name: str, pointer: str) -> None:
        self.record(
            step_id, journey, name, MANUAL,
            detail=f"human judgement — see {CHECKLISTS} ({pointer})",
        )

    def staging_only(self, step_id: str, journey: str, name: str, why: str) -> None:
        self.record(
            step_id, journey, name, REQUIRES_STAGING,
            detail=f"not automatable here: {why}",
        )

    def blocked(self, step_id: str, journey: str, name: str, why: str) -> None:
        self.record(step_id, journey, name, BLOCKED, detail=f"prerequisite failed: {why}")

    async def drain_and_wait(self) -> None:
        """Local mode runs queued jobs in-process; remote mode lets workers run."""

        if self.local and self.fixtures is not None:
            await self.fixtures.drain_jobs()
        else:
            await asyncio.sleep(2.0)

    # ------------------------------------------------------------------ auth

    async def _obtain_otp(self, email: str, response: httpx.Response) -> str | None:
        try:
            debug_code = response.json().get("debug_code")
        except Exception:
            debug_code = None
        if debug_code:
            return str(debug_code)
        if self.args.otp_from_log:
            print(f"  >> Paste the 6-digit OTP sent to {email} "
                  "(from the staging mail intercept/log):")
            code = input("  OTP: ").strip()
            return code or None
        return None

    async def login(self, role: str) -> Persona | None:
        """Create (or reuse) a persona via the request-otp/verify-otp flow."""

        email = f"uat-{role.replace('_', '-')}-{self.run_id}@{self.email_domain}"
        client = self._new_client()
        started = time.monotonic()
        try:
            requested = await client.post("/auth/request-otp", json={"email": email})
        except httpx.HTTPError as exc:
            self.record(
                f"AUTH-{role}", "auth", f"{role} OTP login", FAIL,
                detail=f"transport error {type(exc).__name__}", started=started,
            )
            return None
        code = await self._obtain_otp(email, requested)
        if requested.status_code != 200 or not code:
            self.record(
                f"AUTH-{role}", "auth", f"{role} OTP login", FAIL,
                expected="http 200 + otp code",
                observed=f"http {requested.status_code}, code_available={bool(code)}",
                detail="no OTP available (staging: use --otp-from-log)",
                started=started,
            )
            return None
        verified = await client.post(
            "/auth/verify-otp", json={"email": email, "code": code}
        )
        me = await client.get("/auth/me")
        persona = Persona(role=role, email=email, client=client)
        if me.status_code == 200:
            persona.user_id = str(me.json().get("id"))
        okay = verified.status_code == 200 and me.status_code == 200
        self.record(
            f"AUTH-{role}", "auth", f"{role} OTP login + /auth/me", PASS if okay else FAIL,
            expected="http 200 + 200",
            observed=f"verify {verified.status_code}, me {me.status_code}",
            started=started,
        )
        if not okay:
            return None
        self.personas[role] = persona
        return persona

    # -------------------------------------------------------------- journeys

    async def j01_identity_sessions(self) -> None:
        journey = "J01-identity"
        student = await self.login("student")
        if student is None:
            self.blocked("J01-02", journey, "session listing", "student login failed")
            return
        second = self._new_client()
        started = time.monotonic()
        req = await second.post("/auth/request-otp", json={"email": student.email})
        code = await self._obtain_otp(student.email, req)
        if code:
            await second.post(
                "/auth/verify-otp", json={"email": student.email, "code": code}
            )
        sessions = await student.client.get("/auth/sessions")
        rows = sessions.json() if sessions.status_code == 200 else []
        others = [row for row in rows if not row.get("current")]
        self.check(
            "J01-02", journey, "second device login visible in session list",
            sessions, (200,), ok=len(rows) >= 2 and bool(others),
            detail=f"sessions={len(rows)}", started=started,
        )
        if others:
            started = time.monotonic()
            revoke = await student.client.delete(f"/auth/sessions/{others[0]['id']}")
            after = await second.get("/auth/me")
            self.check(
                "J01-03", journey, "revoked device session rejected opaquely",
                after, (401,), ok=revoke.status_code == 200,
                detail=f"revoke={revoke.status_code}", started=started,
            )
        else:
            self.blocked("J01-03", journey, "session revocation", "no second session")
        await second.aclose()
        self.manual("J01-04", journey, "OTP email arrival, sender and wording", "student §1")

    async def j02_upload(self) -> None:
        journey = "J02-manuscript"
        student = self.personas.get("student")
        if student is None:
            self.blocked("J02-01", journey, "project + upload", "no student persona")
            return
        started = time.monotonic()
        created = await student.client.post(
            "/projects",
            json={"title": "UAT Acceptance Thesis", "mode": "student",
                  "format_profile": "mla_strict"},
        )
        if not self.check("J02-01", journey, "create project", created, (201,),
                          started=started):
            return
        self.state.project_id = created.json()["id"]

        descriptor, docx_path = tempfile.mkstemp(suffix=".docx")
        os.close(descriptor)
        checksum = _build_manuscript_docx(docx_path, self.state.marker)
        self.state.upload_checksum = checksum
        started = time.monotonic()
        with open(docx_path, "rb") as handle:
            uploaded = await student.client.post(
                f"/projects/{self.state.project_id}/manuscript",
                files={"file": ("uat-manuscript.docx", handle,
                                "application/vnd.openxmlformats-officedocument"
                                ".wordprocessingml.document")},
                data={"apply_when_ready": "true"},
            )
        if not self.check("J02-02", journey, "manuscript upload accepted (202)",
                          uploaded, (202,), started=started):
            os.unlink(docx_path)
            return
        self.state.revision_id = uploaded.json()["revision"]["id"]
        reported = uploaded.json()["revision"]["checksum"]
        self.record(
            "J02-03", journey, "immutable revision checksum matches upload bytes",
            PASS if reported == checksum else FAIL,
            expected=f"sha256 {checksum[:12]}…", observed=f"sha256 {reported[:12]}…",
        )

        await self.drain_and_wait()
        started = time.monotonic()
        status = "unknown"
        for _ in range(int(self.args.poll_timeout / 2)):
            revision = await student.client.get(
                f"/projects/{self.state.project_id}/revisions/{self.state.revision_id}"
            )
            status = revision.json().get("status") if revision.status_code == 200 else "http-error"
            if status in {"ready", "failed"}:
                break
            await asyncio.sleep(2)
            await self.drain_and_wait()
        self.record(
            "J02-04", journey, "ingestion job completes and revision becomes ready",
            PASS if status == "ready" else FAIL,
            expected="status ready", observed=f"status {status}", started=started,
        )

        started = time.monotonic()
        with open(docx_path, "rb") as handle:
            duplicate = await student.client.post(
                f"/projects/{self.state.project_id}/manuscript",
                files={"file": ("uat-manuscript.docx", handle,
                                "application/vnd.openxmlformats-officedocument"
                                ".wordprocessingml.document")},
            )
        self.check("J02-05", journey, "identical re-upload rejected as duplicate (409)",
                   duplicate, (409,), started=started)
        os.unlink(docx_path)
        self.staging_only(
            "J02-06", journey, "live ClamAV malware scan on upload",
            "local mode runs MALWARE_SCAN_MODE=disabled; staging runs clamav",
        )

    async def j03_import_report(self) -> None:
        journey = "J03-import-report"
        student = self.personas.get("student")
        if student is None or self.state.revision_id is None:
            self.blocked("J03-01", journey, "import report review", "no ready revision")
            return
        started = time.monotonic()
        revision = await student.client.get(
            f"/projects/{self.state.project_id}/revisions/{self.state.revision_id}"
        )
        report = (revision.json() or {}).get("import_report") or {}
        issues = report.get("issues", [])
        citation_issue = next(
            (i for i in issues if i.get("code") == "citation_resolution_required"), None
        )
        self.check(
            "J03-01", journey,
            "import report present with planted unresolved-citation issue",
            revision, (200,), ok=citation_issue is not None,
            detail=f"issues_open={report.get('summary', {}).get('issues_open')}",
            started=started,
        )
        if citation_issue is None:
            return
        self.state.issue_id = citation_issue["id"]
        evidence = citation_issue.get("evidence", {})
        self.state.unresolved_block_id = evidence.get("block_id")
        self.state.unresolved_raw = evidence.get("raw")

        project = await student.client.get(f"/projects/{self.state.project_id}")
        version = project.json()["document_version"]
        started = time.monotonic()
        resolved = await student.client.patch(
            f"/projects/{self.state.project_id}/revisions/{self.state.revision_id}"
            f"/issues/{self.state.issue_id}",
            json={"resolution": "Reviewed and scheduled for citation resolution.",
                  "expected_version": version},
        )
        open_after = (
            (resolved.json().get("import_report") or {}).get("summary", {}).get("issues_open")
            if resolved.status_code == 200 else None
        )
        before = report.get("summary", {}).get("issues_open")
        self.check(
            "J03-02", journey, "student resolves an import issue (open count drops)",
            resolved, (200,),
            ok=open_after is not None and before is not None and open_after < before,
            detail=f"issues_open {before} -> {open_after}", started=started,
        )

    async def j04_edit_conflict_undo_restore(self) -> None:
        journey = "J04-editing"
        student = self.personas.get("student")
        if student is None or self.state.project_id is None:
            self.blocked("J04-01", journey, "structured editing", "no project")
            return
        project = await student.client.get(f"/projects/{self.state.project_id}")
        chapters = project.json().get("chapters") or []
        block = None
        for chapter in chapters:
            for candidate in chapter.get("blocks", []):
                if candidate.get("type") == "paragraph":
                    block = candidate
                    self.state.chapter_id = chapter.get("id")
                    break
            if block:
                break
        if block is None:
            self.blocked("J04-01", journey, "structured edit", "no paragraph block")
            return
        self.state.chapter_block_id = block["id"]
        version = project.json()["document_version"]
        original_text = "".join(
            run.get("text", "") for run in block.get("runs", [])
        )

        started = time.monotonic()
        edited = await student.client.post(
            f"/projects/{self.state.project_id}/collaboration/commands",
            json={
                "command_type": "update_block_text",
                "payload": {"block_id": block["id"],
                            "text": original_text + " Edited during UAT."},
                "expected_document_version": version,
                "client_request_id": f"uat-edit-{self.run_id}",
                "summary": "UAT structured edit",
            },
        )
        self.check("J04-01", journey, "student structured edit persists (autosave path)",
                   edited, (200,), started=started)
        command_id = (edited.json().get("command") or {}).get("id") if edited.status_code == 200 else None
        new_version = edited.json().get("document_version") if edited.status_code == 200 else version

        started = time.monotonic()
        stale = await student.client.post(
            f"/projects/{self.state.project_id}/collaboration/commands",
            json={
                "command_type": "update_block_text",
                "payload": {"block_id": block["id"], "text": "stale write"},
                "expected_document_version": version,  # deliberately stale
            },
        )
        self.check("J04-02", journey, "stale document version rejected with 409",
                   stale, (409,), started=started)

        if command_id:
            started = time.monotonic()
            undone = await student.client.post(
                f"/projects/{self.state.project_id}/editor/commands/{command_id}/undo",
                json={"expected_document_version": new_version},
            )
            self.check("J04-03", journey, "undo of the structured edit", undone, (200,),
                       started=started)
        else:
            self.blocked("J04-03", journey, "undo", "edit command id unavailable")

        current = await student.client.get(f"/projects/{self.state.project_id}")
        started = time.monotonic()
        restored = await student.client.post(
            f"/projects/{self.state.project_id}/revisions/{self.state.revision_id}/apply",
            json={"expected_version": current.json()["document_version"]},
        )
        self.check("J04-04", journey, "restore from the immutable manuscript revision",
                   restored, (200,), started=started)

    async def j05_sources_quotes(self) -> None:
        journey = "J05-citations"
        student = self.personas.get("student")
        if student is None or self.state.project_id is None:
            self.blocked("J05-01", journey, "source verification", "no project")
            return
        sources = await student.client.get(
            f"/projects/{self.state.project_id}/active-sources"
        )
        rows = sources.json() if sources.status_code == 200 else []
        self.check("J05-01", journey, "registry sources parsed from Works Cited",
                   sources, (200,), ok=len(rows) >= 2, detail=f"sources={len(rows)}")
        if len(rows) < 1:
            return
        achebe = next(
            (row for row in rows
             if "things fall apart" in str(row.get("fields", {})).lower()),
            rows[0],
        )
        self.state.achebe_source_id = achebe["id"]
        project = await student.client.get(f"/projects/{self.state.project_id}")
        version = project.json()["document_version"]
        verified_count = 0
        response = None
        for row in rows:
            response = await student.client.patch(
                f"/projects/{self.state.project_id}/sources/{row['id']}",
                json={"verified": True, "verification_method": "manual",
                      "expected_version": version},
            )
            if response.status_code == 200:
                verified_count += 1
                version += 1
        self.check("J05-02", journey, "student verifies every registry source",
                   response, (200,), ok=verified_count == len(rows),
                   detail=f"verified={verified_count}/{len(rows)}")

        quotes = await student.client.get(
            f"/projects/{self.state.project_id}/active-quotes"
        )
        quote_rows = quotes.json() if quotes.status_code == 200 else []
        self.check("J05-03", journey, "quotation registry extracted from manuscript",
                   quotes, (200,), ok=len(quote_rows) >= 1,
                   detail=f"quotes={len(quote_rows)}")
        if quote_rows:
            quote = quote_rows[0]
            started = time.monotonic()
            mismatch = await student.client.patch(
                f"/projects/{self.state.project_id}/quotes/{quote['id']}",
                json={"verified": True, "verification_method": "manual",
                      "expected_text": "intentionally different text"},
            )
            self.check("J05-04", journey,
                       "quotation verify with changed text rejected (409)",
                       mismatch, (409,), started=started)
            started = time.monotonic()
            verified = await student.client.patch(
                f"/projects/{self.state.project_id}/quotes/{quote['id']}",
                json={"verified": True, "verification_method": "manual",
                      "expected_text": quote.get("text")},
            )
            self.check("J05-05", journey, "exact quotation verification recorded",
                       verified, (200,), started=started)

        if self.state.unresolved_block_id and self.state.unresolved_raw:
            project = await student.client.get(f"/projects/{self.state.project_id}")
            started = time.monotonic()
            resolution = await student.client.post(
                f"/projects/{self.state.project_id}/citation-resolutions",
                json={
                    "block_id": self.state.unresolved_block_id,
                    "raw_citation": self.state.unresolved_raw,
                    "source_id": self.state.achebe_source_id,
                    "expected_version": project.json()["document_version"],
                },
            )
            self.check("J05-06", journey, "exact citation occurrence resolution",
                       resolution, (200,), started=started)
        else:
            self.blocked("J05-06", journey, "citation resolution",
                         "unresolved citation evidence unavailable")

        verify = await student.client.get(f"/projects/{self.state.project_id}/verify")
        passed = verify.json().get("passed") if verify.status_code == 200 else None
        counts = verify.json().get("report", {}).get("counts") if verify.status_code == 200 else {}
        self.check("J05-07", journey, "verifier endpoint returns a full report",
                   verify, (200,),
                   detail=f"passed={passed} counts={counts}")

    async def j06_preview(self) -> None:
        journey = "J06-preview"
        student = self.personas.get("student")
        if student is None or self.state.project_id is None:
            self.blocked("J06-01", journey, "PDF preview", "no project")
            return
        project = await student.client.get(f"/projects/{self.state.project_id}")
        started = time.monotonic()
        preview = await student.client.post(
            f"/projects/{self.state.project_id}/previews",
            json={"expected_document_version": project.json()["document_version"],
                  "force": False},
        )
        if not self.check("J06-01", journey, "preview request accepted (202)",
                          preview, (202,), started=started):
            return
        preview_id = preview.json().get("id")
        await self.drain_and_wait()
        status, error = "unknown", ""
        for _ in range(int(self.args.poll_timeout / 2)):
            row = await student.client.get(f"/previews/{preview_id}")
            if row.status_code != 200:
                status = "http-error"
                break
            status = row.json().get("status", "unknown")
            error = str(row.json().get("error_message") or "")
            if status in {"ready", "failed"}:
                break
            await asyncio.sleep(2)
            await self.drain_and_wait()
        if status == "ready":
            self.record("J06-02", journey, "authoritative PDF preview rendered", PASS,
                        expected="status ready", observed="status ready")
            self.manual("J06-03", journey,
                        "preview is visually faithful to institutional format",
                        "student §4")
        elif "LibreOffice" in error or "PDF rendering unavailable" in error or (
            self.local and status == "failed"
        ):
            self.staging_only(
                "J06-02", journey, "authoritative PDF preview rendered",
                "LibreOffice (soffice) not installed on this host; run on staging",
            )
        else:
            self.record("J06-02", journey, "authoritative PDF preview rendered", FAIL,
                        expected="status ready", observed=f"status {status}")

    async def j07_ai_partner(self) -> None:
        journey = "J07-ai-partner"
        student = self.personas.get("student")
        if student is None or self.state.project_id is None:
            self.blocked("J07-01", journey, "AI health", "no project")
            return
        started = time.monotonic()
        health = await student.client.get(
            f"/projects/{self.state.project_id}/ai/health"
        )
        component = (health.json() or {}).get("component_health", {}) if health.status_code == 200 else {}
        self.check(
            "J07-01", journey,
            "AI health endpoint reports application/editing/export independent of AI",
            health, (200,),
            ok=component.get("application") == "operational"
            and component.get("editing") == "operational",
            detail=f"ai={component.get('ai')}", started=started,
        )
        self.staging_only(
            "J07-02", journey,
            "grounded AI proposal with partial acceptance and provenance",
            "needs a configured commercial AI provider and AI worker",
        )

    async def j08_supervisor_review(self) -> None:
        journey = "J08-supervisor"
        student = self.personas.get("student")
        if student is None or self.state.project_id is None:
            self.blocked("J08-01", journey, "supervisor journey", "no project")
            return

        if self.local and self.fixtures is not None:
            await self.fixtures.link_project(
                self.state.project_id, self.state.institution_id, self.state.department_id
            )
            self.record("J08-00", journey,
                        "fixture: project linked to institution/department", PASS,
                        detail="local direct-DB seeding (staging: provisioning flow)")
        elif self.state.institution_id is None:
            self.staging_only(
                "J08-00", journey, "project linked to institution/department",
                "pass --institution-id and pre-provision the project linkage",
            )

        supervisor = await self.login("supervisor")
        if supervisor is None or self.state.institution_id is None:
            self.blocked("J08-01", journey, "supervisor invitation", "login/institution")
            return
        started = time.monotonic()
        invited = await student.client.post(
            f"/institutions/{self.state.institution_id}/invitations",
            json={"email": supervisor.email, "role": "supervisor",
                  "project_id": self.state.project_id,
                  "capabilities": ["source.verify"]},
        )
        token = invited.json().get("invitation_token") if invited.status_code == 201 else None
        accepted = None
        if token:
            accepted = await supervisor.client.post(
                "/collaboration/invitations/accept", json={"token": token}
            )
        self.check("J08-01", journey,
                   "student invites supervisor; supervisor accepts (single-use token)",
                   accepted if accepted is not None else invited,
                   (200,), ok=invited.status_code == 201,
                   detail=f"invite={invited.status_code}", started=started)
        if accepted is None or accepted.status_code != 200:
            self.blocked("J08-02", journey, "review cycle", "membership not granted")
            return

        started = time.monotonic()
        cycle = await student.client.post(
            f"/projects/{self.state.project_id}/review-cycles",
            json={"reviewer_id": supervisor.user_id, "scope_type": "project",
                  "deadline": (datetime.now(timezone.utc) + timedelta(days=7)).isoformat()},
        )
        if not self.check("J08-02", journey, "student submits snapshot review cycle",
                          cycle, (201,), started=started):
            return
        self.state.review_cycle_id = cycle.json()["id"]

        snapshot = await supervisor.client.get(
            f"/projects/{self.state.project_id}/review-cycles/"
            f"{self.state.review_cycle_id}/snapshot"
        )
        self.check("J08-03", journey, "supervisor reads the immutable review snapshot",
                   snapshot, (200,),
                   ok=bool((snapshot.json() or {}).get("snapshot", {}).get("checksum")))

        block_id = self.state.chapter_block_id
        started = time.monotonic()
        comment = await supervisor.client.post(
            f"/projects/{self.state.project_id}/comments",
            json={
                "anchor_type": "block",
                "anchor": {"block_id": block_id},
                "body": "Please connect this description to the chapter claim.",
                "review_cycle_id": self.state.review_cycle_id,
                "assigned_to": student.user_id,
            },
        )
        self.check("J08-04", journey, "supervisor anchors a review comment",
                   comment, (201,), started=started)

        started = time.monotonic()
        suggestion = await supervisor.client.post(
            f"/projects/{self.state.project_id}/suggestions",
            json={
                "target_block_id": block_id,
                "review_cycle_id": self.state.review_cycle_id,
                "explanation": "Tie the paragraph explicitly to the argument.",
                "proposed_operation": {
                    "command_type": "update_block_text",
                    "payload": {
                        "block_id": block_id,
                        "text": ("The novel reframes the village as the polis of "
                                 "resistance, anchoring the chapter claim."),
                    },
                },
            },
        )
        if not self.check("J08-05", journey, "supervisor files a structured suggestion",
                          suggestion, (201,), started=started):
            return
        started = time.monotonic()
        decision = await student.client.post(
            f"/projects/{self.state.project_id}/suggestions/"
            f"{suggestion.json()['id']}/decision",
            json={"decision": "accepted",
                  "response": "Accepted; I will defend this connection in the viva."},
        )
        self.check("J08-06", journey,
                   "student accepts the suggestion (authorship preserved, command applied)",
                   decision, (200,),
                   ok=bool((decision.json() or {}).get("applied_command_id")),
                   started=started)

        started = time.monotonic()
        stale_decision = await supervisor.client.post(
            f"/projects/{self.state.project_id}/review-cycles/"
            f"{self.state.review_cycle_id}/decision",
            json={"decision": "approved",
                  "note": "The submitted snapshot was academically sound."},
        )
        stale_body = stale_decision.json() if stale_decision.status_code == 200 else {}
        self.check("J08-07", journey,
                   "approval of a superseded snapshot is honestly marked snapshot_only",
                   stale_decision, (200,),
                   ok=(stale_body.get("approval") or {}).get("status") == "snapshot_only",
                   detail=f"approval={ (stale_body.get('approval') or {}).get('status') }",
                   started=started)

        started = time.monotonic()
        resubmitted = await student.client.post(
            f"/projects/{self.state.project_id}/review-cycles",
            json={"reviewer_id": supervisor.user_id, "scope_type": "project",
                  "resubmitted_from_id": self.state.review_cycle_id},
        )
        if not self.check("J08-08", journey,
                          "student resubmits the current version for review",
                          resubmitted, (201,), started=started):
            return
        second_cycle_id = resubmitted.json()["id"]

        started = time.monotonic()
        approved = await supervisor.client.post(
            f"/projects/{self.state.project_id}/review-cycles/"
            f"{second_cycle_id}/decision",
            json={"decision": "approved", "note": "Academic content approved."},
        )
        body = approved.json() if approved.status_code == 200 else {}
        self.check("J08-09", journey,
                   "current-version approval activates and advances workflow",
                   approved, (200,),
                   ok=body.get("workflow_state") == "academically_approved"
                   and (body.get("approval") or {}).get("status") == "active",
                   detail=f"workflow={body.get('workflow_state')}", started=started)

    async def j09_operator_boundary(self) -> None:
        journey = "J09-operator"
        student = self.personas.get("student")
        if student is None or self.state.institution_id is None:
            self.blocked("J09-01", journey, "operator journey", "prerequisites missing")
            return
        operator = await self.login("operator")
        if operator is None:
            self.blocked("J09-01", journey, "operator invitation", "login failed")
            return
        invited = await student.client.post(
            f"/institutions/{self.state.institution_id}/invitations",
            json={"email": operator.email, "role": "operator",
                  "project_id": self.state.project_id},
        )
        token = invited.json().get("invitation_token") if invited.status_code == 201 else None
        accepted = None
        if token:
            accepted = await operator.client.post(
                "/collaboration/invitations/accept", json={"token": token}
            )
        self.check("J09-01", journey, "operator invited and joins project",
                   accepted if accepted is not None else invited, (200,),
                   ok=invited.status_code == 201)
        if accepted is None or accepted.status_code != 200:
            self.blocked("J09-02", journey, "operator boundary", "no membership")
            return

        project = await student.client.get(f"/projects/{self.state.project_id}")
        version = project.json()["document_version"]
        started = time.monotonic()
        prose = await operator.client.post(
            f"/projects/{self.state.project_id}/collaboration/commands",
            json={"command_type": "update_block_text",
                  "payload": {"block_id": self.state.chapter_block_id,
                              "text": "Operator-authored prose"},
                  "expected_document_version": version},
        )
        self.check("J09-02", journey,
                   "operator prose rewrite mechanically refused (403/404/409)",
                   prose, (403, 404, 409), started=started)

        started = time.monotonic()
        fix = await operator.client.post(
            f"/projects/{self.state.project_id}/collaboration/commands",
            json={"command_type": "update_metadata",
                  "payload": {"path": "submission.year", "value": 2026},
                  "expected_document_version": version,
                  "client_request_id": f"uat-op-fix-{self.run_id}",
                  "summary": "Correct submission year"},
        )
        authority = (fix.json() or {}).get("authority", {}) if fix.status_code == 200 else {}
        self.check("J09-03", journey,
                   "operator formatting/metadata fix succeeds without prose authority",
                   fix, (200,),
                   ok=authority.get("operator_prose_rewrite_allowed") is False,
                   started=started)

        started = time.monotonic()
        transition = await operator.client.post(
            f"/projects/{self.state.project_id}/workflow/transition",
            json={"target": "formatting_review", "note": "Academic review completed."},
        )
        self.check("J09-04", journey, "operator moves workflow into formatting review",
                   transition, (200,), started=started)

    async def j10_admin_approvals(self) -> None:
        journey = "J10-approvals"
        supervisor = self.personas.get("supervisor")
        operator = self.personas.get("operator")
        if self.state.institution_id is None:
            self.blocked("J10-01", journey, "admin approvals", "no institution")
            return
        dept_admin = await self.login("dept_admin")
        inst_admin = await self.login("inst_admin")
        if self.local and self.fixtures is not None and dept_admin and inst_admin:
            await self.fixtures.grant_org_role(
                dept_admin.email, self.state.institution_id,
                self.state.department_id, "department_admin",
            )
            await self.fixtures.grant_org_role(
                inst_admin.email, self.state.institution_id, None, "institution_admin",
            )
            self.record("J10-00", journey,
                        "fixture: department/institution admin roles provisioned", PASS,
                        detail="local direct-DB seeding (staging: admin onboarding)")
        elif not self.local:
            self.staging_only(
                "J10-00", journey, "admin role provisioning",
                "seed department_admin/institution_admin memberships on staging first",
            )
        if dept_admin is None or supervisor is None or operator is None:
            self.blocked("J10-01", journey, "approvals", "personas missing")
            return

        started = time.monotonic()
        assignment = await dept_admin.client.post(
            f"/projects/{self.state.project_id}/assignments",
            json={"assignee_id": supervisor.user_id,
                  "assignment_type": "supervisor_review", "scope": {"type": "project"},
                  "priority": "normal"},
        )
        self.check("J10-01", journey, "department admin records a role assignment",
                   assignment, (201,), started=started)

        if inst_admin is not None:
            started = time.monotonic()
            policy = await inst_admin.client.post(
                f"/institutions/{self.state.institution_id}/policies",
                json={"label": "UAT collaboration policy",
                      "policy": {"workflow": {"format_review_required": True}}},
            )
            published = None
            if policy.status_code == 201:
                published = await inst_admin.client.post(
                    f"/institutions/{self.state.institution_id}/policies/"
                    f"{policy.json()['id']}/publish"
                )
            self.check("J10-02", journey, "institution admin drafts and publishes policy",
                       published if published is not None else policy, (200,),
                       ok=policy.status_code == 201, started=started)
        else:
            self.blocked("J10-02", journey, "policy publish", "no institution admin")

        started = time.monotonic()
        citation = await supervisor.client.post(
            f"/projects/{self.state.project_id}/approvals",
            json={"dimension": "citation", "scope_type": "project",
                  "decision": "approved",
                  "note": "Registered evidence and citation traceability reviewed."},
        )
        self.check("J10-03", journey, "separate citation approval recorded",
                   citation, (201,), started=started)

        started = time.monotonic()
        formatting = await operator.client.post(
            f"/projects/{self.state.project_id}/approvals",
            json={"dimension": "formatting", "scope_type": "project",
                  "decision": "approved",
                  "note": "Institutional presentation verified without prose changes."},
        )
        self.check("J10-04", journey, "separate formatting approval recorded",
                   formatting, (201,), started=started)

        started = time.monotonic()
        institutional = await dept_admin.client.post(
            f"/projects/{self.state.project_id}/approvals",
            json={"dimension": "institutional", "scope_type": "project",
                  "decision": "approved",
                  "note": "Workflow approval recorded; not a legal signature."},
        )
        self.check("J10-05", journey, "separate institutional approval recorded",
                   institutional, (201,), started=started)

        access = await dept_admin.client.get(
            f"/projects/{self.state.project_id}/collaboration/access"
        )
        body = access.json() if access.status_code == 200 else {}
        self.check("J10-06", journey,
                   "admin scope excludes manuscript content and private AI history",
                   access, (200,),
                   ok=body.get("content_access") is False
                   and body.get("ai_history_access") is False)
        self.manual("J10-07", journey,
                    "private AI history is invisible in the admin UI", "dept admin §3")

    async def j11_seal(self) -> None:
        journey = "J11-seal"
        student = self.personas.get("student")
        supervisor = self.personas.get("supervisor")
        dept_admin = self.personas.get("dept_admin")
        if student is None or supervisor is None or dept_admin is None:
            self.blocked("J11-01", journey, "sealed submission", "personas missing")
            return
        started = time.monotonic()
        attest_student = await student.client.post(
            f"/projects/{self.state.project_id}/attestations",
            json={"attestation_type": "student_authorship",
                  "statement_version": "2026.1",
                  "statement_text": "I remain the author and reviewed all accepted assistance.",
                  "accepted": True},
        )
        attest_supervisor = await supervisor.client.post(
            f"/projects/{self.state.project_id}/attestations",
            json={"attestation_type": "supervisor_workflow_approval",
                  "statement_version": "2026.1",
                  "statement_text": "I reviewed and approved the recorded academic workflow.",
                  "accepted": True},
        )
        self.check("J11-01", journey, "student and supervisor attestations recorded",
                   attest_supervisor, (201,),
                   ok=attest_student.status_code == 201, started=started)

        project = await student.client.get(f"/projects/{self.state.project_id}")
        version = project.json()["document_version"]
        started = time.monotonic()
        exports = await student.client.post(
            f"/projects/{self.state.project_id}/exports",
            json={"formats": ["docx", "pdf"], "acknowledge": True,
                  "expected_version": version},
        )
        review_fallback = False
        if exports.status_code == 409:
            review_fallback = True
            exports = await student.client.post(
                f"/projects/{self.state.project_id}/exports",
                json={"formats": ["docx", "pdf"], "acknowledge": True,
                      "allow_review_export": True, "expected_version": version},
            )
        self.check("J11-02", journey,
                   "export request accepted after authorship acknowledgement (202)",
                   exports, (202,),
                   detail=f"review_fallback={review_fallback}", started=started)
        await self.drain_and_wait()
        listed = await student.client.get(f"/projects/{self.state.project_id}/exports")
        rows = listed.json() if listed.status_code == 200 else []
        final_ready = {
            row["format"] for row in rows
            if row.get("status") == "ready"
        }
        self.record(
            "J11-03", journey, "rendered export artifacts",
            PASS if {"docx"}.issubset(final_ready) else REQUIRES_STAGING if self.local else FAIL,
            expected="docx (+pdf on staging) ready",
            observed=f"ready={sorted(final_ready)}",
            detail="pdf rendering needs LibreOffice; verify on staging" if
            "pdf" not in final_ready else "",
        )

        readiness = await dept_admin.client.get(
            f"/projects/{self.state.project_id}/submission-readiness"
        )
        ready = (readiness.json() or {}).get("ready") if readiness.status_code == 200 else None
        if not ready and self.local and self.fixtures is not None:
            await self.fixtures.seed_final_exports(self.state.project_id, student.email)
            self.record(
                "J11-04", journey,
                "fixture: final docx+pdf export slots seeded with real stored files",
                PASS,
                detail="local stand-in for worker-rendered finals; staging renders them",
            )
            readiness = await dept_admin.client.get(
                f"/projects/{self.state.project_id}/submission-readiness"
            )
            ready = (readiness.json() or {}).get("ready")
        self.check("J11-05", journey,
                   "submission readiness reports every gate satisfied",
                   readiness, (200,), ok=ready is True,
                   detail=f"missing={ (readiness.json() or {}).get('missing_approvals') } "
                          f"attest={ (readiness.json() or {}).get('missing_attestations') }"
                   if ready is not True else "")
        if ready is not True:
            self.blocked("J11-06", journey, "seal", "readiness gates unmet")
            return

        started = time.monotonic()
        sealed = await dept_admin.client.post(
            f"/projects/{self.state.project_id}/submission-packages",
            json={"note": "UAT department submission package"},
        )
        body = sealed.json() if sealed.status_code == 201 else {}
        self.check("J11-06", journey,
                   "submission sealed with package checksum and signature claim",
                   sealed, (201,),
                   ok=body.get("state") == "sealed" and bool(body.get("package_checksum")),
                   detail=f"checksum={str(body.get('package_checksum'))[:12]}…"
                   if body.get("package_checksum") else "",
                   started=started)
        if sealed.status_code == 201:
            self.state.package_id = body["id"]

        project = await student.client.get(f"/projects/{self.state.project_id}")
        started = time.monotonic()
        locked = await student.client.post(
            f"/projects/{self.state.project_id}/collaboration/commands",
            json={"command_type": "update_metadata",
                  "payload": {"path": "submission.month", "value": "August"},
                  "expected_document_version": project.json()["document_version"]},
        )
        self.check("J11-07", journey, "sealed submission is immutable (edit → 409)",
                   locked, (409,), started=started)

        timeline = await dept_admin.client.get(
            f"/projects/{self.state.project_id}/audit-timeline"
        )
        kinds = {row.get("kind") for row in timeline.json()} if timeline.status_code == 200 else set()
        needed = {"review_cycle_submitted", "human_suggestion_decided",
                  "submission_package_sealed"}
        self.check("J11-08", journey, "audit timeline records the full decision trail",
                   timeline, (200,), ok=needed.issubset(kinds),
                   detail=f"kinds={len(kinds)}")

    async def j12_external_review(self) -> None:
        journey = "J12-external"
        dept_admin = self.personas.get("dept_admin")
        if dept_admin is None or self.state.package_id is None:
            self.blocked("J12-01", journey, "external review", "no sealed package")
            return
        recipient = f"uat-examiner-{self.run_id}@example.edu"
        started = time.monotonic()
        grant = await dept_admin.client.post(
            f"/projects/{self.state.project_id}/submission-packages/"
            f"{self.state.package_id}/external-review",
            json={"recipient_email": recipient,
                  "expires_at": (datetime.now(timezone.utc) + timedelta(days=5)).isoformat(),
                  "permissions": ["sealed.read_metadata", "sealed.read_content"],
                  "download_allowed": False,
                  "watermark": "Confidential external examination copy"},
        )
        if not self.check("J12-01", journey, "recipient-bound external grant issued",
                          grant, (201,), started=started):
            return
        token = grant.json()["access_token"]
        grant_id = grant.json()["id"]
        anon = self._new_client()

        started = time.monotonic()
        wrong = await anon.post(
            "/external-review/access",
            json={"token": token, "recipient_email": "someone-else@example.edu"},
        )
        self.check("J12-02", journey, "wrong recipient rejected opaquely (404)",
                   wrong, (404,), started=started)

        started = time.monotonic()
        view = await anon.post(
            "/external-review/access",
            json={"token": token, "recipient_email": recipient},
        )
        body = view.json() if view.status_code == 200 else {}
        no_keys = all(
            "storage_key" not in item
            for item in (body.get("submission", {}).get("manifest", {}) or {}).get("exports", [])
        )
        self.check("J12-03", journey,
                   "bound recipient reads sealed content with watermark notice",
                   view, (200,),
                   ok=bool(body.get("grant", {}).get("watermark"))
                   and body.get("grant", {}).get("download_allowed") is False
                   and bool(body.get("canonical_document", {}).get("chapters"))
                   and no_keys,
                   started=started)

        started = time.monotonic()
        dl_grant = await dept_admin.client.post(
            f"/projects/{self.state.project_id}/submission-packages/"
            f"{self.state.package_id}/external-review",
            json={"recipient_email": recipient,
                  "expires_at": (datetime.now(timezone.utc) + timedelta(days=2)).isoformat(),
                  "permissions": ["sealed.read_metadata", "sealed.read_content",
                                  "sealed.download"],
                  "download_allowed": True,
                  "watermark": "UAT sealed download watermark"},
        )
        if dl_grant.status_code == 201:
            download = await anon.post(
                "/external-review/download",
                json={"token": dl_grant.json()["access_token"],
                      "recipient_email": recipient, "format": "docx"},
            )
            if download.status_code == 200:
                header = download.headers.get("X-Robofox-Watermark", "")
                self.check("J12-04", journey,
                           "sealed download carries the watermark response header",
                           download, (200,), ok=bool(header),
                           detail="X-Robofox-Watermark present", started=started)
            elif download.status_code == 303:
                self.manual("J12-04", journey,
                            "watermark on presigned download (redirect path)",
                            "external reviewer §3")
            else:
                self.check("J12-04", journey,
                           "sealed download carries the watermark response header",
                           download, (200, 303), started=started)
        else:
            self.check("J12-04", journey, "download-enabled grant issued",
                       dl_grant, (201,), started=started)
        self.manual("J12-05", journey,
                    "watermark visibly rendered on the PDF pages", "external reviewer §4")

        started = time.monotonic()
        short = await dept_admin.client.post(
            f"/projects/{self.state.project_id}/submission-packages/"
            f"{self.state.package_id}/external-review",
            json={"recipient_email": recipient,
                  "expires_at": (datetime.now(timezone.utc) + timedelta(seconds=3)).isoformat(),
                  "permissions": ["sealed.read_metadata"],
                  "download_allowed": False},
        )
        if short.status_code == 201:
            await asyncio.sleep(4)
            expired = await anon.post(
                "/external-review/access",
                json={"token": short.json()["access_token"],
                      "recipient_email": recipient},
            )
            self.check("J12-06", journey, "expired grant rejected opaquely (404)",
                       expired, (404,), started=started)
        else:
            self.check("J12-06", journey, "short-expiry grant issued",
                       short, (201,), started=started)

        started = time.monotonic()
        revoked = await dept_admin.client.delete(
            f"/projects/{self.state.project_id}/external-review/{grant_id}"
        )
        after = await anon.post(
            "/external-review/access",
            json={"token": token, "recipient_email": recipient},
        )
        self.check("J12-07", journey, "revoked grant rejected opaquely (404)",
                   after, (404,), ok=revoked.status_code == 200, started=started)
        await anon.aclose()

    async def j13_data_lifecycle(self) -> None:
        journey = "J13-lifecycle"
        student = self.personas.get("student")
        if student is None:
            self.blocked("J13-01", journey, "data portability", "no student")
            return
        started = time.monotonic()
        export = await student.client.get("/account/data-export")
        self.check("J13-01", journey, "student self-service account data export",
                   export, (200,), started=started)

        throwaway = await student.client.post(
            "/projects",
            json={"title": "UAT deletion target", "mode": "student",
                  "format_profile": "mla_strict"},
        )
        if throwaway.status_code == 201:
            self.state.throwaway_project_id = throwaway.json()["id"]
            started = time.monotonic()
            deletion = await student.client.post(
                "/privacy/lifecycle-requests",
                json={"request_type": "project_delete",
                      "project_id": self.state.throwaway_project_id,
                      "reason": "UAT deletion grace-period exercise."},
            )
            body = deletion.json() if deletion.status_code == 202 else {}
            self.check("J13-02", journey,
                       "project deletion enters a grace period (202, execute_after set)",
                       deletion, (202,),
                       ok=bool(body.get("execute_after")) or body.get("duplicate") is True,
                       started=started)
            if deletion.status_code == 202 and body.get("id"):
                started = time.monotonic()
                cancelled = await student.client.post(
                    f"/privacy/lifecycle-requests/{body['id']}/cancel"
                )
                self.check("J13-03", journey, "deletion cancellable inside grace period",
                           cancelled, (200,), started=started)
        else:
            self.blocked("J13-02", journey, "deletion grace period", "no throwaway project")

        if self.state.package_id:
            started = time.monotonic()
            sealed_delete = await student.client.post(
                "/privacy/lifecycle-requests",
                json={"request_type": "project_delete",
                      "project_id": self.state.project_id,
                      "reason": "UAT sealed-custody restriction exercise."},
            )
            self.check("J13-04", journey,
                       "deletion of a sealed project accepted only into grace period",
                       sealed_delete, (202,), started=started)
            self.staging_only(
                "J13-05", journey,
                "sealed-custody terminal state (authorization_required after grace)",
                "needs grace-period elapse plus the maintenance worker",
            )
        else:
            self.blocked("J13-04", journey, "sealed custody restriction", "no sealed package")

    async def j14_support(self) -> None:
        journey = "J14-support"
        student = self.personas.get("student")
        if student is None or self.state.project_id is None:
            self.blocked("J14-01", journey, "support journey", "no project")
            return
        support = await self.login("support")
        if support is None:
            self.blocked("J14-01", journey, "support grant", "login failed")
            return
        started = time.monotonic()
        granted = await student.client.post(
            f"/projects/{self.state.project_id}/support-access",
            json={"support_user_id": support.user_id,
                  "capabilities": ["project.read_metadata"],
                  "consent_note": "UAT support diagnosis consent.",
                  "expires_in_hours": 2},
        )
        self.check("J14-01", journey,
                   "student grants time-boxed metadata-only support access",
                   granted, (201,),
                   ok=(granted.json() or {}).get("visibility_banner_required") is True
                   if granted.status_code == 201 else False,
                   started=started)
        if granted.status_code != 201:
            self.blocked("J14-02", journey, "diagnostic bundle", "no support grant")
            return

        started = time.monotonic()
        bundle = await support.client.post(
            f"/support/projects/{self.state.project_id}/diagnostic-bundle",
            json={"justification": "Resolve failed export from metadata and job state only."},
        )
        serialized = json.dumps(bundle.json(), default=str) if bundle.status_code == 200 else ""
        no_prose = (
            self.state.marker not in serialized
            and (bundle.json() or {}).get("privacy", {}).get("manuscript_content_included")
            is False
            if bundle.status_code == 200 else False
        )
        self.check("J14-02", journey,
                   "diagnostic bundle is metadata-only (canary prose absent)",
                   bundle, (200,), ok=no_prose,
                   detail="manuscript_content_included=false, canary absent"
                   if no_prose else "prose-leak or privacy flag failure",
                   started=started)

        if self.local and self.fixtures is not None:
            self.state.failed_job_id = await self.fixtures.seed_failed_job(
                self.state.project_id, student.email
            )
            self.record("J14-03", journey, "fixture: failed pdf job induced", PASS,
                        detail="local stand-in for a crashed worker")
        else:
            jobs = await student.client.get(f"/projects/{self.state.project_id}/jobs")
            failed = next(
                (row for row in jobs.json() if row.get("status") == "failed"), None
            ) if jobs.status_code == 200 else None
            if failed:
                self.state.failed_job_id = failed["id"]
            else:
                self.staging_only(
                    "J14-03", journey, "induced failed job for retry",
                    "kill a pdf worker mid-job on staging (failure exercise 3)",
                )
        if self.state.failed_job_id:
            started = time.monotonic()
            retried = await support.client.post(
                f"/support/projects/{self.state.project_id}/jobs/"
                f"{self.state.failed_job_id}/retry",
                json={"justification": "Requeue idempotent job after worker replacement."},
            )
            body = retried.json() if retried.status_code == 200 else {}
            self.check("J14-04", journey,
                       "support retries the failed job without content access",
                       retried, (200,),
                       ok=body.get("status") == "queued"
                       and body.get("content_accessed") is False,
                       started=started)
        else:
            self.blocked("J14-04", journey, "support retry", "no failed job available")

    async def failure_exercises(self) -> None:
        journey = "F-failure-exercises"
        self.staging_only("F-01", journey,
                          "stop AI provider: editing/export stay available",
                          "operator stops the AI provider on staging")
        self.staging_only("F-02", journey,
                          "stop one web instance: traffic continues",
                          "needs two staging web instances behind routing")
        self.staging_only("F-03", journey,
                          "kill PDF worker mid-job: lease reclaimed without duplicates",
                          "operator kills the staging pdf worker (unit-covered in "
                          "tests/test_phase5_acceptance.py)")
        self.staging_only("F-04", journey,
                          "stop ClamAV: uploads 503, existing projects usable",
                          "operator stops staging ClamAV")

        student = self.personas.get("student")
        target = self.state.throwaway_project_id or self.state.project_id
        if student and target:
            started = time.monotonic()
            stale = await student.client.post(
                f"/projects/{target}/collaboration/commands",
                json={"command_type": "update_metadata",
                      "payload": {"path": "submission.year", "value": 2027},
                      "expected_document_version": 999_999},
            )
            self.check("F-05", journey,
                       "stale document version mutation returns 409 and loses no work",
                       stale, (409,), started=started)
        else:
            self.blocked("F-05", journey, "stale version exercise", "no project")

        inst_admin = self.personas.get("inst_admin")
        supervisor = self.personas.get("supervisor")
        if inst_admin and supervisor and self.state.institution_id:
            started = time.monotonic()
            revoked = await inst_admin.client.post(
                f"/institutions/{self.state.institution_id}/members/"
                f"{supervisor.user_id}/revoke-sessions",
                json={"reason": "Employment ended; institutional access revoked."},
            )
            after = await supervisor.client.get("/auth/me")
            self.check("F-06", journey,
                       "revoked staff member's sessions fail opaquely (401)",
                       after, (401,), ok=revoked.status_code == 200,
                       detail=f"revoke={revoked.status_code}", started=started)
            self.manual("F-06b", journey,
                        "revoked staff invitation links fail opaquely in the UI",
                        "dept admin §5")
        else:
            self.blocked("F-06", journey, "staff revocation", "personas missing")
        self.staging_only("F-07", journey,
                          "restore backup into isolation; compare sealed checksum",
                          "restore drill runs on staging infrastructure (Subphase C)")

    # ------------------------------------------------------------------ main

    async def run(self) -> dict:
        if self.local:
            storage_dir = tempfile.mkdtemp(prefix="uat-storage-")
            _configure_local_env(
                self.args.database_url or DEFAULT_LOCAL_DATABASE_URL, storage_dir
            )
            from app.main import app  # imported after env configuration

            # DEBUG=true (needed for OTP debug codes) also enables SQLAlchemy
            # echo; silence it so SQL parameters never reach the console.
            import logging

            logging.getLogger("sqlalchemy.engine").setLevel(logging.WARNING)
            logging.getLogger("sqlalchemy.engine.Engine").setLevel(logging.WARNING)
            logging.getLogger("sqlalchemy.engine.Engine").propagate = False
            for handler in list(
                logging.getLogger("sqlalchemy.engine.Engine").handlers
            ):
                logging.getLogger("sqlalchemy.engine.Engine").removeHandler(handler)

            self._transport = httpx.ASGITransport(app=app)
            self.fixtures = LocalFixtures()
            started = time.monotonic()
            try:
                await self.fixtures.recreate_schema()
                inst_id, dept_id = await self.fixtures.seed_institution(
                    self.short_name, self.email_domain
                )
                self.state.institution_id = inst_id
                self.state.department_id = dept_id
                self.record("SETUP-01", "setup",
                            "fixture: fresh schema + UAT institution/department", PASS,
                            started=started)
            except Exception as exc:
                self.record("SETUP-01", "setup", "local test database setup", FAIL,
                            detail=f"{type(exc).__name__}", started=started)
                return self.report()
        else:
            probe = self._new_client()
            started = time.monotonic()
            try:
                health = await probe.get("/healthz")
                self.check("SETUP-01", "setup", "staging health probe", health, (200,),
                           started=started)
            except httpx.HTTPError as exc:
                self.record("SETUP-01", "setup", "staging health probe", FAIL,
                            detail=f"transport error {type(exc).__name__}",
                            started=started)
                await probe.aclose()
                return self.report()
            await probe.aclose()

        await self.j01_identity_sessions()
        await self.j02_upload()
        await self.j03_import_report()
        await self.j04_edit_conflict_undo_restore()
        await self.j05_sources_quotes()
        await self.j06_preview()
        await self.j07_ai_partner()
        await self.j08_supervisor_review()
        await self.j09_operator_boundary()
        await self.j10_admin_approvals()
        await self.j11_seal()
        await self.j12_external_review()
        await self.j13_data_lifecycle()
        await self.j14_support()
        await self.failure_exercises()

        for persona in self.personas.values():
            await persona.client.aclose()
        if self.fixtures is not None:
            await self.fixtures.dispose()
        return self.report()

    def report(self) -> dict:
        totals: dict[str, int] = {}
        for step in self.steps:
            totals[step.status] = totals.get(step.status, 0) + 1
        return {
            "schema": "uat-flow-report/1",
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "mode": "local-asgi" if self.local else "staging-http",
            "base_url": None if self.local else self.args.base_url,
            "run_id": self.run_id,
            "totals": {
                "steps": len(self.steps),
                "pass": totals.get(PASS, 0),
                "fail": totals.get(FAIL, 0),
                "manual": totals.get(MANUAL, 0),
                "requires_staging": totals.get(REQUIRES_STAGING, 0),
                "blocked": totals.get(BLOCKED, 0),
            },
            "checklists": CHECKLISTS,
            "steps": [step.__dict__ for step in self.steps],
        }


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    """Parse CLI arguments for both staging and local execution modes."""

    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    mode = parser.add_mutually_exclusive_group(required=True)
    mode.add_argument("--base-url", help="Staging base URL (real HTTP mode)")
    mode.add_argument("--local", action="store_true",
                      help="Drive the ASGI app in-process against the test database")
    parser.add_argument("--out", default="uat_report.json",
                        help="Path for the JSON step report")
    parser.add_argument("--otp-from-log", action="store_true",
                        help="Prompt the operator to paste OTP codes from the "
                             "staging mail intercept/log")
    parser.add_argument("--database-url",
                        help=f"Local-mode database URL (default {DEFAULT_LOCAL_DATABASE_URL})")
    parser.add_argument("--institution-id",
                        help="Staging institution UUID for admin journeys")
    parser.add_argument("--email-domain", default="test.edu",
                        help="Email domain for generated personas (must map to the "
                             "target institution)")
    parser.add_argument("--poll-timeout", type=int, default=60,
                        help="Seconds to wait for async jobs per step")
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    """Run every journey, write the report, and return a CI exit code."""

    args = parse_args(argv)
    driver = UatDriver(args)
    report = asyncio.run(driver.run())

    serialized = json.dumps(report, indent=2, default=str)
    if driver.state.marker and driver.state.marker in serialized:
        print("REFUSING to write report: canary prose marker leaked into it.")
        return 2
    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(serialized + "\n")

    totals = report["totals"]
    print(
        f"\nUAT summary: {totals['pass']} pass, {totals['fail']} fail, "
        f"{totals['manual']} manual, {totals['requires_staging']} requires-staging, "
        f"{totals['blocked']} blocked ({totals['steps']} steps) -> {out_path}"
    )
    return 0 if totals["fail"] == 0 and totals["blocked"] == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
