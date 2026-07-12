"""Generate golden formatting artifacts for every registered formatting profile.

Subphase E (institution profile sign-off tooling). For each profile registered
in ``app.renderers.phase1_profiles`` this script:

1. builds one deterministic golden sample thesis as a canonical
   ``ThesisDocument`` (fixed uuid5 identifiers, synthetic original text) that
   exercises title page, certificate, declaration, acknowledgement,
   AI disclosure, contents, two chapters (headings, paragraph, block quote,
   verse quote) and a Works Cited registry covering book/journal/web kinds;
2. renders the DOCX through the production renderer with that profile;
3. extracts a FORMATTING FINGERPRINT by re-reading the produced DOCX with
   python-docx (margins, fonts and sizes, line spacing, first-line indents,
   page-number configuration, TOC field presence, works-cited hanging indent);
4. computes SHA-256 checksums (raw artifact and a zip-timestamp-independent
   content digest) and writes goldens + ``fingerprint.json`` to
   ``var/profile-goldens/<profile>/<version>/``;
5. attempts PDF conversion only when LibreOffice ``soffice`` is available,
   otherwise records ``pdf: blocked (LibreOffice not installed)``.

The fingerprints feed docs/release/INSTITUTION_PROFILE_SIGNOFF.md ("Golden
artifacts" section). Generating a golden is NOT institutional certification:
sign-off still requires the official source guide and an authorised approver.

Output policy: this script prints identifiers, counts and checksums only —
never document prose.

Usage::

    .venv-validate/bin/python scripts/generate_profile_golden.py
    .venv-validate/bin/python scripts/generate_profile_golden.py --profiles mla_strict
"""

from __future__ import annotations

import argparse
import hashlib
import json
import subprocess
import sys
import zipfile
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from uuid import UUID, uuid5

REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT))

import docx  # noqa: E402  (third-party; needs no repo path but grouped after bootstrap)
from docx import Document as DocxDocument  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

from app.canonical.model import (  # noqa: E402
    AiDisclosure,
    BlockQuoteBlock,
    CandidateMeta,
    ChapterDoc,
    CollegeMeta,
    FrontMatterEntry,
    HeadingBlock,
    ParagraphBlock,
    PersonMeta,
    Run,
    SubmissionMeta,
    ThesisDocument,
    ThesisMeta,
    VerseQuoteBlock,
    WorksCitedRef,
)
from app.renderers.docx_renderer import render_docx  # noqa: E402
from app.renderers.fonts import soffice_path  # noqa: E402
from app.renderers.pdf_renderer import PdfConversionError, convert_to_pdf  # noqa: E402
from app.renderers.phase1_profiles import PROFILE_LABELS, resolve_phase1_profile  # noqa: E402

# Deterministic namespace for every UUID in the golden fixture.
_GOLDEN_NS = uuid5(UUID("6ba7b811-9dad-11d1-80b4-00c04fd430c8"), "thesis-studio:profile-golden")

GOLDENS_ROOT = REPO_ROOT / "var" / "profile-goldens"


@dataclass
class GoldenSource:
    """Minimal SourceLike (kind + fields) for the works-cited renderer."""

    kind: str
    fields: dict[str, Any] = field(default_factory=dict)


def _gid(name: str) -> UUID:
    """Return a stable uuid5 for a named golden fixture object."""
    return uuid5(_GOLDEN_NS, name)


def _para(name: str, *runs: Run) -> ParagraphBlock:
    """Build a deterministic ParagraphBlock with a fixed id."""
    return ParagraphBlock(id=_gid(name), runs=list(runs))


def build_golden_document() -> tuple[ThesisDocument, dict[UUID, GoldenSource]]:
    """Build the deterministic golden sample thesis and its source registry.

    All text is original synthetic placeholder content; all UUIDs are uuid5
    values derived from fixed names so repeated runs produce byte-identical
    canonical JSON.
    """
    meta = ThesisMeta(
        doc_type="ma_dissertation",
        title="Golden Sample: Memory and Landscape in the Placeholder Novel",
        candidate=CandidateMeta(name="Golden Candidate", reg_no="GLD-0001"),
        degree="Master of Arts in English",
        department="Department of English",
        college=CollegeMeta(
            name="Golden Test College",
            affiliation="University of Placeholder",
            city="Chennai",
            pin="600001",
        ),
        guide=PersonMeta(name="Dr. Golden Guide", designation="Associate Professor"),
        hod=PersonMeta(name="Dr. Golden Head", designation="Head of the Department"),
        submission=SubmissionMeta(month="April", year=2026),
        ai_disclosure=AiDisclosure(
            enabled=True,
            text=(
                "The candidate used an AI writing assistant for grammar review "
                "and outline discussion. All arguments, close readings, and "
                "final prose are the candidate's own work."
            ),
            tools=["Thesis Studio coaching assistant"],
            assistance_types=["grammar review", "outline discussion"],
        ),
    )

    front_matter = [
        FrontMatterEntry(id=_gid("fm:title_page"), kind="title_page", status="approved"),
        FrontMatterEntry(id=_gid("fm:certificate"), kind="certificate", status="approved"),
        FrontMatterEntry(id=_gid("fm:declaration"), kind="declaration", status="approved"),
        FrontMatterEntry(
            id=_gid("fm:acknowledgement"),
            kind="acknowledgement",
            status="approved",
            body_blocks=[
                _para(
                    "fm:ack:p1",
                    Run(
                        text=(
                            "I thank my research supervisor and the Department of "
                            "English for their patient guidance during this golden "
                            "sample dissertation."
                        )
                    ),
                )
            ],
        ),
        FrontMatterEntry(
            id=_gid("fm:ai_disclosure"),
            kind="ai_disclosure",
            status="approved",
            body_blocks=[_para("fm:ai:p1", Run(text=meta.ai_disclosure.text))],
        ),
        FrontMatterEntry(id=_gid("fm:contents"), kind="contents", status="approved"),
    ]

    chapter_one = ChapterDoc(
        id=_gid("ch:1"),
        number=1,
        title="Introduction",
        status="approved",
        blocks=[
            _para(
                "ch1:p1",
                Run(text="This golden chapter opens the sample dissertation and cites "),
                Run(text="The Placeholder Novel", italic=True),
                Run(text=" to exercise the italic run path of every renderer."),
            ),
            HeadingBlock(id=_gid("ch1:h2"), level=2, text="Scope of the Study"),
            _para(
                "ch1:p2",
                Run(
                    text=(
                        "The paragraph below the section heading confirms that body "
                        "text keeps its first-line indent, line spacing, and font "
                        "after a level-two heading."
                    )
                ),
            ),
            BlockQuoteBlock(
                id=_gid("ch1:bq"),
                text=(
                    "A long prose quotation of more than four rendered lines is set "
                    "off from the body text as an indented block without quotation "
                    "marks, so this synthetic passage is deliberately padded until "
                    "it comfortably exceeds the block threshold configured by the "
                    "active formatting profile."
                ),
                citation="Placeholder 112",
            ),
        ],
    )

    chapter_two = ChapterDoc(
        id=_gid("ch:2"),
        number=2,
        title="Landscape as Memory",
        status="approved",
        blocks=[
            _para(
                "ch2:p1",
                Run(
                    text=(
                        "The second golden chapter verifies chapter breaks, the "
                        "chapter label format, and verse quotation handling."
                    )
                ),
            ),
            HeadingBlock(id=_gid("ch2:h3"), level=3, text="A Sub-Section in Italics"),
            VerseQuoteBlock(
                id=_gid("ch2:verse"),
                lines=[
                    "Synthetic verse line one stands here,",
                    "a second placeholder line appears,",
                    "and a third completes the golden strophe.",
                ],
                citation="Sample, lines 1-3",
            ),
            _para(
                "ch2:p2",
                Run(text="A closing paragraph returns to normal body formatting."),
            ),
        ],
    )

    sources: dict[UUID, GoldenSource] = {
        _gid("src:book"): GoldenSource(
            kind="book",
            fields={
                "author": "Placeholder, Anita",
                "title": "The Placeholder Novel",
                "publisher": "Golden House Press",
                "year": "2019",
            },
        ),
        _gid("src:journal"): GoldenSource(
            kind="journal",
            fields={
                "author": "Sample, Ravi",
                "title": "Reading Memory in Synthetic Landscapes",
                "container": "Journal of Golden Studies",
                "volume": "12",
                "number": "3",
                "year": "2021",
                "pages": "45-62",
            },
        ),
        _gid("src:web"): GoldenSource(
            kind="web",
            fields={
                "author": "Fixture, Devi",
                "title": "Placeholder Criticism Online",
                "site": "Golden Review",
                "url": "https://example.org/golden-review/placeholder-criticism",
                "pub_date": "12 Jan. 2022",
                "access_date": "1 Mar. 2026",
            },
        ),
    }

    document = ThesisDocument(
        meta=meta,
        front_matter=front_matter,
        chapters=[chapter_one, chapter_two],
        works_cited=[WorksCitedRef(source_id=sid) for sid in sorted(sources, key=str)],
    )
    return document, sources


# ---------------------------------------------------------------------------
# Checksums
# ---------------------------------------------------------------------------


def sha256_file(path: Path) -> str:
    """Return the SHA-256 hex digest of *path*'s raw bytes."""
    h = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def zip_content_sha256(path: Path) -> str:
    """Return a zip-timestamp-independent SHA-256 over sorted member contents.

    DOCX is a zip archive whose member headers embed the local mtime, so the
    raw file digest changes on every run even when the document content is
    identical. This digest hashes (name, bytes) pairs in sorted order and is
    therefore stable across re-renders of the same content.
    """
    h = hashlib.sha256()
    with zipfile.ZipFile(path) as zf:
        for name in sorted(zf.namelist()):
            h.update(name.encode("utf-8"))
            h.update(b"\x00")
            h.update(zf.read(name))
    return h.hexdigest()


def git_commit() -> str:
    """Return the current repo HEAD commit hash, or 'unknown'."""
    try:
        out = subprocess.run(
            ["git", "-C", str(REPO_ROOT), "rev-parse", "HEAD"],
            capture_output=True,
            text=True,
            timeout=10,
        )
    except OSError:
        return "unknown"
    return out.stdout.strip() if out.returncode == 0 else "unknown"


# ---------------------------------------------------------------------------
# Fingerprint extraction (reads the produced DOCX back with python-docx)
# ---------------------------------------------------------------------------


def _classify_page_size(width_in: float, height_in: float) -> str:
    """Map physical page dimensions back to a named paper size."""
    if abs(width_in - 8.5) < 0.02 and abs(height_in - 11.0) < 0.02:
        return "Letter"
    if abs(width_in - 8.268) < 0.02 and abs(height_in - 11.693) < 0.02:
        return "A4"
    return f"custom({width_in:.3f}x{height_in:.3f}in)"


def _para_has_page_field(paragraph: Any) -> bool:
    """True when the paragraph carries a ``w:fldSimple`` PAGE field."""
    for fld in paragraph._p.findall(qn("w:fldSimple")):
        if "PAGE" in (fld.get(qn("w:instr")) or ""):
            return True
    return False


def _part_page_field(part: Any) -> bool:
    """True when any paragraph of a header/footer part has a PAGE field."""
    return any(_para_has_page_field(p) for p in part.paragraphs)


def _style_fingerprint(document: Any, name: str) -> dict[str, Any] | None:
    """Extract font/size/spacing/indent facts for one named TS-* style."""
    try:
        style = document.styles[name]
    except KeyError:
        return None
    font = style.font
    pf = style.paragraph_format
    return {
        "font": font.name,
        "size_pt": font.size.pt if font.size is not None else None,
        "bold": font.bold,
        "italic": font.italic,
        "line_spacing": pf.line_spacing,
        "first_line_indent_in": (
            round(pf.first_line_indent.inches, 3) if pf.first_line_indent is not None else None
        ),
        "left_indent_in": (
            round(pf.left_indent.inches, 3) if pf.left_indent is not None else None
        ),
        "alignment": str(pf.alignment) if pf.alignment is not None else None,
    }


def extract_fingerprint(docx_path: Path) -> dict[str, Any]:
    """Read the rendered DOCX and return the formatting fingerprint dict."""
    document = DocxDocument(str(docx_path))

    sections: list[dict[str, Any]] = []
    for section in document.sections:
        sect_pr = section._sectPr
        pg_num = sect_pr.find(qn("w:pgNumType"))
        width_in = round(section.page_width.inches, 3)
        height_in = round(section.page_height.inches, 3)
        sections.append(
            {
                "page_size": _classify_page_size(width_in, height_in),
                "page_width_in": width_in,
                "page_height_in": height_in,
                "margins_in": {
                    "top": round(section.top_margin.inches, 3),
                    "bottom": round(section.bottom_margin.inches, 3),
                    "left": round(section.left_margin.inches, 3),
                    "right": round(section.right_margin.inches, 3),
                },
                "page_number_format": pg_num.get(qn("w:fmt")) if pg_num is not None else None,
                "page_number_start": pg_num.get(qn("w:start")) if pg_num is not None else None,
                "different_first_page": bool(section.different_first_page_header_footer),
                "footer_page_field": _part_page_field(section.footer),
                "header_page_field": _part_page_field(section.header),
            }
        )

    toc_field = False
    for fld in document.element.body.iter(qn("w:fldSimple")):
        if "TOC" in (fld.get(qn("w:instr")) or ""):
            toc_field = True
            break

    styles = {
        name: _style_fingerprint(document, name)
        for name in (
            "TS-Normal",
            "TS-FrontCenter",
            "TS-ChapterLabel",
            "TS-ChapterTitle",
            "TS-Heading2",
            "TS-Heading3",
            "TS-BlockQuote",
            "TS-Verse",
            "TS-WorksCited",
        )
    }

    wc_style = styles.get("TS-WorksCited") or {}
    hanging = None
    if wc_style.get("left_indent_in") is not None and wc_style.get("first_line_indent_in") is not None:
        # A hanging indent renders as left_indent=+X with first_line_indent=-X.
        hanging = {
            "left_indent_in": wc_style["left_indent_in"],
            "first_line_indent_in": wc_style["first_line_indent_in"],
            "is_hanging": wc_style["first_line_indent_in"] < 0
            and abs(wc_style["left_indent_in"] + wc_style["first_line_indent_in"]) < 0.01,
        }

    return {
        "sections": sections,
        "styles": styles,
        "toc_native_word_field_present": toc_field,
        "works_cited_hanging_indent": hanging,
    }


# ---------------------------------------------------------------------------
# Per-profile generation
# ---------------------------------------------------------------------------


def version_dir_name(version_label: str) -> str:
    """Directory-safe version segment from a governed label like 'name:v1'."""
    return version_label.split(":", 1)[1] if ":" in version_label else version_label


def ensure_var_ignored() -> None:
    """Add var/.gitignore only when the repo .gitignore does not ignore var/."""
    root_gitignore = REPO_ROOT / ".gitignore"
    if root_gitignore.exists():
        lines = {line.strip() for line in root_gitignore.read_text().splitlines()}
        if "var/" in lines or "var" in lines or "/var/" in lines:
            return
    var_ignore = REPO_ROOT / "var" / ".gitignore"
    if not var_ignore.exists():
        var_ignore.parent.mkdir(parents=True, exist_ok=True)
        var_ignore.write_text("*\n")


def generate_for_profile(profile_name: str) -> dict[str, Any]:
    """Render the golden artifacts for one profile; return the summary record."""
    profile, version_label = resolve_phase1_profile(profile_name)
    out_dir = GOLDENS_ROOT / profile_name / version_dir_name(version_label)
    out_dir.mkdir(parents=True, exist_ok=True)

    document, sources = build_golden_document()

    canonical_json = json.dumps(
        document.model_dump(mode="json"), sort_keys=True, indent=2, ensure_ascii=False
    )
    canonical_path = out_dir / "canonical.json"
    canonical_path.write_text(canonical_json, encoding="utf-8")
    canonical_sha = hashlib.sha256(canonical_json.encode("utf-8")).hexdigest()

    docx_path = out_dir / "golden.docx"
    render_docx(document, dict(sources), profile, str(docx_path))

    fingerprint = extract_fingerprint(docx_path)

    pdf_record: Any
    if soffice_path() is None:
        pdf_record = "blocked (LibreOffice not installed)"
    else:
        try:
            pdf_path = convert_to_pdf(str(docx_path), str(out_dir))
            pdf_record = {
                "file": Path(pdf_path).name,
                "sha256": sha256_file(Path(pdf_path)),
            }
        except (PdfConversionError, OSError) as exc:
            pdf_record = f"failed ({type(exc).__name__})"

    record: dict[str, Any] = {
        "profile": profile_name,
        "profile_label": PROFILE_LABELS.get(profile_name, ""),
        "profile_version": version_label,
        "renderer_git_commit": git_commit(),
        "python_docx_version": docx.__version__,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "canonical_sha256": canonical_sha,
        "docx_sha256": sha256_file(docx_path),
        "docx_content_sha256": zip_content_sha256(docx_path),
        "pdf": pdf_record,
        "profile_notes": profile.notes,
        "formatting": fingerprint,
        "counts": {
            "front_matter_entries": len(document.front_matter),
            "chapters": len(document.chapters),
            "works_cited_sources": len(sources),
        },
        "certification": (
            "NOT CERTIFIED — golden fingerprint only; institutional sign-off "
            "requires the official source guide and an authorised approver "
            "(docs/release/INSTITUTION_PROFILE_SIGNOFF.md)."
        ),
    }
    (out_dir / "fingerprint.json").write_text(
        json.dumps(record, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
    )
    return record


def _print_summary(record: dict[str, Any]) -> None:
    """Print the identifier/checksum summary for one profile (no prose)."""
    fmt = record["formatting"]
    body_section = fmt["sections"][-1]
    normal = fmt["styles"].get("TS-Normal") or {}
    print(f"profile:            {record['profile']}  [{record['profile_version']}]")
    print(f"  canonical sha256: {record['canonical_sha256']}")
    print(f"  docx sha256:      {record['docx_sha256']}")
    print(f"  docx content sha: {record['docx_content_sha256']}")
    print(f"  pdf:              {record['pdf']}")
    print(
        "  page:             "
        f"{body_section['page_size']} margins(in) T{body_section['margins_in']['top']} "
        f"B{body_section['margins_in']['bottom']} L{body_section['margins_in']['left']} "
        f"R{body_section['margins_in']['right']}"
    )
    print(
        "  body type:        "
        f"{normal.get('font')} {normal.get('size_pt')}pt spacing {normal.get('line_spacing')} "
        f"first-line indent {normal.get('first_line_indent_in')}in"
    )
    print(
        "  pagination:       "
        + " | ".join(
            f"s{i}: fmt={s['page_number_format']} start={s['page_number_start']} "
            f"footer_field={s['footer_page_field']} header_field={s['header_page_field']}"
            for i, s in enumerate(fmt["sections"])
        )
    )
    print(f"  toc native field: {fmt['toc_native_word_field_present']}")
    print(f"  wc hanging:       {fmt['works_cited_hanging_indent']}")
    print()


def main(argv: list[str] | None = None) -> int:
    """CLI entry point; returns a process exit code."""
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument(
        "--profiles",
        nargs="*",
        default=sorted(PROFILE_LABELS),
        help="Subset of registered profiles (default: all registered profiles).",
    )
    args = parser.parse_args(argv)

    unknown = [p for p in args.profiles if p not in PROFILE_LABELS]
    if unknown:
        print(f"ERROR: unknown profile(s): {', '.join(unknown)}")
        print(f"Registered profiles: {', '.join(sorted(PROFILE_LABELS))}")
        return 2

    ensure_var_ignored()

    failures = 0
    print(f"Golden output root: {GOLDENS_ROOT}")
    print()
    for name in args.profiles:
        try:
            record = generate_for_profile(name)
        except Exception as exc:  # a broken profile must not hide the others
            failures += 1
            print(f"profile:            {name}")
            print(f"  FAILED:           {type(exc).__name__}: {exc}")
            print()
            continue
        _print_summary(record)

    print(
        "NOTE: goldens are engineering evidence only. Certification is blocked "
        "until an official institutional guide and an authorised approver are "
        "recorded in docs/release/INSTITUTION_PROFILE_SIGNOFF.md."
    )
    return 1 if failures else 0


if __name__ == "__main__":
    raise SystemExit(main())
