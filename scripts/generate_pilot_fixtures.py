"""Generate synthetic, anonymised DOCX fixtures for the real-manuscript pilot harness.

Every fixture is built from deterministic synthetic prose ("Synthetic paragraph N
about postcolonial narrative form.") — never real thesis text. The set covers the
corpus requirements in docs/release/REAL_MANUSCRIPT_PILOT.md that can be exercised
synthetically: tables, embedded images, OMML equations, comments, tracked changes,
footnotes/endnotes, verse and block quotations, broken/duplicated headings, a
complex Works Cited, non-English (Tamil/Hindi) paragraphs, a large ~150-page
document, a malformed (truncated) zip and a zip-bomb-like package.

Usage:
    python scripts/generate_pilot_fixtures.py [--out tests/fixtures/pilot]

Only fixture file names, sizes and checksums are printed — never fixture text.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import struct
import zipfile
import zlib
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.shared import Inches

REPO_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUT = REPO_ROOT / "tests" / "fixtures" / "pilot"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

_SYN = "Synthetic paragraph {n} about postcolonial narrative form"
_FILLER = (
    " It rehearses a synthetic argument about narrative form, memory, and the"
    " archive without naming any real writer, student, or institution."
)
_FIXED_DATE = (1980, 1, 1, 0, 0, 0)


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------


def _sha256(path: Path) -> str:
    """Return the SHA-256 hex digest of *path*."""
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1 << 20), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _write_zip(path: Path, entries: dict[str, bytes]) -> None:
    """Write *entries* as a deflated zip with fixed timestamps (deterministic)."""
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, blob in entries.items():
            info = zipfile.ZipInfo(name, date_time=_FIXED_DATE)
            info.compress_type = zipfile.ZIP_DEFLATED
            zf.writestr(info, blob)


def _read_zip(source: Path | BytesIO) -> dict[str, bytes]:
    """Read every member of a zip into an ordered name → bytes mapping."""
    with zipfile.ZipFile(source) as zf:
        return {info.filename: zf.read(info.filename) for info in zf.infolist()}


def _normalise_zip(path: Path) -> None:
    """Rewrite *path* in place with fixed entry timestamps."""
    _write_zip(path, _read_zip(path))


def _save(doc: Document, path: Path) -> None:
    """Save *doc* to *path* and normalise zip metadata for determinism."""
    doc.save(str(path))
    _normalise_zip(path)


def _inject_parts(
    path: Path,
    parts: dict[str, bytes],
    overrides: dict[str, str],
    rels: list[tuple[str, str, str]],
) -> None:
    """Add OPC *parts* (with content-type *overrides* and document *rels*) to a docx."""
    entries = _read_zip(path)
    content_types = entries["[Content_Types].xml"].decode("utf-8")
    additions = "".join(
        f'<Override PartName="/{name}" ContentType="{ctype}"/>'
        for name, ctype in overrides.items()
    )
    entries["[Content_Types].xml"] = content_types.replace(
        "</Types>", additions + "</Types>"
    ).encode("utf-8")
    rel_xml = entries["word/_rels/document.xml.rels"].decode("utf-8")
    rel_additions = "".join(
        f'<Relationship Id="{rid}" Type="{rtype}" Target="{target}"/>'
        for rid, rtype, target in rels
    )
    entries["word/_rels/document.xml.rels"] = rel_xml.replace(
        "</Relationships>", rel_additions + "</Relationships>"
    ).encode("utf-8")
    entries.update(parts)
    _write_zip(path, entries)


def _append_raw_paragraph(doc: Document, xml: str) -> None:
    """Insert a raw WordprocessingML paragraph before the body sectPr."""
    doc.element.body.get_or_add_sectPr().addprevious(parse_xml(xml))


def _png_1px() -> bytes:
    """Return a deterministic 1x1 red-pixel PNG."""

    def chunk(tag: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + tag
            + data
            + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)
        )

    signature = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00\xff\x00\x00", 9)
    return signature + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


# ---------------------------------------------------------------------------
# Synthetic content helpers
# ---------------------------------------------------------------------------


def _syn(n: int, citation: str = "") -> str:
    """Return synthetic paragraph *n*, optionally ending with an in-text citation."""
    base = _SYN.format(n=n)
    return f"{base} ({citation})." if citation else f"{base}."


def _centered(doc: Document, text: str, bold: bool = False) -> None:
    """Add a centered paragraph (title-page / heading style content)."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = bold or None


def _title_page(doc: Document) -> None:
    """Add two synthetic title-page paragraphs (pre-chapter front matter)."""
    _centered(doc, "A SYNTHETIC STUDY OF NARRATIVE FORM", bold=True)
    _centered(doc, "A synthetic pilot fixture for the Robofox ingestion harness")


def _front_matter(doc: Document) -> None:
    """Add DECLARATION / ACKNOWLEDGEMENT / CONTENTS front-matter sections."""
    _centered(doc, "DECLARATION", bold=True)
    doc.add_paragraph(
        "I hereby declare that this synthetic fixture contains no real"
        " manuscript text and exists only to exercise the ingestion pipeline."
    )
    _centered(doc, "ACKNOWLEDGEMENT", bold=True)
    doc.add_paragraph(
        "The author of this synthetic fixture acknowledges the deterministic"
        " generator that produced it."
    )
    _centered(doc, "CONTENTS", bold=True)
    doc.add_paragraph("Chapter listing omitted in this synthetic fixture.")


def _chapter(doc: Document, label: str) -> None:
    """Add a chapter heading paragraph, e.g. 'CHAPTER I: THE SYNTHETIC ARCHIVE'."""
    doc.add_paragraph(label)


def _body(doc: Document, count: int, start: int = 1, cite_every: int = 0,
          surname: str = "Iyer") -> int:
    """Add *count* synthetic body paragraphs; returns the next paragraph number."""
    n = start
    for _ in range(count):
        citation = f"{surname} 12" if cite_every and n % cite_every == 0 else ""
        doc.add_paragraph(_syn(n, citation))
        n += 1
    return n


def _wc_entry(doc: Document, parts: list[tuple[str, bool]]) -> None:
    """Add one Works Cited entry paragraph from (text, italic) run parts."""
    para = doc.add_paragraph()
    for text, italic in parts:
        run = para.add_run(text)
        if italic:
            run.italic = True


def _simple_wc(doc: Document) -> None:
    """Add a minimal two-entry Works Cited section."""
    doc.add_paragraph("WORKS CITED")
    _wc_entry(doc, [
        ("Iyer, Kamala. ", False),
        ("A Synthetic Companion. ", True),
        ("Imagined House, 2005.", False),
    ])
    _wc_entry(doc, [
        ("Rao, Devika. ", False),
        ("Synthetic Tides. ", True),
        ("Imagined House, 2011.", False),
    ])


def _base_doc(front_matter: bool = False) -> Document:
    """Return a new document with a synthetic title page (and optional front matter)."""
    doc = Document()
    _title_page(doc)
    if front_matter:
        _front_matter(doc)
    return doc


# ---------------------------------------------------------------------------
# Fixture builders
# ---------------------------------------------------------------------------


def build_tables(path: Path) -> None:
    """Tables plus full front matter; table text is expected unsupported content."""
    doc = _base_doc(front_matter=True)
    _chapter(doc, "CHAPTER I: THE SYNTHETIC ARCHIVE")
    n = _body(doc, 4, cite_every=3)
    table = doc.add_table(rows=3, cols=3)
    for r in range(3):
        for c in range(3):
            table.cell(r, c).text = f"Synthetic cell r{r + 1}c{c + 1}"
    _body(doc, 2, start=n)
    _simple_wc(doc)
    _save(doc, path)


def build_embedded_image(path: Path) -> None:
    """An embedded 1x1 PNG figure with a caption paragraph."""
    doc = _base_doc()
    _chapter(doc, "CHAPTER I: THE SYNTHETIC IMAGE")
    n = _body(doc, 3, cite_every=2)
    doc.add_picture(BytesIO(_png_1px()), width=Inches(1.0))
    doc.add_paragraph("Figure 1. A synthetic placeholder figure.")
    _body(doc, 2, start=n)
    _simple_wc(doc)
    _save(doc, path)


def build_equation_omml(path: Path) -> None:
    """An OMML equation-like placeholder (m:oMathPara/m:oMath)."""
    doc = _base_doc()
    _chapter(doc, "CHAPTER I: THE SYNTHETIC EQUATION")
    n = _body(doc, 2)
    _append_raw_paragraph(
        doc,
        f'<w:p xmlns:w="{W_NS}" xmlns:m="{M_NS}">'
        "<m:oMathPara><m:oMath><m:r><m:t>x = y + 1</m:t></m:r></m:oMath></m:oMathPara>"
        "</w:p>",
    )
    _body(doc, 2, start=n)
    _simple_wc(doc)
    _save(doc, path)


def build_comments(path: Path) -> None:
    """Two anchored reviewer comments injected as a proper word/comments.xml part."""
    doc = _base_doc()
    _chapter(doc, "CHAPTER I: THE SYNTHETIC MARGIN")
    n = _body(doc, 2)
    for cid in (1, 2):
        _append_raw_paragraph(
            doc,
            f'<w:p xmlns:w="{W_NS}">'
            f'<w:commentRangeStart w:id="{cid}"/>'
            f"<w:r><w:t>{_syn(n)}</w:t></w:r>"
            f'<w:commentRangeEnd w:id="{cid}"/>'
            f'<w:r><w:commentReference w:id="{cid}"/></w:r>'
            "</w:p>",
        )
        n += 1
    _body(doc, 1, start=n)
    _simple_wc(doc)
    _save(doc, path)
    comments_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<w:comments xmlns:w="{W_NS}">'
        '<w:comment w:id="1" w:author="Synthetic Reviewer" w:initials="SR"'
        ' w:date="2026-01-01T00:00:00Z">'
        "<w:p><w:r><w:t>Synthetic reviewer comment one.</w:t></w:r></w:p></w:comment>"
        '<w:comment w:id="2" w:author="Synthetic Reviewer" w:initials="SR"'
        ' w:date="2026-01-01T00:00:00Z">'
        "<w:p><w:r><w:t>Synthetic reviewer comment two.</w:t></w:r></w:p></w:comment>"
        "</w:comments>"
    ).encode("utf-8")
    _inject_parts(
        path,
        parts={"word/comments.xml": comments_xml},
        overrides={
            "word/comments.xml": (
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.comments+xml"
            )
        },
        rels=[("rId1001", f"{REL_NS}/comments", "comments.xml")],
    )


def build_tracked_changes(path: Path) -> None:
    """Tracked insertions (w:ins) and deletions (w:del) via raw WordprocessingML."""
    doc = _base_doc()
    _chapter(doc, "CHAPTER I: THE SYNTHETIC REVISION")
    n = _body(doc, 2)
    _append_raw_paragraph(
        doc,
        f'<w:p xmlns:w="{W_NS}">'
        f'<w:r><w:t xml:space="preserve">{_syn(n)} </w:t></w:r>'
        '<w:ins w:id="101" w:author="Synthetic Reviewer" w:date="2026-01-01T00:00:00Z">'
        '<w:r><w:t xml:space="preserve">Synthetic inserted clause pending'
        " acceptance.</w:t></w:r></w:ins>"
        '<w:del w:id="102" w:author="Synthetic Reviewer" w:date="2026-01-01T00:00:00Z">'
        '<w:r><w:delText xml:space="preserve"> Synthetic deleted clause pending'
        " rejection.</w:delText></w:r></w:del>"
        "</w:p>",
    )
    n += 1
    _append_raw_paragraph(
        doc,
        f'<w:p xmlns:w="{W_NS}">'
        f'<w:r><w:t xml:space="preserve">{_syn(n)} </w:t></w:r>'
        '<w:del w:id="103" w:author="Synthetic Reviewer" w:date="2026-01-01T00:00:00Z">'
        '<w:r><w:delText xml:space="preserve"> A second synthetic deleted'
        " clause.</w:delText></w:r></w:del>"
        "</w:p>",
    )
    _body(doc, 1, start=n + 1)
    _simple_wc(doc)
    _save(doc, path)


def build_footnotes_endnotes(path: Path) -> None:
    """Footnote and endnote references with injected footnotes/endnotes parts."""
    doc = _base_doc()
    _chapter(doc, "CHAPTER I: THE SYNTHETIC APPARATUS")
    n = _body(doc, 2)
    _append_raw_paragraph(
        doc,
        f'<w:p xmlns:w="{W_NS}">'
        f'<w:r><w:t xml:space="preserve">{_syn(n)} </w:t></w:r>'
        '<w:r><w:footnoteReference w:id="1"/></w:r>'
        '<w:r><w:t xml:space="preserve"> A synthetic clause after the footnote'
        " reference.</w:t></w:r>"
        '<w:r><w:endnoteReference w:id="1"/></w:r>'
        "</w:p>",
    )
    n += 1
    _append_raw_paragraph(
        doc,
        f'<w:p xmlns:w="{W_NS}">'
        f'<w:r><w:t xml:space="preserve">{_syn(n)} </w:t></w:r>'
        '<w:r><w:footnoteReference w:id="2"/></w:r>'
        "</w:p>",
    )
    _body(doc, 1, start=n + 1)
    _simple_wc(doc)
    _save(doc, path)
    footnotes_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<w:footnotes xmlns:w="{W_NS}">'
        '<w:footnote w:type="separator" w:id="-1">'
        "<w:p><w:r><w:separator/></w:r></w:p></w:footnote>"
        '<w:footnote w:type="continuationSeparator" w:id="0">'
        "<w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>"
        '<w:footnote w:id="1"><w:p><w:r>'
        "<w:t>Synthetic footnote text one.</w:t></w:r></w:p></w:footnote>"
        '<w:footnote w:id="2"><w:p><w:r>'
        "<w:t>Synthetic footnote text two.</w:t></w:r></w:p></w:footnote>"
        "</w:footnotes>"
    ).encode("utf-8")
    endnotes_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<w:endnotes xmlns:w="{W_NS}">'
        '<w:endnote w:type="separator" w:id="-1">'
        "<w:p><w:r><w:separator/></w:r></w:p></w:endnote>"
        '<w:endnote w:type="continuationSeparator" w:id="0">'
        "<w:p><w:r><w:continuationSeparator/></w:r></w:p></w:endnote>"
        '<w:endnote w:id="1"><w:p><w:r>'
        "<w:t>Synthetic endnote text one.</w:t></w:r></w:p></w:endnote>"
        "</w:endnotes>"
    ).encode("utf-8")
    _inject_parts(
        path,
        parts={"word/footnotes.xml": footnotes_xml, "word/endnotes.xml": endnotes_xml},
        overrides={
            "word/footnotes.xml": (
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.footnotes+xml"
            ),
            "word/endnotes.xml": (
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.endnotes+xml"
            ),
        },
        rels=[
            ("rId1002", f"{REL_NS}/footnotes", "footnotes.xml"),
            ("rId1003", f"{REL_NS}/endnotes", "endnotes.xml"),
        ],
    )


def build_verse_quotation(path: Path) -> None:
    """An indented multi-line verse quotation with a trailing citation."""
    doc = _base_doc()
    _chapter(doc, "CHAPTER I: THE SYNTHETIC VERSE")
    n = _body(doc, 2)
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.5)
    lines = [
        "Synthetic verse line one of the harbour,",
        "synthetic verse line two of the tide,",
        "synthetic verse line three returns the form",
    ]
    for i, line in enumerate(lines):
        run = para.add_run(line)
        if i < len(lines) - 1:
            run.add_break()
    para.add_run(" (Poet 12-13)")
    _body(doc, 2, start=n)
    _simple_wc(doc)
    _save(doc, path)


def build_block_quotes(path: Path) -> None:
    """One cited and one uncited indented block quotation (uncited → ambiguity)."""
    doc = _base_doc()
    _chapter(doc, "CHAPTER I: THE SYNTHETIC QUOTATION")
    n = _body(doc, 1)
    cited = doc.add_paragraph(
        "This synthetic block quotation runs long enough to be prose rather than"
        " verse, and it rehearses a purely synthetic claim about narrative form"
        " and the archive. (Iyer 88)"
    )
    cited.paragraph_format.left_indent = Inches(0.5)
    n = _body(doc, 1, start=n)
    uncited = doc.add_paragraph(
        "This second synthetic block quotation also runs long enough to be prose,"
        " but it deliberately carries no trailing citation so the parser must"
        " raise a human-review ambiguity."
    )
    uncited.paragraph_format.left_indent = Inches(0.5)
    _body(doc, 1, start=n)
    _simple_wc(doc)
    _save(doc, path)


def build_broken_headings(path: Path) -> None:
    """Broken/duplicated heading styles and mixed Roman/Arabic chapter numbering."""
    doc = _base_doc()
    _chapter(doc, "CHAPTER I: THE FIRST SYNTHETIC CHAPTER")
    n = _body(doc, 2)

    def bold_heading(text: str) -> None:
        para = doc.add_paragraph()
        para.add_run(text).bold = True

    bold_heading("The Weight of Synthetic Names")
    n = _body(doc, 2, start=n)
    bold_heading("The Weight of Synthetic Names")  # duplicated heading
    n = _body(doc, 1, start=n)
    italic_para = doc.add_paragraph()
    italic_para.add_run("reading the synthetic shore").italic = True
    n = _body(doc, 1, start=n)
    doc.add_paragraph("A Question Of Synthetic Salt")  # unstyled title-like line
    n = _body(doc, 1, start=n)
    doc.add_paragraph("Misused Word Heading Style", style="Heading 1")
    n = _body(doc, 1, start=n)
    doc.add_paragraph("CHAPTER II")
    doc.add_paragraph("CHAPTER II")  # duplicated chapter boundary
    _centered(doc, "THE DOUBLED SYNTHETIC HEADING")
    n = _body(doc, 2, start=n)
    doc.add_paragraph("CHAPTER 3: ARABIC NUMBERED SYNTHETIC CHAPTER")
    _body(doc, 2, start=n)
    _simple_wc(doc)
    _save(doc, path)


def build_complex_works_cited(path: Path) -> None:
    """Journal-in-database, translated book, film, and same-author '---.' repeats."""
    doc = _base_doc()
    _chapter(doc, "CHAPTER I: THE SYNTHETIC BIBLIOGRAPHY")
    doc.add_paragraph(_syn(1, "Sundar 45"))          # ambiguous surname (two Sundars)
    doc.add_paragraph(_syn(2, "Kumar, The Invented Shore 12"))  # title-hint resolution
    doc.add_paragraph(_syn(3, "qtd. in Rao 88"))     # secondary quotation, no source
    doc.add_paragraph(_syn(4, "Sundar, Letters to a Synthetic City 7"))
    doc.add_paragraph(_syn(5))
    doc.add_paragraph("WORKS CITED")
    _wc_entry(doc, [
        ("Kumar, Ravi. ", False),
        ("The Invented Shore. ", True),
        ("Translated by Anita Rao, Synthetic Press, 2001.", False),
    ])
    _wc_entry(doc, [
        ("The Paper Harbour. ", True),
        ("Directed by Sona Illustrator, Synthetic Studios, 2010.", False),
    ])
    _wc_entry(doc, [
        ("Sundar, Anil. ", False),
        ("Coastal Grammars. ", True),
        ("Imagined House, 2003.", False),
    ])
    _wc_entry(doc, [
        ("Sundar, Meera. ", False),
        ('"Archive Fever in the Postcolonial Novel." ', False),
        ("Journal of Synthetic Letters", True),
        (", vol. 12, no. 3, 2019, pp. 45-60. ", False),
        ("SyntheticDB", True),
        (", doi.org/10.0000/synthetic.12345.", False),
    ])
    _wc_entry(doc, [
        ("---. ", False),
        ("Letters to a Synthetic City. ", True),
        ("Imagined House, 2015.", False),
    ])
    _wc_entry(doc, [
        ("---. ", False),
        ("The Second Synthetic City. ", True),
        ("Imagined House, 2018.", False),
    ])
    _save(doc, path)


def build_non_english(path: Path) -> None:
    """Tamil and Hindi unicode paragraphs alongside English synthetic prose."""
    doc = _base_doc()
    _chapter(doc, "CHAPTER I: THE SYNTHETIC POLYGLOT")
    n = _body(doc, 2)
    tamil = [
        "செயற்கைப் பத்தி ஒன்று: பின்காலனித்துவக் கதை வடிவம் குறித்த செயற்கை உரை.",
        "செயற்கைப் பத்தி இரண்டு: நினைவும் ஆவணமும் பற்றிய செயற்கை வாக்கியம்.",
        "செயற்கைப் பத்தி மூன்று: எந்த உண்மையான ஆய்வுரையும் இதில் இல்லை.",
    ]
    hindi = [
        "कृत्रिम अनुच्छेद एक: उपनिवेशोत्तर कथा रूप पर एक कृत्रिम पाठ।",
        "कृत्रिम अनुच्छेद दो: स्मृति और अभिलेख पर एक कृत्रिम वाक्य।",
        "कृत्रिम अनुच्छेद तीन: इसमें कोई वास्तविक शोध-पाठ नहीं है।",
    ]
    for line in tamil + hindi:
        doc.add_paragraph(line)
    _body(doc, 1, start=n)
    _simple_wc(doc)
    _save(doc, path)


def build_large_150p(path: Path) -> None:
    """A large ~150-page document generated by deterministic repetition."""
    doc = _base_doc(front_matter=True)
    romans = ["I", "II", "III", "IV", "V", "VI"]
    titles = [
        "THE SYNTHETIC ARCHIVE", "THE SYNTHETIC SHORE", "THE SYNTHETIC CITY",
        "THE SYNTHETIC TIDE", "THE SYNTHETIC MARGIN", "THE SYNTHETIC RETURN",
    ]
    n = 1
    for roman, title in zip(romans, titles):
        _chapter(doc, f"CHAPTER {roman}: {title}")
        for _ in range(184):
            citation = "Iyer 12" if n % 25 == 0 else ""
            doc.add_paragraph(_syn(n, citation) + _FILLER * 2)
            n += 1
    _simple_wc(doc)
    _save(doc, path)


def _tiny_doc_bytes() -> bytes:
    """Return the bytes of a minimal valid synthetic docx (in memory)."""
    doc = Document()
    _title_page(doc)
    _chapter(doc, "CHAPTER I: THE SYNTHETIC STUB")
    _body(doc, 2)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_malformed_zip(path: Path) -> None:
    """A truncated docx whose zip central directory is missing (must be rejected)."""
    blob = _tiny_doc_bytes()
    path.write_bytes(blob[: len(blob) // 2])


def build_zip_bomb(path: Path) -> None:
    """A zip-bomb-like docx: ~60MB document.xml of spaces compressing tiny."""
    entries = _read_zip(BytesIO(_tiny_doc_bytes()))
    document = entries["word/document.xml"].decode("utf-8")
    filler = " " * (60 * 1024 * 1024)
    bomb_paragraph = (
        '<w:p><w:r><w:t xml:space="preserve">' + filler + "</w:t></w:r></w:p>"
    )
    entries["word/document.xml"] = document.replace(
        "<w:body>", "<w:body>" + bomb_paragraph, 1
    ).encode("utf-8")
    _write_zip(path, entries)


# ---------------------------------------------------------------------------
# Registry and entry point
# ---------------------------------------------------------------------------

FIXTURES: list[tuple[str, object, bool, list[str]]] = [
    ("tables", build_tables, False, ["tables"]),
    ("embedded_image", build_embedded_image, False, ["images"]),
    ("equation_omml", build_equation_omml, False, ["equations"]),
    ("comments", build_comments, False, []),
    ("tracked_changes", build_tracked_changes, False, []),
    ("footnotes_endnotes", build_footnotes_endnotes, False, ["footnotes", "endnotes"]),
    ("verse_quotation", build_verse_quotation, False, []),
    ("block_quotes", build_block_quotes, False, []),
    ("broken_headings", build_broken_headings, False, []),
    ("complex_works_cited", build_complex_works_cited, False, []),
    ("non_english", build_non_english, False, []),
    ("large_150p", build_large_150p, False, []),
    ("malformed_zip", build_malformed_zip, True, []),
    ("zip_bomb", build_zip_bomb, True, []),
]


def main() -> int:
    """Generate every fixture plus a manifest; print names, sizes and checksums."""
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT,
                        help="Output directory for the .docx fixtures")
    args = parser.parse_args()
    out_dir: Path = args.out
    out_dir.mkdir(parents=True, exist_ok=True)
    gitignore = out_dir / ".gitignore"
    if not gitignore.exists():
        gitignore.write_text("*.docx\n", encoding="utf-8")

    manifest: dict[str, object] = {
        "generated_by": "scripts/generate_pilot_fixtures.py",
        "content": "deterministic synthetic prose only — never real thesis text",
        "fixtures": [],
    }
    for case, builder, expect_reject, expected_unsupported in FIXTURES:
        fixture_path = out_dir / f"{case}.docx"
        builder(fixture_path)  # type: ignore[operator]
        record = {
            "file": fixture_path.name,
            "case": case,
            "expect_preflight_reject": expect_reject,
            "expected_unsupported": expected_unsupported,
            "size_bytes": fixture_path.stat().st_size,
            "sha256": _sha256(fixture_path),
        }
        manifest["fixtures"].append(record)  # type: ignore[union-attr]
        print(f"{fixture_path.name}: {record['size_bytes']} bytes"
              f" sha256={record['sha256'][:16]}…"
              f"{' (expect preflight reject)' if expect_reject else ''}")

    manifest_path = out_dir / "manifest.json"
    manifest_path.write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")
    print(f"wrote {len(FIXTURES)} fixtures + manifest to {out_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
