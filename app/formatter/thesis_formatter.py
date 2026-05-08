"""
Thesis Formatter — produces MCC/UoM-format MA dissertations as .docx files.

Drop-in module for the Robofox Thesis Studio FastAPI backend.

Usage:
    from thesis_formatter import render_thesis_docx, ThesisInput

    thesis = ThesisInput(
        front_matter=FrontMatter(...),
        abstract="...",
        keywords=[...],
        acknowledgement="...",
        chapters=[...],
        works_cited=[...],
    )
    render_thesis_docx(thesis, logo_path="/path/to/logo.png", output_path="thesis.docx")

Format spec (verified against MCC English Dept reference thesis):
- Page: 8.5"x11", margins 1.5"L / 1"T-R-B (binding-friendly)
- Body: Times New Roman 12pt, 1.5 line spacing, justified, first-line indent
- Headings: Chapter 14pt bold centered ALL CAPS; sections 12pt bold left-aligned ALL CAPS
- Page header: <LastName> <PageNum> top-right, body pages only (front matter unnumbered)
- Citation style: MLA 9 in-text, hanging-indent works cited
- TOC: auto-generated Word field (updates on open)
"""

from __future__ import annotations

from dataclasses import dataclass, field
from typing import Optional
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

from app.formatter.inline_markdown import parse_inline_runs


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------

@dataclass
class FrontMatter:
    """All institution-specific and student-specific data for the front matter."""

    # Institution
    college_name: str                # "Madras Christian College (Autonomous)"
    college_address: str             # "Tambaram, Chennai – 600 059."
    university_name: str             # "University of Madras"
    department_name: str             # "PG & Research Department of English"
    department_aided: bool = False   # whether to add "(Aided)" in signature block
    short_address: str = ""          # short address for sig block, e.g. "Tambaram, Chennai – 59"

    # Degree
    degree_program_line_1: str = "MASTER OF ARTS"
    degree_program_line_2: str = "IN ENGLISH LANGUAGE AND LITERATURE"

    # Candidate
    student_full_name: str = ""      # "R. Chris Immanuel"
    register_number: str = ""        # "2301712006044"

    # Supervisor
    supervisor_full_name: str = ""   # "Prof. J. Arun Kumar"
    supervisor_designation: str = "" # "Assistant Professor"

    # HOD (for certificate signature block)
    hod_full_name: str = ""          # "Dr. S. Franklin Daniel"
    hod_designation: str = "Head of the Department"

    # Dates / period
    study_period: str = ""           # "2023 - 2025"
    submission_month_year: str = ""  # "April 2025"
    submission_date: str = ""        # "01/04/2025"
    place: str = ""                  # "Tambaram, Chennai – 59."

    # Thesis title
    thesis_title: str = ""           # ALL CAPS

    @property
    def surname(self) -> str:
        """Last word of the student's full name, used in body page header."""
        if not self.student_full_name:
            return ""
        return self.student_full_name.strip().split()[-1]


@dataclass
class BlockQuotation:
    """A block (long) quotation — indented, no quote marks, citation after."""
    text: str
    citation: str  # e.g. "(Boal 141)"


@dataclass
class SubSection:
    """1.3.1-style sub-section."""
    number: str          # "1.3.1"
    heading: str         # "HERMENEUTICS"
    paragraphs: list[str] = field(default_factory=list)
    block_quotations: list[tuple[int, BlockQuotation]] = field(default_factory=list)
    # block_quotations is list of (after_paragraph_index, BlockQuotation)


@dataclass
class Section:
    """1.1-style section within a chapter."""
    number: str          # "1.1"
    heading: str         # "LITERATURE REVIEW AND RESEARCH GAP"
    paragraphs: list[str] = field(default_factory=list)
    block_quotations: list[tuple[int, BlockQuotation]] = field(default_factory=list)
    subsections: list[SubSection] = field(default_factory=list)


@dataclass
class Chapter:
    """A thesis chapter."""
    number_roman: str    # "I"
    title: str           # "INTRODUCTION"
    intro_paragraphs: list[str] = field(default_factory=list)
    intro_block_quotations: list[tuple[int, BlockQuotation]] = field(default_factory=list)
    sections: list[Section] = field(default_factory=list)


@dataclass
class ThesisInput:
    """The complete structured input the formatter consumes."""
    front_matter: FrontMatter
    abstract: str
    keywords: list[str]
    acknowledgement: str
    chapters: list[Chapter]
    works_cited: list[str]  # MLA-formatted strings, ready to print


# ---------------------------------------------------------------------------
# Low-level helpers (XML manipulation for things python-docx doesn't expose)
# ---------------------------------------------------------------------------

def _add_page_number_field(run):
    """Insert a Word PAGE field into a run. Word evaluates it on open."""
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _add_toc_field(paragraph):
    """Insert a Word TOC field. Updates when user opens doc and presses F9."""
    run = paragraph.add_run()

    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    fld_begin.set(qn('w:dirty'), 'true')

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    # \o "1-3" = include heading levels 1-3
    # \h     = make TOC entries hyperlinks
    # \z     = hide tab leader and page numbers in Web layout
    # \u     = use applied paragraph outline level
    instr.text = r'TOC \o "1-3" \h \z \u'

    fld_separate = OxmlElement('w:fldChar')
    fld_separate.set(qn('w:fldCharType'), 'separate')

    # Placeholder text shown until field is updated
    placeholder = OxmlElement('w:r')
    placeholder_text = OxmlElement('w:t')
    placeholder_text.text = 'Right-click here and select "Update Field" to populate the Table of Contents.'
    placeholder.append(placeholder_text)

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(placeholder)
    run._r.append(fld_end)


def _set_cell_borders_none(cell):
    """Remove all borders from a table cell (for invisible signature blocks)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'nil')
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _set_table_borders_none(table):
    """Remove all borders from a table."""
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders_none(cell)


def _add_horizontal_line(paragraph):
    """Add a horizontal line below a paragraph (used as visual separator)."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    p_borders.append(bottom)
    p_pr.append(p_borders)


# ---------------------------------------------------------------------------
# Style configuration
# ---------------------------------------------------------------------------

FONT = "Times New Roman"
BODY_PT = 12
HEADING_PT = 14
LINE_SPACING = 1.5


# Words that stay lowercase in title-case (articles, short prepositions, conjunctions)
_LOWERCASE_WORDS = {
    "a", "an", "the", "and", "or", "but", "nor", "for", "yet", "so",
    "in", "on", "at", "by", "to", "of", "as", "from", "with", "into",
}


def _proper_case_degree(text: str) -> str:
    """Convert ALL CAPS degree program text to Title Case for use in prose.

    'MASTER OF ARTS' -> 'Master of Arts'
    'IN ENGLISH LANGUAGE AND LITERATURE' -> 'in English Language and Literature'
    """
    words = text.split()
    out = []
    for i, w in enumerate(words):
        lower = w.lower()
        # Keep small words lowercase except as the very first word
        if i > 0 and lower in _LOWERCASE_WORDS:
            out.append(lower)
        else:
            out.append(w.capitalize())
    return " ".join(out)


def _configure_default_style(doc: Document):
    """Set Normal style: TNR 12pt, 1.5 line spacing, justified."""
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(BODY_PT)
    pf = normal.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    # Force the East-Asian font slot too so Word doesn't substitute
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rfonts.set(qn(attr), FONT)


def _configure_heading_styles(doc: Document):
    """Configure Heading 1, 2, 3 styles to match the spec.

    Heading 1: chapter titles  — 14pt bold, centered, ALL CAPS, outline level 0
    Heading 2: section (1.1)   — 12pt bold, left-aligned, ALL CAPS, outline level 1
    Heading 3: subsection      — 12pt bold, left-aligned, ALL CAPS, outline level 2

    Outline levels matter for the TOC field to pick up entries.
    """
    for level, size, align, outline in (
        (1, HEADING_PT, WD_ALIGN_PARAGRAPH.CENTER, 0),
        (2, BODY_PT, WD_ALIGN_PARAGRAPH.LEFT, 1),
        (3, BODY_PT, WD_ALIGN_PARAGRAPH.LEFT, 2),
    ):
        style = doc.styles[f'Heading {level}']
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        pf = style.paragraph_format
        pf.alignment = align
        pf.line_spacing = LINE_SPACING
        pf.space_before = Pt(12)
        pf.space_after = Pt(12)
        pf.keep_with_next = True
        # Force font in XML
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.insert(0, rfonts)
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
            rfonts.set(qn(attr), FONT)


# ---------------------------------------------------------------------------
# Page setup
# ---------------------------------------------------------------------------

def _configure_page(section):
    """8.5x11, 1.5" left binding margin, 1" elsewhere."""
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.gutter = Inches(0)


# ---------------------------------------------------------------------------
# Paragraph helpers
# ---------------------------------------------------------------------------

def _add_centered(doc_or_section, text: str, *, bold=False, italic=False, size_pt=BODY_PT,
                  space_after_pt=0, all_caps=False):
    """Add a centered paragraph with given formatting."""
    p = doc_or_section.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_after = Pt(space_after_pt)
    run = p.add_run(text.upper() if all_caps else text)
    run.font.name = FONT
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    return p


def _add_body_paragraph(doc, text: str, *, indent_first_line=True, justify=True):
    """Add a justified body paragraph with first-line indent.

    Text may contain markdown italic markers (*foo* or _foo_) which will be
    rendered as italic runs within the paragraph. This is how MLA work titles
    (books, plays, films, journals, long poems) get italicized inline.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    if indent_first_line:
        pf.first_line_indent = Inches(0.5)

    for run_text, is_italic in parse_inline_runs(text):
        run = p.add_run(run_text)
        run.font.name = FONT
        run.font.size = Pt(BODY_PT)
        if is_italic:
            run.italic = True
    return p


def _add_block_quotation(doc, bq: BlockQuotation):
    """Add an MLA-style block quotation: indented, no quote marks, citation after.

    The quotation text may contain markdown italics, which are honored.
    The citation is rendered as plain text (no italics).
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.left_indent = Inches(0.5)
    pf.right_indent = Inches(0.0)
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)

    for run_text, is_italic in parse_inline_runs(bq.text):
        run = p.add_run(run_text)
        run.font.name = FONT
        run.font.size = Pt(BODY_PT)
        if is_italic:
            run.italic = True

    # Trailing space + citation (citation in plain text)
    cit_run = p.add_run(' ' + bq.citation)
    cit_run.font.name = FONT
    cit_run.font.size = Pt(BODY_PT)
    return p


def _add_blank_line(doc, size_pt=BODY_PT):
    """Add an empty paragraph for vertical spacing."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = LINE_SPACING
    run = p.add_run('')
    run.font.size = Pt(size_pt)
    return p


# ---------------------------------------------------------------------------
# Front matter renderers
# ---------------------------------------------------------------------------

def _add_title_page_paragraph(doc, text, *, bold=False, italic=False, size_pt=BODY_PT):
    """Add a centered paragraph with single-line spacing (for title page density)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.0  # single-spaced on title page so everything fits
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    return p


def _add_title_page_gap(doc, lines=1):
    """Add a small vertical gap on the title page (compact)."""
    for _ in range(lines):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run('')
        run.font.size = Pt(BODY_PT)


def _render_title_page(doc: Document, fm: FrontMatter, logo_path: Optional[str]):
    """Page 1: Title, dissertation declaration, candidate, supervisor, logo, dept, college, date.

    Uses single-line spacing throughout to ensure everything fits on one page,
    matching the reference PDF density.
    """

    # Title (bold, 14pt, centered, with mild line spacing for readability)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(fm.thesis_title)
    run.font.name = FONT
    run.font.size = Pt(HEADING_PT)
    run.bold = True

    _add_title_page_gap(doc, 1)

    # "A Dissertation" (italic)
    _add_title_page_paragraph(doc, "A Dissertation", italic=True)
    _add_title_page_paragraph(doc, f"submitted to the {fm.university_name}", italic=True)
    _add_title_page_paragraph(doc,
        "in partial fulfillment of the requirement for the degree of", italic=True)

    _add_title_page_gap(doc, 1)

    # Degree (bold)
    _add_title_page_paragraph(doc, fm.degree_program_line_1, bold=True)
    _add_title_page_paragraph(doc, fm.degree_program_line_2, bold=True)

    _add_title_page_gap(doc, 1)

    # Submitted by
    _add_title_page_paragraph(doc, "submitted by", italic=True)
    _add_title_page_paragraph(doc, fm.student_full_name, bold=True)
    _add_title_page_paragraph(doc, f"Register Number: {fm.register_number}", bold=True)

    _add_title_page_gap(doc, 1)

    # Supervisor block
    _add_title_page_paragraph(doc, "Under the supervision and guidance of", bold=True)
    _add_title_page_paragraph(doc, fm.supervisor_full_name, bold=True)
    _add_title_page_paragraph(doc, fm.supervisor_designation, bold=True)

    _add_title_page_gap(doc, 1)

    # Logo (centered, fit within 2" square)
    if logo_path and Path(logo_path).exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(2), height=Inches(2))

    _add_title_page_gap(doc, 1)

    # Department / College / Address (bold, centered)
    _add_title_page_paragraph(doc, fm.department_name, bold=True)
    _add_title_page_paragraph(doc, fm.college_name, bold=True)
    _add_title_page_paragraph(doc, fm.college_address, bold=True)

    _add_title_page_gap(doc, 1)

    # Submission date
    _add_title_page_paragraph(doc, fm.submission_month_year, bold=True)

    doc.add_page_break()


def _render_certificate(doc: Document, fm: FrontMatter):
    """Certificate page with two-column signature block (HOD left, Supervisor right)."""

    _add_centered(doc, "CERTIFICATE", bold=True, size_pt=HEADING_PT, space_after_pt=12)
    _add_blank_line(doc)

    # Justified certifying paragraph
    aided = " (Aided)" if fm.department_aided else ""
    cert_text = (
        f"This is to certify that this dissertation titled \u201C{fm.thesis_title}\u201D "
        f"is a record of original research work done by {fm.student_full_name.upper()} "
        f"({fm.register_number}), a full-time student of {fm.college_name.upper()}, "
        f"doing his/her {fm.degree_program_line_1} {fm.degree_program_line_2}, "
        f"in the {fm.department_name.upper()}, during the period {fm.study_period}. "
        f"This dissertation represents an entirely independent work on the part of the "
        f"candidate under the supervision and guidance of {fm.supervisor_full_name.upper()}, "
        f"{fm.supervisor_designation}, {fm.department_name}, {fm.college_name}."
    )
    _add_body_paragraph(doc, cert_text, indent_first_line=False)
    _add_blank_line(doc)
    _add_blank_line(doc)
    _add_blank_line(doc)

    # Two-column signature block via borderless table
    aided_str = " (Aided)" if fm.department_aided else ""
    sig_lines = [
        (fm.hod_full_name, fm.supervisor_full_name),
        (fm.hod_designation, "Research Supervisor"),
        (f"Department of English{aided_str}", fm.supervisor_designation),
        (fm.college_name, f"Department of English{aided_str}"),
        (fm.short_address or fm.college_address, fm.college_name),
        ("", fm.short_address or fm.college_address),
    ]

    table = doc.add_table(rows=len(sig_lines), cols=2)
    table.autofit = False
    for row_idx, (left, right) in enumerate(sig_lines):
        cells = table.rows[row_idx].cells
        for cell, text, is_bold in [(cells[0], left, row_idx == 0), (cells[1], right, row_idx == 0)]:
            cell.width = Inches(3)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = LINE_SPACING
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.name = FONT
            run.font.size = Pt(BODY_PT)
            run.bold = is_bold

    _set_table_borders_none(table)

    _add_blank_line(doc)
    _add_blank_line(doc)

    # Place + Date (left-aligned)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = LINE_SPACING
    run = p.add_run(f"Place: {fm.place}")
    run.font.name = FONT
    run.font.size = Pt(BODY_PT)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = LINE_SPACING
    run = p.add_run(f"Date: {fm.submission_date}")
    run.font.name = FONT
    run.font.size = Pt(BODY_PT)

    doc.add_page_break()


def _render_declaration(doc: Document, fm: FrontMatter):
    """Declaration page: candidate's own statement of originality."""

    _add_centered(doc, "DECLARATION", bold=True, size_pt=HEADING_PT, space_after_pt=12)
    _add_blank_line(doc)

    full_degree = _proper_case_degree(
        f"{fm.degree_program_line_1} {fm.degree_program_line_2}"
    )
    decl_text = (
        f"I, {fm.student_full_name.upper()}, hereby declare that the dissertation titled "
        f"\u201C{fm.thesis_title}\u201D, submitted by me for the partial fulfilment of the "
        f"degree of {full_degree} "
        f"is a record of original research work done by me under the guidance and supervision "
        f"of {fm.supervisor_full_name.upper()}, {fm.supervisor_designation}, {fm.department_name}, "
        f"{fm.college_name} during the period {fm.study_period}. No part of this dissertation "
        f"has been presented for the award of any degree from any other university or institution."
    )
    _add_body_paragraph(doc, decl_text, indent_first_line=False)
    _add_blank_line(doc)
    _add_blank_line(doc)
    _add_blank_line(doc)
    _add_blank_line(doc)

    # Place left, Name right (via table for clean alignment)
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    rows = [
        (f"Place: {fm.place}", fm.student_full_name.upper()),
        (f"Date: {fm.submission_date}", ""),
    ]
    for row_idx, (left, right) in enumerate(rows):
        cells = table.rows[row_idx].cells
        for cell in cells:
            cell.width = Inches(3)
        # Left
        p = cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = LINE_SPACING
        run = p.add_run(left)
        run.font.name = FONT
        run.font.size = Pt(BODY_PT)
        # Right
        p = cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.line_spacing = LINE_SPACING
        run = p.add_run(right)
        run.font.name = FONT
        run.font.size = Pt(BODY_PT)
        run.bold = True if row_idx == 0 and right else False

    _set_table_borders_none(table)
    doc.add_page_break()


def _render_acknowledgement(doc: Document, fm: FrontMatter, ack_text: str):
    """Acknowledgement page — student-written prose."""

    _add_centered(doc, "ACKNOWLEDGEMENT", bold=True, size_pt=HEADING_PT, space_after_pt=12)
    _add_blank_line(doc)

    # Split on double-newlines to allow multiple paragraphs
    paragraphs = [p.strip() for p in ack_text.split('\n\n') if p.strip()]
    for para in paragraphs:
        _add_body_paragraph(doc, para)

    _add_blank_line(doc)
    _add_blank_line(doc)

    # Date left, Name right
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    cells = table.rows[0].cells
    for cell in cells:
        cell.width = Inches(3)
    p = cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = LINE_SPACING
    run = p.add_run(f"Date: {fm.submission_date}")
    run.font.name = FONT
    run.font.size = Pt(BODY_PT)

    p = cells[1].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = LINE_SPACING
    run = p.add_run(fm.student_full_name.upper())
    run.font.name = FONT
    run.font.size = Pt(BODY_PT)

    _set_table_borders_none(table)
    doc.add_page_break()


def _render_abstract(doc: Document, abstract_text: str, keywords: list[str]):
    """Abstract page — single dense paragraph + Key Words line."""

    _add_centered(doc, "ABSTRACT", bold=True, size_pt=HEADING_PT, space_after_pt=12)
    _add_blank_line(doc)

    _add_body_paragraph(doc, abstract_text)
    _add_blank_line(doc)

    # Key Words line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = LINE_SPACING
    bold_run = p.add_run("Key Words: ")
    bold_run.font.name = FONT
    bold_run.font.size = Pt(BODY_PT)
    bold_run.bold = True
    body_run = p.add_run(", ".join(keywords) + ".")
    body_run.font.name = FONT
    body_run.font.size = Pt(BODY_PT)

    doc.add_page_break()


def _render_toc(doc: Document, thesis: Optional['ThesisInput'] = None,
                static_for_preview: bool = False):
    """Table of Contents.

    Default: inserts a Word TOC field that auto-populates when the doc is opened in Word.
    static_for_preview=True: renders a static TOC built from the thesis structure
                             (used for demo PDF previews where Word fields don't update).
    """

    _add_centered(doc, "TABLE OF CONTENTS", bold=True, size_pt=HEADING_PT, space_after_pt=12)
    _add_blank_line(doc)

    if static_for_preview and thesis is not None:
        # Build a static TOC from the chapter list. Page numbers are estimates;
        # the real Word TOC field will compute exact numbers when opened in Word.
        # The header row
        page = 1
        for ch in thesis.chapters:
            entry = f"Chapter {ch.number_roman}: {ch.title}"
            _add_toc_entry(doc, entry, page)
            # Estimate ~6 pages per chapter for the static preview
            # (Real auto-TOC in Word will be exact)
            page += 6
        _add_toc_entry(doc, "Works Cited", page)
    else:
        # Production path: Word TOC field
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = LINE_SPACING
        _add_toc_field(p)

    doc.add_page_break()


def _add_toc_entry(doc, text: str, page_num: int):
    """Add a single TOC line with right-aligned page number using a tab stop."""
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    # Right tab stop near the right margin (page is 8.5", left margin 1.5", right 1" → content 6")
    pf.tab_stops.add_tab_stop(Inches(5.8), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    run = p.add_run(text + "\t" + str(page_num))
    run.font.name = FONT
    run.font.size = Pt(BODY_PT)


# ---------------------------------------------------------------------------
# Body renderers
# ---------------------------------------------------------------------------

def _add_chapter_heading(doc: Document, number_roman: str, title: str):
    """CHAPTER <ROMAN>: <TITLE> — centered, bold, 14pt, ALL CAPS, Heading 1 style."""
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"CHAPTER {number_roman}: {title.upper()}")
    run.font.name = FONT
    run.font.size = Pt(HEADING_PT)
    run.bold = True


def _add_section_heading(doc: Document, number: str, heading: str):
    """1.1-style section heading — left-aligned, bold, 12pt, ALL CAPS, Heading 2 style."""
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 2']
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{number} {heading.upper()}")
    run.font.name = FONT
    run.font.size = Pt(BODY_PT)
    run.bold = True


def _add_subsection_heading(doc: Document, number: str, heading: str):
    """1.3.1-style sub-section heading — Heading 3 style."""
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 3']
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{number} {heading.upper()}")
    run.font.name = FONT
    run.font.size = Pt(BODY_PT)
    run.bold = True


def _render_paragraphs_with_quotes(doc: Document, paragraphs: list[str],
                                    block_quotations: list[tuple[int, BlockQuotation]]):
    """Render body paragraphs with block quotations interspersed at correct positions."""
    # Map paragraph_index -> list of block quotations that come after it
    quote_map: dict[int, list[BlockQuotation]] = {}
    for idx, bq in block_quotations:
        quote_map.setdefault(idx, []).append(bq)

    for i, para_text in enumerate(paragraphs):
        _add_body_paragraph(doc, para_text)
        for bq in quote_map.get(i, []):
            _add_block_quotation(doc, bq)


def _render_chapter(doc: Document, chapter: Chapter, *, repeat_title: bool = False,
                    thesis_title: str = ""):
    """Render a complete chapter with sections and sub-sections."""

    if repeat_title and thesis_title:
        _add_centered(doc, thesis_title, bold=True, size_pt=HEADING_PT,
                      space_after_pt=12, all_caps=True)
        _add_blank_line(doc)

    _add_chapter_heading(doc, chapter.number_roman, chapter.title)
    _add_blank_line(doc)

    # Intro paragraphs (before any section)
    _render_paragraphs_with_quotes(doc, chapter.intro_paragraphs,
                                   chapter.intro_block_quotations)

    # Sections
    for section in chapter.sections:
        _add_section_heading(doc, section.number, section.heading)
        _render_paragraphs_with_quotes(doc, section.paragraphs, section.block_quotations)

        # Subsections within the section
        for subsection in section.subsections:
            _add_subsection_heading(doc, subsection.number, subsection.heading)
            _render_paragraphs_with_quotes(doc, subsection.paragraphs,
                                           subsection.block_quotations)

    doc.add_page_break()


def _render_works_cited(doc: Document, entries: list[str]):
    """Works Cited page — hanging indent on each entry.

    Each entry may contain markdown italics for book/journal titles.
    """

    _add_centered(doc, "WORKS CITED", bold=True, size_pt=HEADING_PT, space_after_pt=12)
    _add_blank_line(doc)

    for entry in entries:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.line_spacing = LINE_SPACING
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.5)  # hanging indent
        pf.space_after = Pt(6)
        for run_text, is_italic in parse_inline_runs(entry):
            run = p.add_run(run_text)
            run.font.name = FONT
            run.font.size = Pt(BODY_PT)
            if is_italic:
                run.italic = True


# ---------------------------------------------------------------------------
# Section / header configuration
# ---------------------------------------------------------------------------

def _setup_body_header(section, surname: str):
    """Configure body-section header: '<surname> <pagenum>' top-right."""
    section.header.is_linked_to_previous = False
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(f"{surname} ")
    run.font.name = FONT
    run.font.size = Pt(BODY_PT)
    # Page number field
    page_run = p.add_run()
    page_run.font.name = FONT
    page_run.font.size = Pt(BODY_PT)
    _add_page_number_field(page_run)


def _clear_header(section):
    """Ensure a section has no header content (front matter pages)."""
    section.header.is_linked_to_previous = False
    header = section.header
    for p in list(header.paragraphs):
        for run in p.runs:
            run.text = ""


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def render_thesis_docx(thesis: ThesisInput, *, logo_path: Optional[str], output_path: str,
                       static_toc_for_preview: bool = False):
    """
    Render a complete MA thesis as a .docx file matching the MCC/UoM format.

    Args:
        thesis: Structured thesis content (front matter, abstract, chapters, etc.)
        logo_path: Path to the college logo image (PNG/JPG). Will be fit within 2"x2".
        output_path: Where to save the .docx file.
        static_toc_for_preview: If True, render a static TOC instead of a Word TOC
            field. Use this only for PDF previews where you can't update fields in Word.
            Production should leave this False so Word auto-updates page numbers.
    """

    doc = Document()

    # Document-level styles
    _configure_default_style(doc)
    _configure_heading_styles(doc)

    # Section 1: front matter (no header)
    _configure_page(doc.sections[0])
    _clear_header(doc.sections[0])

    # Render front matter
    _render_title_page(doc, thesis.front_matter, logo_path)
    _render_certificate(doc, thesis.front_matter)
    _render_declaration(doc, thesis.front_matter)
    _render_acknowledgement(doc, thesis.front_matter, thesis.acknowledgement)
    _render_abstract(doc, thesis.abstract, thesis.keywords)
    _render_toc(doc, thesis=thesis, static_for_preview=static_toc_for_preview)

    # Section break: new section for body, restart page numbering at 1
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    _configure_page(body_section)
    _setup_body_header(body_section, thesis.front_matter.surname)

    # Restart page numbering at 1 for body
    sect_pr = body_section._sectPr
    pg_num_type = OxmlElement('w:pgNumType')
    pg_num_type.set(qn('w:start'), '1')
    sect_pr.append(pg_num_type)

    # Render body chapters
    for i, chapter in enumerate(thesis.chapters):
        _render_chapter(
            doc, chapter,
            repeat_title=(i == 0),  # Only repeat the title above Chapter I
            thesis_title=thesis.front_matter.thesis_title,
        )

    # Works Cited
    _render_works_cited(doc, thesis.works_cited)

    # Save
    doc.save(output_path)
    return output_path
