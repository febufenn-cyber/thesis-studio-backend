"""Raw-XML pagination helpers for python-docx.

python-docx exposes no high-level API for page-number fields or section
page-format/restart control.  This module fills the gap using OOXML elements
constructed via ``docx.oxml.OxmlElement`` and ``docx.oxml.ns.qn``.

Public helpers
--------------
add_page_number_field(paragraph)
    Append a ``w:fldSimple`` PAGE field to an existing paragraph.

set_section_pagenum_format(section, fmt, start)
    Insert/replace ``w:pgNumType`` on a section's ``w:sectPr``.

suppress_first_page_number(section)
    Enable different-first-page and leave the first-page footer empty,
    effectively hiding the page number on page 1 (title page).

add_mla_header(section, surname)
    Add a right-aligned header paragraph reading "{surname} {PAGE}" using
    a ``w:fldSimple`` PAGE field.
"""

from __future__ import annotations

from typing import Literal

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.section import Section
from docx.text.paragraph import Paragraph

# Word namespace URI — used for lxml-level attribute access and etree searches.
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def add_page_number_field(paragraph: Paragraph) -> None:
    """Append a ``w:fldSimple PAGE`` field to *paragraph*.

    OOXML construct::

        <w:fldSimple w:instr=" PAGE ">
          <w:r/>
        </w:fldSimple>

    Word evaluates the ``w:instr`` instruction to the current page number
    when the document is opened or updated.  The empty ``w:r`` child is
    required for the field to render — Word replaces it with the computed
    value at update time.

    Parameters
    ----------
    paragraph:
        The python-docx ``Paragraph`` whose underlying ``w:p`` element
        receives the field.
    """
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    # Empty run — Word populates this with the current page number.
    run = OxmlElement("w:r")
    fld.append(run)
    paragraph._p.append(fld)


def set_section_pagenum_format(
    section: Section,
    fmt: Literal["lowerRoman", "upper_roman", "decimal"],
    start: int | None = None,
) -> None:
    """Set (or replace) ``w:pgNumType`` on *section*'s ``w:sectPr``.

    OOXML construct::

        <w:pgNumType w:fmt="lowerRoman" w:start="1"/>

    ``w:fmt`` controls the numeral style: ``"lowerRoman"`` (i, ii, iii …),
    ``"upperRoman"`` (I, II, III …), or ``"decimal"`` (1, 2, 3 …).

    If *start* is provided the ``w:start`` attribute is set, which restarts
    counting at that number in this section — essential for the body section
    that must restart at arabic 1 after roman front-matter pages.

    Parameters
    ----------
    section:
        A python-docx ``Section`` object.
    fmt:
        The page-number format string as an OOXML ``ST_NumberFormat`` value.
        Pass ``"lowerRoman"`` for front-matter, ``"decimal"`` for body.
    start:
        Integer to restart numbering at.  ``None`` inherits from the
        previous section (no ``w:start`` attribute emitted).
    """
    sectPr = section._sectPr

    # Remove any existing pgNumType to avoid duplicates.
    for existing in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(existing)

    pg_num = OxmlElement("w:pgNumType")
    pg_num.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num.set(qn("w:start"), str(start))

    # Insert after w:pgMar if present, otherwise append.  The OOXML schema
    # places w:pgNumType after w:pgMar in the CT_SectPr sequence.
    pg_mar = sectPr.find(qn("w:pgMar"))
    if pg_mar is not None:
        pg_mar.addnext(pg_num)
    else:
        sectPr.append(pg_num)


def suppress_first_page_number(section: Section) -> None:
    """Suppress the page number on the first page of *section*.

    Sets ``section.different_first_page_header_footer = True``, which inserts
    ``w:titlePg`` into the ``w:sectPr`` element.  Word then uses separate
    first-page header/footer slots.  Clearing the first-page footer (which
    this function does) leaves page 1 with no visible page number while the
    default footer still carries the roman numeral on subsequent pages.

    This matches the FORMAT_SPEC §1 requirement: ``tn_university`` title page
    is counted as roman *i* but the number is suppressed.
    """
    section.different_first_page_header_footer = True

    # Ensure the first-page footer paragraph is empty.
    first_footer = section.first_page_footer
    for para in first_footer.paragraphs:
        para.clear()


def add_mla_header(section: Section, surname: str) -> None:
    """Add an MLA-style right-aligned header reading "{surname} {PAGE}".

    MLA 9 requires a running header top-right: the student's surname followed
    by a space and the current page number, on every page from page 1.

    OOXML structure emitted in ``word/header*.xml``::

        <w:p>
          <w:pPr>
            <w:jc w:val="right"/>
          </w:pPr>
          <w:r><w:t xml:space="preserve">{surname} </w:t></w:r>
          <w:fldSimple w:instr=" PAGE "><w:r/></w:fldSimple>
        </w:p>

    Parameters
    ----------
    section:
        The section whose default (non-first-page) header is configured.
    surname:
        The student's surname as it should appear in the header, e.g. ``"Sharma"``.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # local import — avoids circular

    header = section.header
    # python-docx always provides at least one paragraph in the header part.
    para = header.paragraphs[0]
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Surname run — note the trailing space before the page field.
    para.add_run(f"{surname} ")

    # PAGE field
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    r = OxmlElement("w:r")
    fld.append(r)
    para._p.append(fld)
