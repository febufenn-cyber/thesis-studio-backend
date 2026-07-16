"""DOCX renderer — consumes the canonical ThesisDocument (FORMAT_SPEC §1–§7).

Layout source of truth for exports; PDF is converted from this output.
Synchronous module (python-docx is blocking) — callers run it in a worker
thread (asyncio.to_thread), mirroring the v1 compile pipeline discipline.
"""

from __future__ import annotations

from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt

from app.canonical.model import (
    BlockQuoteBlock,
    ChapterDoc,
    FrontMatterEntry,
    HeadingBlock,
    MarkerBlock,
    ParagraphBlock,
    Run,
    ThesisDocument,
    VerseQuoteBlock,
)
from app.renderers.pagination import (
    add_mla_header,
    add_page_number_field,
    set_section_pagenum_format,
    suppress_first_page_number,
)
from app.renderers.profiles import ResolvedProfile
from app.renderers.styles import get_citation_style
from app.renderers.works_cited import SourceLike


class RenderError(Exception):
    """The document cannot be rendered (collected problems in the message)."""


_ROMAN = ["", "I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X",
          "XI", "XII", "XIII", "XIV", "XV"]

_PGNUM_FMT = {"lower_roman": "lowerRoman", "upper_roman": "upperRoman",
              "arabic": "decimal", "decimal": "decimal"}

# A4 in mm; Letter in inches.
_PAGE_DIMS = {"A4": (Mm(210), Mm(297)), "Letter": (Inches(8.5), Inches(11))}


def _roman(n: int) -> str:
    return _ROMAN[n] if 0 < n < len(_ROMAN) else str(n)


def _setup_section(section: Any, prof: ResolvedProfile) -> None:
    w, h = _PAGE_DIMS.get(prof.page.size, _PAGE_DIMS["A4"])
    section.page_width, section.page_height = w, h
    m = prof.page.margins_in
    section.left_margin = Inches(m.left)
    section.right_margin = Inches(m.right)
    section.top_margin = Inches(m.top)
    section.bottom_margin = Inches(m.bottom)


def _create_styles(doc: Document, prof: ResolvedProfile) -> None:
    """Named TS-* styles per DESIGN §9, parameterised by the profile."""
    from docx.enum.style import WD_STYLE_TYPE

    t = prof.type

    def mk(name: str) -> Any:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = t.font
        style.font.size = Pt(t.size_pt)
        style.paragraph_format.line_spacing = t.line_spacing
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        return style

    normal = mk("TS-Normal")
    normal.paragraph_format.first_line_indent = Inches(t.first_line_indent_in)
    if t.justify_body:
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    center = mk("TS-FrontCenter")
    center.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    label = mk("TS-ChapterLabel")
    label.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = mk("TS-ChapterTitle")
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.font.bold = prof.chapter_label.title_bold

    h2 = mk("TS-Heading2")
    h2.font.bold = prof.headings.h2.bold

    h3 = mk("TS-Heading3")
    h3.font.italic = prof.headings.h3.italic

    bq = mk("TS-BlockQuote")
    bq.paragraph_format.left_indent = Inches(prof.quotes.block_indent_in)
    bq.paragraph_format.first_line_indent = Inches(0)

    verse = mk("TS-Verse")
    verse.paragraph_format.left_indent = Inches(prof.quotes.block_indent_in)
    verse.paragraph_format.first_line_indent = Inches(0)

    wc = mk("TS-WorksCited")
    wc.paragraph_format.left_indent = Inches(prof.works_cited.hanging_indent_in)
    wc.paragraph_format.first_line_indent = Inches(-prof.works_cited.hanging_indent_in)


def _add_runs(para: Any, runs: list[Run]) -> None:
    for r in runs:
        added = para.add_run(r.text)
        added.italic = r.italic or None


def _center_footer_page_field(section: Any) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0]
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(para)


def _p(doc: Document, style: str, text: str = "", bold: bool = False,
       size_pt: int | None = None) -> Any:
    para = doc.add_paragraph(style=style)
    if text:
        run = para.add_run(text)
        run.bold = bold or None
        if size_pt:
            run.font.size = Pt(size_pt)
    return para


def _page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break()
    # Use a hard page break on the last empty paragraph's run.
    doc.paragraphs[-1].runs[0]._r.append(OxmlElement("w:br"))
    doc.paragraphs[-1].runs[0]._r[-1].set(qn("w:type"), "page")


# ---------------------------------------------------------------------------
# Front matter (FORMAT_SPEC §3)
# ---------------------------------------------------------------------------


def _fm_title_page(doc: Document, td: ThesisDocument) -> None:
    """Title page with graceful omission: a line whose metadata slot is empty
    is dropped entirely — never rendered as "submitted to ," artifacts. The
    review/final gate (verifier) still requires the fields before a final
    export; this only keeps drafts presentable (FRICTION_LOG F-title)."""
    m = td.meta
    _p(doc, "TS-FrontCenter", m.title.upper(), bold=True, size_pt=14)
    _p(doc, "TS-FrontCenter")
    if m.college.name.strip():
        _p(doc, "TS-FrontCenter", f"A dissertation submitted to {m.college.name},")
        if m.college.affiliation.strip():
            _p(doc, "TS-FrontCenter", f"affiliated to {m.college.affiliation},")
    else:
        _p(doc, "TS-FrontCenter", "A dissertation submitted")
    _p(doc, "TS-FrontCenter", "in partial fulfilment of the requirements")
    _p(doc, "TS-FrontCenter", "for the award of the degree of")
    _p(doc, "TS-FrontCenter")
    _p(doc, "TS-FrontCenter", m.degree.upper(), bold=True)
    if m.candidate.name.strip():
        _p(doc, "TS-FrontCenter")
        _p(doc, "TS-FrontCenter", "By")
        _p(doc, "TS-FrontCenter", m.candidate.name)
        if m.candidate.reg_no:
            _p(doc, "TS-FrontCenter", f"(Reg. No. {m.candidate.reg_no})")
    if m.guide.name.strip():
        _p(doc, "TS-FrontCenter")
        _p(doc, "TS-FrontCenter", "Under the guidance of")
        guide = m.guide.name + (f", {m.guide.designation}" if m.guide.designation else "")
        _p(doc, "TS-FrontCenter", guide)
    _p(doc, "TS-FrontCenter")
    for line in (m.department, m.college.name):
        if line.strip():
            _p(doc, "TS-FrontCenter", line)
    if m.college.city.strip():
        city = m.college.city + (f" – {m.college.pin}" if m.college.pin else "")
        _p(doc, "TS-FrontCenter", city)
    sub = f"{m.submission.month} {m.submission.year or ''}".strip()
    if sub:
        _p(doc, "TS-FrontCenter")
        _p(doc, "TS-FrontCenter", sub)


def _fm_certificate(doc: Document, td: ThesisDocument) -> None:
    m = td.meta
    _p(doc, "TS-FrontCenter", "CERTIFICATE", bold=True)
    _p(doc, "TS-FrontCenter")
    body = doc.add_paragraph(style="TS-Normal")
    body.add_run(
        "This is to certify that the dissertation entitled "
    )
    body.add_run(f"“{m.title}”").italic = True
    body.add_run(
        f" submitted by {m.candidate.name}"
        + (f" (Reg. No. {m.candidate.reg_no})" if m.candidate.reg_no else "")
        + f" to {m.college.name}, affiliated to {m.college.affiliation}, in partial"
        f" fulfilment of the requirements for the award of the degree of"
        f" {m.degree}, is a record of bonafide research work carried out by the"
        " candidate under my supervision and guidance, and that it has not"
        " previously formed the basis for the award of any degree, diploma,"
        " or similar title."
    )
    _p(doc, "TS-Normal")
    sig = doc.add_paragraph(style="TS-Normal")
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig.paragraph_format.first_line_indent = Inches(0)
    sig.add_run(f"{m.guide.name}\n{m.guide.designation}\nResearch Supervisor")
    place = doc.add_paragraph(style="TS-Normal")
    place.paragraph_format.first_line_indent = Inches(0)
    place.add_run(f"Place: {m.college.city}\nDate:")


def _fm_declaration(doc: Document, td: ThesisDocument) -> None:
    m = td.meta
    _p(doc, "TS-FrontCenter", "DECLARATION", bold=True)
    _p(doc, "TS-FrontCenter")
    body = doc.add_paragraph(style="TS-Normal")
    body.add_run(
        f"I, {m.candidate.name}"
        + (f" (Reg. No. {m.candidate.reg_no})" if m.candidate.reg_no else "")
        + ", hereby declare that the dissertation entitled "
    )
    body.add_run(f"“{m.title}”").italic = True
    body.add_run(
        f", submitted to {m.college.name}, affiliated to {m.college.affiliation},"
        f" in partial fulfilment of the requirements for the award of the degree"
        f" of {m.degree}, is my original work, and that it has not been"
        " submitted, either in part or in full, for the award of any other"
        " degree or diploma of this or any other university."
    )
    _p(doc, "TS-Normal")
    sig = doc.add_paragraph(style="TS-Normal")
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig.paragraph_format.first_line_indent = Inches(0)
    sig.add_run(m.candidate.name)
    place = doc.add_paragraph(style="TS-Normal")
    place.paragraph_format.first_line_indent = Inches(0)
    place.add_run(f"Place: {m.college.city}\nDate:")


def _fm_body_page(doc: Document, heading: str, entry: FrontMatterEntry,
                  fallback: str = "") -> None:
    _p(doc, "TS-FrontCenter", heading, bold=True)
    _p(doc, "TS-FrontCenter")
    if entry.body_blocks:
        for block in entry.body_blocks:
            _render_block(doc, block)
    elif fallback:
        para = doc.add_paragraph(style="TS-Normal")
        para.add_run(fallback)


def _fm_contents(doc: Document, td: ThesisDocument, prof: ResolvedProfile) -> None:
    _p(doc, "TS-FrontCenter", "CONTENTS", bold=True)
    _p(doc, "TS-FrontCenter")
    if prof.toc.native_word_field:
        para = doc.add_paragraph(style="TS-Normal")
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), ' TOC \\o "1-2" \\h \\z \\u ')
        fld.append(OxmlElement("w:r"))
        para._p.append(fld)
        note = doc.add_paragraph(style="TS-Normal")
        note.add_run("(Right-click the table and choose “Update Field” in Word.)").italic = True
    else:
        for ch in td.chapters:
            row = doc.add_paragraph(style="TS-Normal")
            row.paragraph_format.first_line_indent = Inches(0)
            row.add_run(f"CHAPTER {_roman(ch.number)}\t{ch.title.upper()}")
        row = doc.add_paragraph(style="TS-Normal")
        row.paragraph_format.first_line_indent = Inches(0)
        row.add_run("WORKS CITED")


_FM_RENDERERS = {
    "title_page": _fm_title_page,
    "certificate": _fm_certificate,
    "declaration": _fm_declaration,
}


# ---------------------------------------------------------------------------
# Body blocks (FORMAT_SPEC §4–§5)
# ---------------------------------------------------------------------------


def _render_block(doc: Document, block: Any) -> None:
    if isinstance(block, ParagraphBlock):
        para = doc.add_paragraph(style="TS-Normal")
        _add_runs(para, block.runs)
    elif isinstance(block, BlockQuoteBlock):
        para = doc.add_paragraph(style="TS-BlockQuote")
        text = block.text
        citation = f" ({block.citation})" if block.citation else ""
        para.add_run(text + citation)
    elif isinstance(block, VerseQuoteBlock):
        para = doc.add_paragraph(style="TS-Verse")
        for i, line in enumerate(block.lines):
            run = para.add_run(line)
            if i < len(block.lines) - 1:
                run.add_break()
        if block.citation:
            para.add_run(f" ({block.citation})")
    elif isinstance(block, HeadingBlock):
        style = "TS-Heading2" if block.level == 2 else "TS-Heading3"
        doc.add_paragraph(block.text, style=style)
    elif isinstance(block, MarkerBlock):
        para = doc.add_paragraph(style="TS-Normal")
        para.add_run(f"[{block.kind}: {block.note}]").bold = True


def _render_chapter(doc: Document, ch: ChapterDoc, prof: ResolvedProfile,
                    first: bool) -> None:
    if not first and prof.chapter_label.new_page:
        _page_break(doc)
    label = prof.chapter_label.format.replace("{ROMAN}", _roman(ch.number))
    _p(doc, "TS-ChapterLabel", label, bold=prof.chapter_label.title_bold)
    _p(doc, "TS-ChapterLabel")
    title = ch.title.upper() if prof.chapter_label.title_caps else ch.title
    _p(doc, "TS-ChapterTitle", title, bold=prof.chapter_label.title_bold)
    _p(doc, "TS-ChapterLabel")
    for block in ch.blocks:
        _render_block(doc, block)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------


def render_docx(
    doc_model: ThesisDocument,
    sources: dict[Any, SourceLike],
    profile: ResolvedProfile,
    output_path: str,
    *,
    strict: bool = True,
) -> str:
    """Render the canonical document to *output_path*; returns the path.

    strict=True (final exports): raises RenderError listing every works-cited
    problem at once — an incomplete citation can never reach a final document.
    strict=False (review exports): sources with missing required fields render
    as their original imported line behind a loud [UNVERIFIED ...] marker, so a
    draft can always be produced without laundering anything (FRICTION_LOG /
    Priya rule 4).
    """
    from app.renderers.works_cited import MissingCitationField, _sort_key, fallback_entry

    problems: list[str] = []
    used_sources: list[SourceLike] = []
    for ref in doc_model.works_cited:
        src = sources.get(ref.source_id) or sources.get(str(ref.source_id))
        if src is None:
            problems.append(f"works_cited references unknown source {ref.source_id}")
        else:
            used_sources.append(src)

    wc_entries: list[list[Run]] = []
    if used_sources:
        style = get_citation_style(doc_model.meta.citation_style)
        if strict:
            try:
                wc_entries = style.sorted_entries(used_sources)
            except MissingCitationField as exc:
                problems.append(str(exc))
        else:
            good: list[SourceLike] = []
            flagged: list[SourceLike] = []
            for src in used_sources:
                try:
                    style.sorted_entries([src])
                    good.append(src)
                except MissingCitationField:
                    flagged.append(src)
            merged: list[tuple[tuple[str, str], list[Run]]] = []
            if good:
                good_sorted = sorted(good, key=_sort_key)
                for src, runs in zip(good_sorted, style.sorted_entries(good)):
                    merged.append((_sort_key(src), runs))
            for src in flagged:
                merged.append((_sort_key(src), fallback_entry(src)))
            wc_entries = [runs for _, runs in sorted(merged, key=lambda item: item[0])]
    if problems:
        raise RenderError("; ".join(problems))

    doc = Document()
    _create_styles(doc, profile)
    section = doc.sections[0]
    _setup_section(section, profile)

    is_mla = profile.pagination.mla_header_name
    has_front = bool(profile.front_matter_order) and bool(doc_model.front_matter)

    if is_mla:
        # MLA first-page heading block + running header from page 1.
        add_mla_header(section, doc_model.meta.candidate.name.split()[-1]
                       if doc_model.meta.candidate.name else "")
        set_section_pagenum_format(section, "decimal", start=1)
        head = doc.add_paragraph(style="TS-Normal")
        head.paragraph_format.first_line_indent = Inches(0)
        m = doc_model.meta
        sub = f"{m.submission.month} {m.submission.year or ''}".strip()
        head.add_run("\n".join(x for x in (
            m.candidate.name, m.guide.name, m.degree, sub) if x))
        title_para = doc.add_paragraph(style="TS-FrontCenter")
        title_para.add_run(doc_model.meta.title)
        body_doc_start_new_section = False
    elif has_front:
        fmt = _PGNUM_FMT.get(profile.pagination.front.style, "lowerRoman")
        set_section_pagenum_format(section, fmt, start=1)
        _center_footer_page_field(section)
        suppress_first_page_number(section)

        entries = {e.kind: e for e in doc_model.front_matter}
        ordered = [k for k in profile.front_matter_order if k in entries]
        if (doc_model.meta.ai_disclosure.enabled and "ai_disclosure" in entries
                and "ai_disclosure" not in ordered):
            idx = ordered.index("contents") if "contents" in ordered else len(ordered)
            ordered.insert(idx, "ai_disclosure")
        for i, kind in enumerate(ordered):
            if i:
                _page_break(doc)
            entry = entries[kind]
            if kind in _FM_RENDERERS:
                _FM_RENDERERS[kind](doc, doc_model)
            elif kind == "acknowledgement":
                _fm_body_page(doc, "ACKNOWLEDGEMENT", entry)
            elif kind == "ai_disclosure":
                _fm_body_page(doc, "AI-ASSISTANCE DISCLOSURE", entry,
                              fallback=doc_model.meta.ai_disclosure.text)
            elif kind == "contents":
                _fm_contents(doc, doc_model, profile)
            elif kind == "abbreviations":
                _fm_body_page(doc, "LIST OF ABBREVIATIONS", entry)
        body_doc_start_new_section = True
    else:
        set_section_pagenum_format(section, "decimal", start=1)
        _center_footer_page_field(section)
        body_doc_start_new_section = False

    # Body section.
    if body_doc_start_new_section:
        body_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        _setup_section(body_section, profile)
        fmt = _PGNUM_FMT.get(profile.pagination.body.style, "decimal")
        set_section_pagenum_format(
            body_section, fmt, start=profile.pagination.body.restart_at
        )
        if not is_mla:
            _center_footer_page_field(body_section)
            # New section inherits titlePg from section 1 unless reset.
            body_section.different_first_page_header_footer = False

    for i, ch in enumerate(doc_model.chapters):
        _render_chapter(doc, ch, profile, first=(i == 0))

    if wc_entries:
        _page_break(doc)
        heading = profile.works_cited.heading
        _p(doc, "TS-FrontCenter", heading, bold=profile.works_cited.heading_bold)
        _p(doc, "TS-FrontCenter")
        for entry_runs in wc_entries:
            para = doc.add_paragraph(style="TS-WorksCited")
            _add_runs(para, entry_runs)

    doc.save(output_path)
    return output_path
