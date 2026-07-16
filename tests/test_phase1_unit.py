"""Fast unit coverage for Phase 1 contracts that do not require PostgreSQL."""

from __future__ import annotations

from copy import deepcopy
from types import SimpleNamespace
from uuid import uuid4

import pytest
from docx import Document

from app.canonical.model import ParagraphBlock, Run, ThesisDocument
from app.ingest.citations import InTextCitation, resolve_citation
from app.ingest.docx_extract import ExtractedPara, ExtractedRun
from app.ingest.preflight import ManuscriptValidationError, inspect_docx
from app.ingest.structure import parse_manuscript
from app.ingest.verifier import verify
from app.models.project import Project
from app.renderers.phase1_profiles import resolve_phase1_profile
from app.services.export_service import _apply_review_qa_policy


def test_canonical_blocks_receive_stable_ids_and_provenance() -> None:
    revision_id = uuid4()
    block = ParagraphBlock(
        runs=[Run(text="Preserved text")],
        source_revision_id=revision_id,
        source_paragraph_index=17,
    )
    payload = block.model_dump(mode="json")
    restored = ParagraphBlock.model_validate(payload)
    assert restored.id == block.id
    assert restored.source_revision_id == revision_id
    assert restored.source_paragraph_index == 17


def test_docx_preflight_reports_unsupported_table(tmp_path) -> None:
    path = tmp_path / "table.docx"
    doc = Document()
    doc.add_paragraph("CHAPTER I")
    doc.add_table(rows=1, cols=2)
    doc.save(path)

    report = inspect_docx(str(path))
    assert report.counts["tables"] == 1
    assert any(
        issue.code == "unsupported_tables" and issue.severity == "block"
        for issue in report.issues
    )


def test_docx_preflight_rejects_non_docx(tmp_path) -> None:
    path = tmp_path / "fake.docx"
    path.write_text("not a zip package", encoding="utf-8")
    with pytest.raises(ManuscriptValidationError):
        inspect_docx(str(path))


def test_parser_accounts_for_chapter_boundary_title_and_body() -> None:
    revision_id = uuid4()
    paras = [
        ExtractedPara(index=4, runs=[ExtractedRun("CHAPTER I")], all_caps=True),
        ExtractedPara(index=5, runs=[ExtractedRun("INTRODUCTION")], all_caps=True),
        ExtractedPara(index=6, runs=[ExtractedRun("The argument begins here.")]),
    ]
    result = parse_manuscript(paras, revision_id)
    chapter = result.document.chapters[0]
    assert chapter.source_paragraph_index == 4
    assert chapter.title_source_paragraph_index == 5
    assert chapter.blocks[0].source_paragraph_index == 6
    assert set(result.structural_paragraph_indexes) == {4, 5}


def test_citation_resolution_never_guesses_between_same_surname() -> None:
    first, second = uuid4(), uuid4()
    sources = {
        first: SimpleNamespace(fields={"author": "Achebe, Chinua", "title": "Things Fall Apart"}),
        second: SimpleNamespace(fields={"author": "Achebe, Chinua", "title": "No Longer at Ease"}),
    }
    ambiguous = InTextCitation(surname="Achebe", pages="45", raw="(Achebe 45)")
    resolved_id, candidates, reason = resolve_citation(ambiguous, sources)
    assert resolved_id is None
    assert set(candidates) == {first, second}
    assert reason == "ambiguous_surname"

    hinted = InTextCitation(
        surname="Achebe", pages="45", title_hint="Things Fall Apart", raw="(Achebe, Things Fall Apart 45)"
    )
    resolved_id, _, reason = resolve_citation(hinted, sources)
    assert resolved_id == first
    assert reason == "surname_and_title"


def test_verifier_honors_exact_human_resolution() -> None:
    source_id = uuid4()
    block = ParagraphBlock(runs=[Run(text="The claim appears here (Achebe 45).")])
    document = ThesisDocument(
        chapters=[{"number": 1, "title": "Introduction", "blocks": [block]}],
        works_cited=[{"source_id": source_id}],
    )
    source = SimpleNamespace(
        fields={
            "author": "Achebe, Chinua",
            "title": "Things Fall Apart",
            "publisher": "Heinemann",
            "year": "1958",
        },
        verified=True,
        parse_status="fully_structured",
        consulted_flag=False,
        raw_entry="",
    )
    report = verify(
        document,
        {source_id: source},
        {},
        {(str(block.id), "(Achebe 45)"): source_id},
    )
    assert not any(v.rule.startswith("citation_") for v in report.violations)


def test_mcc_profile_uses_verified_spacing_and_native_toc() -> None:
    profile, version = resolve_phase1_profile("mcc_ma_english_2026")
    assert profile.type.line_spacing == 1.5
    assert profile.toc.native_word_field is True
    assert version == "mcc_ma_english_2026:v1"


def test_generic_tn_profile_is_explicitly_unverified() -> None:
    profile, version = resolve_phase1_profile("tn_university")
    assert profile.toc.native_word_field is True
    assert "not institution-certified" in profile.notes
    assert "unverified" in version


def test_active_revision_fk_is_named_and_breaks_metadata_cycle() -> None:
    foreign_key = next(iter(Project.__table__.c.active_revision_id.foreign_keys))
    assert foreign_key.target_fullname == "manuscript_revisions.id"
    assert foreign_key.constraint.name == "fk_projects_active_revision"
    assert foreign_key.constraint.use_alter is True


def test_review_export_downgrades_only_visible_markers() -> None:
    marker = {
        "rule": "unresolved_marker_rendered",
        "severity": "block",
        "found": "[QUOTE_NEEDED:",
        "expected": "no unresolved marker in final output",
        "location": {"section": "rendered_output"},
    }
    final_result = _apply_review_qa_policy(
        {"pass": False, "violations": [deepcopy(marker)]},
        review_export=False,
    )
    assert final_result["pass"] is False
    assert final_result["violations"][0]["severity"] == "block"

    review_result = _apply_review_qa_policy(
        {"pass": False, "violations": [deepcopy(marker)]},
        review_export=True,
    )
    assert review_result["pass"] is True
    assert review_result["review_export"] is True
    assert review_result["violations"][0]["severity"] == "warn"

    structural = {
        "rule": "margin_mismatch",
        "severity": "block",
        "found": "1.0",
        "expected": "1.5",
        "location": {"section": "rendered_output"},
    }
    mixed = _apply_review_qa_policy(
        {"pass": False, "violations": [deepcopy(marker), structural]},
        review_export=True,
    )
    assert mixed["pass"] is False
    assert structural in mixed["violations"]


def test_heading_recovery_on_messy_student_formatting() -> None:
    """FRICTION_LOG F2: real students format chapter headings chaotically.

    The parser must recover a Word Heading 1, a spelled-out 'Chapter Two' in
    bold, and a standalone ALL-CAPS 'INTRODUCTION'/'CONCLUSION' — and must NOT
    promote a plain lowercase paragraph to a chapter.
    """
    from app.ingest.docx_extract import ExtractedPara, ExtractedRun
    from app.ingest.structure import parse_manuscript

    def para(idx, text, *, style="Normal", bold=False, caps=False, center=False):
        return ExtractedPara(
            index=idx,
            runs=[ExtractedRun(text=text, bold=bold)],
            style_name=style,
            alignment="center" if center else "",
            all_caps=caps,
            mostly_bold=bold,
        )

    paras = [
        para(0, "A Study of Voice", center=True),          # title page
        para(1, "By Priya Ramesh", center=True),
        para(2, "INTRODUCTION", style="Heading 1"),        # real H1
        para(3, "The introduction discusses the argument at length here."),
        para(4, "Chapter Two: Memory as Method", bold=True),  # bold, spelled-out
        para(5, "This chapter turns to memory as a formal device in the novel."),
        para(6, "CONCLUSION", caps=True),                  # standalone caps word
        para(7, "In conclusion the argument is drawn together for the reader."),
        para(8, "WORKS CITED", caps=True),
        para(9, "Ishiguro, Kazuo. The Remains of the Day. Faber, 1989."),
    ]
    result = parse_manuscript(paras)
    titles = [c.title for c in result.document.chapters]
    assert titles == ["INTRODUCTION", "Memory as Method", "CONCLUSION"]
    # The plain lowercase body paragraphs were NOT promoted to chapters.
    assert all("introduction discusses" not in t for t in titles)
    assert len(result.wc_raw_entries) == 1
