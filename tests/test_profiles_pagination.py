"""Tests for app/renderers/profiles.py and app/renderers/pagination.py.

Covers:
- resolve_profile merge semantics (partial override, unknown keys → notes)
- TN_UNIVERSITY and MLA_STRICT built-in spot-checks
- Pagination XML helpers verified via lxml etree inspection of a saved .docx
"""

from __future__ import annotations

import io
import zipfile

import lxml.etree
import pytest
from docx import Document

from app.renderers.pagination import (
    add_mla_header,
    add_page_number_field,
    set_section_pagenum_format,
    suppress_first_page_number,
)
from app.renderers.profiles import (
    MLA_STRICT,
    TN_UNIVERSITY,
    ResolvedProfile,
    resolve_profile,
)

# ---------------------------------------------------------------------------
# Helper
# ---------------------------------------------------------------------------

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W = f"{{{_W_NS}}}"  # Clark-notation prefix shorthand


def _docx_bytes(doc: Document) -> bytes:
    """Render *doc* to an in-memory bytes buffer and return the bytes."""
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _read_xml_part(docx_bytes: bytes, part_path: str) -> lxml.etree._Element:
    """Extract and parse an XML part from a .docx (zip) file."""
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        raw = zf.read(part_path)
    return lxml.etree.fromstring(raw)


def _qn(local: str) -> str:
    """Return Clark-notation tag ``{w_namespace}local``."""
    return f"{_W}{local}"


# ---------------------------------------------------------------------------
# profiles — built-in spot-checks
# ---------------------------------------------------------------------------


class TestBuiltinProfiles:
    def test_tn_university_page_size(self) -> None:
        assert TN_UNIVERSITY.page.size == "A4"

    def test_tn_university_left_margin(self) -> None:
        assert TN_UNIVERSITY.page.margins_in.left == 1.5

    def test_tn_university_other_margins_one_inch(self) -> None:
        m = TN_UNIVERSITY.page.margins_in
        assert m.top == 1.0
        assert m.bottom == 1.0
        assert m.right == 1.0

    def test_tn_university_font_and_size(self) -> None:
        assert TN_UNIVERSITY.type.font == "Times New Roman"
        assert TN_UNIVERSITY.type.size_pt == 12

    def test_tn_university_double_spacing(self) -> None:
        assert TN_UNIVERSITY.type.line_spacing == 2.0

    def test_tn_university_justified(self) -> None:
        assert TN_UNIVERSITY.type.justify_body is True

    def test_tn_university_front_pagination(self) -> None:
        assert TN_UNIVERSITY.pagination.front.style == "lower_roman"
        assert TN_UNIVERSITY.pagination.front.position == "footer_center"

    def test_tn_university_body_pagination(self) -> None:
        assert TN_UNIVERSITY.pagination.body.style == "arabic"
        assert TN_UNIVERSITY.pagination.body.position == "footer_center"
        assert TN_UNIVERSITY.pagination.body.restart_at == 1

    def test_tn_university_mla_header_name_false(self) -> None:
        assert TN_UNIVERSITY.pagination.mla_header_name is False

    def test_tn_university_front_matter_order(self) -> None:
        assert list(TN_UNIVERSITY.front_matter_order) == [
            "title_page",
            "certificate",
            "declaration",
            "acknowledgement",
            "contents",
        ]

    def test_tn_university_chapter_label(self) -> None:
        cl = TN_UNIVERSITY.chapter_label
        assert cl.format == "CHAPTER {ROMAN}"
        assert cl.title_caps is True
        assert cl.title_bold is True
        assert cl.new_page is True

    def test_tn_university_works_cited_heading_bold(self) -> None:
        wc = TN_UNIVERSITY.works_cited
        assert wc.heading == "WORKS CITED"
        assert wc.heading_bold is True

    # --- MLA_STRICT ---

    def test_mla_strict_page_size(self) -> None:
        assert MLA_STRICT.page.size == "Letter"

    def test_mla_strict_all_margins_one_inch(self) -> None:
        m = MLA_STRICT.page.margins_in
        assert m.top == 1.0
        assert m.bottom == 1.0
        assert m.left == 1.0
        assert m.right == 1.0

    def test_mla_strict_not_justified(self) -> None:
        assert MLA_STRICT.type.justify_body is False

    def test_mla_strict_mla_header_name_true(self) -> None:
        assert MLA_STRICT.pagination.mla_header_name is True

    def test_mla_strict_empty_front_matter(self) -> None:
        assert len(MLA_STRICT.front_matter_order) == 0

    def test_mla_strict_works_cited_not_bold(self) -> None:
        wc = MLA_STRICT.works_cited
        assert wc.heading == "Works Cited"
        assert wc.heading_bold is False

    def test_mla_strict_page_header_right(self) -> None:
        assert MLA_STRICT.pagination.body.position == "header_right"


# ---------------------------------------------------------------------------
# profiles — resolve_profile
# ---------------------------------------------------------------------------


class TestResolveProfile:
    def test_no_override_returns_equivalent_to_builtin(self) -> None:
        p = resolve_profile("tn_university", None)
        assert isinstance(p, ResolvedProfile)
        assert p.page.size == TN_UNIVERSITY.page.size
        assert p.type.justify_body == TN_UNIVERSITY.type.justify_body

    def test_empty_dict_returns_base(self) -> None:
        p = resolve_profile("mla_strict", {})
        assert p.page.size == "Letter"

    def test_override_single_margin(self) -> None:
        """A partial margins_in override touches only the supplied sub-key."""
        override = {"page": {"margins_in": {"top": 1.25}}}
        p = resolve_profile("tn_university", override)
        assert p.page.margins_in.top == 1.25
        # Other margins untouched
        assert p.page.margins_in.left == 1.5
        assert p.page.margins_in.bottom == 1.0
        assert p.page.margins_in.right == 1.0

    def test_override_page_size(self) -> None:
        p = resolve_profile("tn_university", {"page": {"size": "Letter"}})
        assert p.page.size == "Letter"
        # margins untouched
        assert p.page.margins_in.left == 1.5

    def test_override_line_spacing(self) -> None:
        p = resolve_profile("tn_university", {"type": {"line_spacing": 1.5}})
        assert p.type.line_spacing == 1.5
        assert p.type.font == "Times New Roman"  # untouched

    def test_override_front_matter_order(self) -> None:
        override = {
            "front_matter_order": ["title_page", "acknowledgement", "contents"]
        }
        p = resolve_profile("tn_university", override)
        assert list(p.front_matter_order) == [
            "title_page",
            "acknowledgement",
            "contents",
        ]

    def test_override_works_cited_heading(self) -> None:
        p = resolve_profile("tn_university", {"works_cited": {"heading": "Bibliography"}})
        assert p.works_cited.heading == "Bibliography"
        assert p.works_cited.heading_bold is True  # untouched

    def test_unknown_top_level_key_lands_in_notes(self) -> None:
        p = resolve_profile("tn_university", {"fancy_new_option": "yes"})
        assert "fancy_new_option" in p.notes

    def test_unknown_nested_key_lands_in_notes(self) -> None:
        p = resolve_profile("tn_university", {"page": {"paper_colour": "cream"}})
        assert "paper_colour" in p.notes

    def test_multiple_unknown_keys_all_in_notes(self) -> None:
        override = {"alpha": 1, "beta": 2, "page": {"gamma": 3}}
        p = resolve_profile("tn_university", override)
        assert "alpha" in p.notes
        assert "beta" in p.notes
        assert "gamma" in p.notes

    def test_metadata_keys_not_in_notes(self) -> None:
        """id, name, base are metadata — silently ignored, not treated as unknown."""
        p = resolve_profile(
            "tn_university",
            {"id": "abc", "name": "My profile", "base": "tn_university"},
        )
        assert "id" not in p.notes
        assert "name" not in p.notes
        assert "base" not in p.notes

    def test_notes_string_in_override_preserved(self) -> None:
        p = resolve_profile("tn_university", {"notes": "Dept says use 1.5 spacing."})
        assert "Dept says use 1.5 spacing." in p.notes

    def test_unknown_key_appended_to_overridden_notes(self) -> None:
        p = resolve_profile(
            "tn_university",
            {"notes": "Custom note.", "weird_key": 42},
        )
        assert "Custom note." in p.notes
        assert "weird_key" in p.notes

    def test_singletons_not_mutated_after_resolve(self) -> None:
        """resolve_profile must not mutate TN_UNIVERSITY or MLA_STRICT."""
        original_left = TN_UNIVERSITY.page.margins_in.left
        resolve_profile("tn_university", {"page": {"margins_in": {"left": 2.0}}})
        assert TN_UNIVERSITY.page.margins_in.left == original_left

    def test_invalid_base_raises_value_error(self) -> None:
        with pytest.raises(ValueError, match="Unknown base profile"):
            resolve_profile("chicago", None)

    def test_mla_base(self) -> None:
        p = resolve_profile("mla_strict", {"type": {"justify_body": True}})
        assert p.type.justify_body is True
        assert p.page.size == "Letter"


# ---------------------------------------------------------------------------
# pagination — add_page_number_field
# ---------------------------------------------------------------------------


class TestAddPageNumberField:
    def test_fld_simple_instr_in_document_xml(self) -> None:
        """add_page_number_field inserts w:fldSimple with PAGE instr into document body."""
        doc = Document()
        para = doc.add_paragraph()
        add_page_number_field(para)

        root = _read_xml_part(_docx_bytes(doc), "word/document.xml")
        flds = root.findall(f".//{_qn('fldSimple')}")
        instrs = [f.get(_qn("instr"), "") for f in flds]
        assert any("PAGE" in instr for instr in instrs), (
            f"Expected a fldSimple with PAGE instr; found instrs: {instrs}"
        )

    def test_multiple_calls_add_multiple_fields(self) -> None:
        doc = Document()
        para = doc.add_paragraph()
        add_page_number_field(para)
        add_page_number_field(para)

        root = _read_xml_part(_docx_bytes(doc), "word/document.xml")
        flds = [
            f
            for f in root.findall(f".//{_qn('fldSimple')}")
            if "PAGE" in f.get(_qn("instr"), "")
        ]
        assert len(flds) == 2


# ---------------------------------------------------------------------------
# pagination — set_section_pagenum_format
# ---------------------------------------------------------------------------


class TestSetSectionPagenumFormat:
    def test_lower_roman_with_start(self) -> None:
        doc = Document()
        sec = doc.sections[0]
        set_section_pagenum_format(sec, "lowerRoman", start=1)

        root = _read_xml_part(_docx_bytes(doc), "word/document.xml")
        pgt_list = root.findall(f".//{_qn('pgNumType')}")
        assert pgt_list, "Expected at least one w:pgNumType element"
        pgt = pgt_list[-1]
        assert pgt.get(_qn("fmt")) == "lowerRoman"
        assert pgt.get(_qn("start")) == "1"

    def test_decimal_without_start(self) -> None:
        doc = Document()
        sec = doc.sections[0]
        set_section_pagenum_format(sec, "decimal")

        root = _read_xml_part(_docx_bytes(doc), "word/document.xml")
        pgt_list = root.findall(f".//{_qn('pgNumType')}")
        assert pgt_list
        pgt = pgt_list[-1]
        assert pgt.get(_qn("fmt")) == "decimal"
        assert pgt.get(_qn("start")) is None

    def test_calling_twice_does_not_duplicate(self) -> None:
        """Re-calling replaces the existing pgNumType rather than adding a second one."""
        doc = Document()
        sec = doc.sections[0]
        set_section_pagenum_format(sec, "lowerRoman", start=1)
        set_section_pagenum_format(sec, "decimal", start=1)

        root = _read_xml_part(_docx_bytes(doc), "word/document.xml")
        pgt_list = root.findall(f".//{_qn('pgNumType')}")
        # There should be exactly one pgNumType in this single-section document.
        assert len(pgt_list) == 1
        assert pgt_list[0].get(_qn("fmt")) == "decimal"


# ---------------------------------------------------------------------------
# pagination — suppress_first_page_number
# ---------------------------------------------------------------------------


class TestSuppressFirstPageNumber:
    def test_titlePg_present_in_sectPr(self) -> None:
        """suppress_first_page_number must set w:titlePg on the sectPr."""
        doc = Document()
        sec = doc.sections[0]
        suppress_first_page_number(sec)

        root = _read_xml_part(_docx_bytes(doc), "word/document.xml")
        # titlePg lives inside sectPr
        title_pgs = root.findall(f".//{_qn('titlePg')}")
        assert title_pgs, "Expected w:titlePg in sectPr after suppress_first_page_number"

    def test_idempotent(self) -> None:
        """Calling twice should not raise and should still produce exactly one titlePg."""
        doc = Document()
        sec = doc.sections[0]
        suppress_first_page_number(sec)
        suppress_first_page_number(sec)

        root = _read_xml_part(_docx_bytes(doc), "word/document.xml")
        title_pgs = root.findall(f".//{_qn('titlePg')}")
        assert len(title_pgs) == 1


# ---------------------------------------------------------------------------
# pagination — add_mla_header
# ---------------------------------------------------------------------------


class TestAddMlaHeader:
    def _header_xml(self, doc: Document) -> lxml.etree._Element:
        """Return the parsed XML of the first header part in the docx."""
        raw = _docx_bytes(doc)
        with zipfile.ZipFile(io.BytesIO(raw)) as zf:
            header_names = [
                n for n in zf.namelist() if "header" in n and n.endswith(".xml")
            ]
        assert header_names, "No header XML part found in the docx"
        return _read_xml_part(raw, header_names[0])

    def test_page_field_in_header(self) -> None:
        """add_mla_header inserts a PAGE fldSimple into the header part."""
        doc = Document()
        add_mla_header(doc.sections[0], "Sharma")

        root = self._header_xml(doc)
        flds = root.findall(f".//{_qn('fldSimple')}")
        instrs = [f.get(_qn("instr"), "") for f in flds]
        assert any("PAGE" in i for i in instrs), (
            f"Expected PAGE fldSimple in header; found: {instrs}"
        )

    def test_surname_text_in_header(self) -> None:
        """The surname string appears as a text run in the header."""
        doc = Document()
        add_mla_header(doc.sections[0], "Nair")

        root = self._header_xml(doc)
        texts = [
            t.text or "" for t in root.findall(f".//{_qn('t')}")
        ]
        combined = "".join(texts)
        assert "Nair" in combined, f"Surname 'Nair' not found in header; got: {texts}"

    def test_header_right_aligned(self) -> None:
        """The header paragraph uses right alignment (w:jc w:val='right')."""
        doc = Document()
        add_mla_header(doc.sections[0], "Iyer")

        root = self._header_xml(doc)
        jc_elements = root.findall(f".//{_qn('jc')}")
        vals = [jc.get(_qn("val"), "") for jc in jc_elements]
        assert "right" in vals, f"Expected right-aligned paragraph; jc vals: {vals}"
