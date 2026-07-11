"""M3 ingestion engine tests — round-trip + the DESIGN §10 verifier traps."""

from __future__ import annotations

from dataclasses import dataclass, field
from uuid import uuid4

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from app.canonical.model import (
    BlockQuoteBlock,
    HeadingBlock,
    MarkerBlock,
    ParagraphBlock,
    ThesisDocument,
    WorksCitedRef,
)
from app.ingest.citations import VERIFY, parse_wc_entries, scan_document
from app.ingest.docx_extract import extract_paragraphs
from app.ingest.structure import parse_manuscript
from app.ingest.verifier import verify


pytestmark = pytest.mark.asyncio


def _build_manuscript(path: str) -> None:
    d = Document()
    # Title page-ish + front matter
    t = d.add_paragraph("COLONIAL VOICES IN THE AFRICAN NOVEL")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph("A dissertation by Meena Krishnan")
    d.add_paragraph("CERTIFICATE")
    d.add_paragraph("This is to certify that this work is bonafide.")
    d.add_paragraph("DECLARATION")
    d.add_paragraph("I hereby declare this dissertation is my original work.")
    # Chapter I
    d.add_paragraph("CHAPTER I")
    ch_title = d.add_paragraph("INTRODUCTION")
    ch_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = d.add_paragraph()
    p.add_run("Achebe's ")
    r = p.add_run("Things Fall Apart")
    r.italic = True
    p.add_run(" resists the colonial gaze (Achebe 45).")
    bq = d.add_paragraph(
        "The white man is very clever. He has put a knife on the things "
        "that held us together and we have fallen apart. (Achebe 176)"
    )
    bq.paragraph_format.left_indent = Inches(0.5)
    h = d.add_paragraph()
    hr = h.add_run("The Village as Polis")
    hr.bold = True
    d.add_paragraph("Umuofia functions as a civic organism (Gikandi 12).")
    # Chapter II
    d.add_paragraph("CHAPTER II: SUMMATION")
    d.add_paragraph("The argument closes where it began.")
    # Works Cited
    d.add_paragraph("WORKS CITED")
    wc1 = d.add_paragraph()
    wc1.add_run("Achebe, Chinua. ")
    wr = wc1.add_run("Things Fall Apart")
    wr.italic = True
    wc1.add_run(". Heinemann, 1958.")
    wc2 = d.add_paragraph()
    wc2.add_run("Gikandi, Simon. “Chinua Achebe and the Invention of African Culture.” ")
    wj = wc2.add_run("Research in African Literatures")
    wj.italic = True
    wc2.add_run(", vol. 32, no. 3, 2001, pp. 3-8.")
    d.save(path)


async def test_ingest_round_trip(tmp_path) -> None:
    path = str(tmp_path / "ms.docx")
    _build_manuscript(path)

    paras = extract_paragraphs(path)
    assert paras[0].all_caps and paras[0].alignment == "center"

    result = parse_manuscript(paras)
    doc = result.document
    assert isinstance(doc, ThesisDocument)

    # Front matter detected
    kinds = [e.kind for e in doc.front_matter]
    assert kinds[0] == "title_page"
    assert "certificate" in kinds and "declaration" in kinds

    # Chapters: I Introduction (title from following caps line), II inline title
    assert [c.number for c in doc.chapters] == [1, 2]
    assert doc.chapters[0].title == "INTRODUCTION"
    assert doc.chapters[1].title == "SUMMATION"

    blocks = doc.chapters[0].blocks
    # Paragraph keeps italics byte-for-byte
    para = next(b for b in blocks if isinstance(b, ParagraphBlock))
    assert any(r.italic and r.text == "Things Fall Apart" for r in para.runs)
    # Indented paragraph became a block quote with split citation
    bq = next(b for b in blocks if isinstance(b, BlockQuoteBlock))
    assert bq.citation == "Achebe 176"
    assert bq.text.endswith("fallen apart.")
    # Bold short line became a level-2 heading
    assert any(isinstance(b, HeadingBlock) and b.level == 2 for b in blocks)

    # Works Cited raw entries preserved with italics
    assert len(result.wc_raw_entries) == 2

    # Citation extraction
    intext = scan_document(doc)
    surnames = {c.surname for c in intext}
    assert {"Achebe", "Gikandi"} <= surnames

    cands = parse_wc_entries(result.wc_raw_entries)
    assert cands[0].kind == "book"
    assert cands[0].fields == {
        "author": "Achebe, Chinua", "title": "Things Fall Apart",
        "publisher": "Heinemann", "year": "1958",
    }
    assert cands[1].kind == "journal"
    assert cands[1].fields["container"] == "Research in African Literatures"
    assert cands[1].fields["volume"] == "32" and cands[1].fields["pages"] == "3-8"


# --------------------------------------------------------------------------
# Verifier traps (DESIGN §10: fabricated source, orphan citation,
# unverified quote, WC entry never cited — must catch all four)
# --------------------------------------------------------------------------


@dataclass
class _Src:
    kind: str
    fields: dict
    verified: bool = True
    consulted_flag: bool = False


@dataclass
class _Quote:
    text: str
    verified: bool = True


def _doc_with(blocks, wc_ids) -> ThesisDocument:
    return ThesisDocument.model_validate({
        "meta": {}, "front_matter": [],
        "chapters": [{"number": 1, "title": "One", "blocks": blocks}],
        "works_cited": [{"source_id": str(i)} for i in wc_ids],
    })


async def test_verifier_traps() -> None:
    achebe = uuid4()
    unused = uuid4()
    qid = uuid4()
    sources = {
        achebe: _Src("book", {"author": "Achebe, Chinua", "title": "Things Fall Apart",
                              "publisher": "Heinemann", "year": "1958"}),
        unused: _Src("book", {"author": "Zombie, Fake", "title": "Never Cited",
                              "publisher": "Nowhere", "year": "2001"}),
    }
    quotes = {qid: _Quote(text="We have fallen apart.")}

    doc = _doc_with(
        [
            {"type": "paragraph", "runs": [{"text": "A claim (Achebe 45). A fabricated one (Phantom 99)."}]},
            {"type": "block_quote", "text": "We have fallen apart.", "citation": "Achebe 176",
             "quote_id": str(qid)},
            {"type": "block_quote", "text": "No id here.", "citation": "Achebe 12"},
            {"type": "marker", "kind": "QUOTE_NEEDED", "note": "p. 84"},
        ],
        [achebe, unused],
    )
    report = verify(doc, sources, quotes)
    rules = {v.rule for v in report.violations}
    assert not report.passed
    assert "citation_without_source" in rules          # fabricated (Phantom 99)
    assert "quote_missing_id" in rules                 # quote without registry id
    assert "unresolved_marker" in rules
    assert "wc_entry_uncited" in rules                 # Zombie entry never cited
    assert report.counts["block"] >= 3

    # Unverified quote trap
    quotes[qid].verified = False
    report2 = verify(doc, sources, quotes)
    assert "quote_unverified" in {v.rule for v in report2.violations}

    # Text divergence trap
    quotes[qid].verified = True
    quotes[qid].text = "Entirely different words."
    report3 = verify(doc, sources, quotes)
    assert "quote_text_divergence" in {v.rule for v in report3.violations}


async def test_verifier_passes_clean_document() -> None:
    achebe = uuid4()
    qid = uuid4()
    sources = {achebe: _Src("book", {"author": "Achebe, Chinua",
                                     "title": "Things Fall Apart",
                                     "publisher": "Heinemann", "year": "1958"})}
    quotes = {qid: _Quote(text="We have fallen apart.")}
    doc = _doc_with(
        [
            {"type": "paragraph", "runs": [{"text": "A claim (Achebe 45)."}]},
            {"type": "block_quote", "text": "We have fallen apart.",
             "citation": "Achebe 176", "quote_id": str(qid)},
        ],
        [achebe],
    )
    report = verify(doc, sources, quotes)
    assert report.passed and report.counts == {"block": 0, "warn": 0}
