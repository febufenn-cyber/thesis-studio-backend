"""v2 M1+M2 tests — canonical model, renderers, project API, export flow."""

from __future__ import annotations

from uuid import uuid4

import pytest
from docx import Document
from httpx import AsyncClient
from sqlalchemy import delete, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.canonical.model import ThesisDocument
from app.models.project import Project
from app.models.source import Source
from app.models.user import User
from app.renderers.docx_renderer import render_docx
from app.renderers.md_renderer import render_md
from app.renderers.profiles import resolve_profile
from app.renderers.txt_renderer import render_txt
from app.renderers.works_cited import MissingCitationField, format_entry
from tests.conftest import auth_cookie


pytestmark = pytest.mark.asyncio

_SRC_BOOK = str(uuid4())
_SRC_JOURNAL = str(uuid4())
_SRC_FILM = str(uuid4())

MINI_DOC = {
    "meta": {
        "title": "Colonial Consciousness in Nervous Conditions",
        "candidate": {"name": "Meena Krishnan", "reg_no": "MK2026"},
        "college": {
            "name": "St. Ignatius College",
            "affiliation": "University of Madras",
            "city": "Chennai",
            "pin": "600 034",
        },
        "guide": {"name": "Dr. R. Devi", "designation": "Associate Professor"},
        "submission": {"month": "April", "year": 2026},
        "ai_disclosure": {"enabled": True, "text": "AI assisted with formatting only."},
    },
    "front_matter": [
        {"kind": "title_page"},
        {"kind": "certificate"},
        {"kind": "declaration"},
        {"kind": "acknowledgement", "body_blocks": [
            {"type": "paragraph", "runs": [{"text": "The candidate thanks her guide."}]},
        ]},
        {"kind": "ai_disclosure"},
        {"kind": "contents"},
    ],
    "chapters": [
        {"number": 1, "title": "Introduction", "blocks": [
            {"type": "paragraph", "runs": [
                {"text": "Tsitsi Dangarembga's "},
                {"text": "Nervous Conditions", "italic": True},
                {"text": " interrogates the colonial encounter."},
            ]},
            {"type": "block_quote", "text": "The condition of native is a nervous condition.",
             "citation": "Sartre 20"},
        ]},
        {"number": 2, "title": "Education and Alienation", "blocks": [
            {"type": "heading", "level": 2, "text": "The Mission School"},
            {"type": "paragraph", "runs": [{"text": "Education promises mobility."}]},
            {"type": "verse_quote", "lines": ["Quick quick", "the mission bell"],
             "citation": "12-13"},
            {"type": "marker", "kind": "QUOTE_NEEDED", "note": "p. 84 passage"},
        ]},
    ],
    "works_cited": [
        {"source_id": _SRC_BOOK},
        {"source_id": _SRC_JOURNAL},
        {"source_id": _SRC_FILM},
    ],
}


class _Src:
    def __init__(self, kind: str, fields: dict) -> None:
        self.kind = kind
        self.fields = fields


MINI_SOURCES = {
    _SRC_BOOK: _Src("book", {
        "author": "Dangarembga, Tsitsi", "title": "Nervous Conditions",
        "publisher": "Seal Press", "year": "1988",
    }),
    _SRC_JOURNAL: _Src("journal", {
        "author": "Aegerter, Lindsay", "title": "A Dialectic of Autonomy",
        "container": "Tulsa Studies", "volume": "15", "number": "2",
        "year": "1996", "pages": "231-240",
    }),
    _SRC_FILM: _Src("film", {
        "title": "Neria", "director": "Godwin Mawuru",
        "studio": "Media for Development Trust", "year": "1993",
    }),
}


def _doc() -> ThesisDocument:
    return ThesisDocument.model_validate(MINI_DOC)


def _sources_by_uuid() -> dict:
    from uuid import UUID as _UUID
    return {_UUID(k): v for k, v in MINI_SOURCES.items()}


# ---------------------------------------------------------------------------
# Canonical model + renderers
# ---------------------------------------------------------------------------


async def test_canonical_roundtrip() -> None:
    doc = _doc()
    assert doc.meta.candidate.name == "Meena Krishnan"
    assert doc.chapters[1].blocks[0].type == "heading"
    assert len(doc.works_cited) == 3
    # invalid discriminator rejected
    bad = dict(MINI_DOC, chapters=[{"number": 1, "title": "X",
                                    "blocks": [{"type": "nope"}]}])
    with pytest.raises(Exception):
        ThesisDocument.model_validate(bad)


async def test_render_docx_tn_university(tmp_path) -> None:
    prof = resolve_profile("tn_university", None)
    out = str(tmp_path / "tn.docx")
    render_docx(_doc(), _sources_by_uuid(), prof, out)
    d = Document(out)
    sec = d.sections[0]
    assert round(sec.left_margin.inches, 2) == 1.5
    assert d.styles["TS-Normal"].font.name == "Times New Roman"
    texts = [p.text for p in d.paragraphs]
    assert any("CHAPTER I" in t for t in texts)
    assert any("CERTIFICATE" in t for t in texts)
    assert any("DECLARATION" in t for t in texts)
    assert any("AI-ASSISTANCE DISCLOSURE" in t for t in texts)
    assert any("WORKS CITED" in t for t in texts)
    # block quote: 0.5" indent, no first-line indent
    bqs = [p for p in d.paragraphs if p.style.name == "TS-BlockQuote"]
    assert bqs
    assert round(d.styles["TS-BlockQuote"].paragraph_format.left_indent.inches, 2) == 0.5
    assert d.styles["TS-BlockQuote"].paragraph_format.first_line_indent.inches == 0
    # WC entries alphabetised (Aegerter before Dangarembga before Neria)
    wc = [p.text for p in d.paragraphs if p.style.name == "TS-WorksCited"]
    assert len(wc) == 3 and wc[0].startswith("Aegerter") and wc[1].startswith("Dangarembga")
    # WC hanging indent
    wcp = [p for p in d.paragraphs if p.style.name == "TS-WorksCited"][0]
    assert round(d.styles["TS-WorksCited"].paragraph_format.first_line_indent.inches, 2) == -0.5
    # two sections (front matter + body), roman then decimal
    assert len(d.sections) == 2
    xml = d.element.xml
    assert 'w:fmt="lowerRoman"' in xml and 'w:fmt="decimal"' in xml
    # marker rendered visibly
    assert any("[QUOTE_NEEDED:" in t for t in texts)


async def test_render_docx_mla(tmp_path) -> None:
    prof = resolve_profile("mla_strict", None)
    out = str(tmp_path / "mla.docx")
    render_docx(_doc(), _sources_by_uuid(), prof, out)
    d = Document(out)
    assert round(d.sections[0].left_margin.inches, 2) == 1.0
    header_text = "".join(p.text for p in d.sections[0].header.paragraphs)
    assert "Krishnan" in header_text
    # no tn front matter
    texts = [p.text for p in d.paragraphs]
    assert not any("CERTIFICATE" == t.strip() for t in texts)


async def test_works_cited_missing_field_never_guesses() -> None:
    with pytest.raises(MissingCitationField):
        format_entry("book", {"author": "X, Y", "title": "T", "year": "2001"})
    with pytest.raises(MissingCitationField):
        format_entry("journal", {"author": "X, Y", "title": "T",
                                 "container": "J", "volume": "1",
                                 "number": "[VERIFY]", "year": "2001", "pages": "1-2"})


async def test_render_md_and_txt() -> None:
    prof = resolve_profile("tn_university", None)
    md = render_md(_doc(), _sources_by_uuid(), prof)
    assert md.startswith("# Colonial Consciousness")
    assert "## CHAPTER I: INTRODUCTION" in md
    assert "> The condition of native" in md
    assert "*Nervous Conditions*" in md
    assert "## WORKS CITED" in md
    txt = render_txt(_doc(), _sources_by_uuid(), prof)
    assert "CHAPTER II — EDUCATION AND ALIENATION" in txt
    assert all(len(line) <= 80 for line in txt.splitlines())


# ---------------------------------------------------------------------------
# Project API + isolation
# ---------------------------------------------------------------------------


async def _mk_project(client: AsyncClient, user: User, title: str = "Op job") -> dict:
    r = await client.post("/projects", json={"title": title}, cookies=auth_cookie(user))
    assert r.status_code == 201
    return r.json()


async def test_project_crud_and_isolation(
    client: AsyncClient, user_a: User, user_b: User
) -> None:
    proj = await _mk_project(client, user_a)
    # list isolation
    r = await client.get("/projects", cookies=auth_cookie(user_b))
    assert r.status_code == 200 and r.json() == []
    # cross-user 404
    r = await client.get(f"/projects/{proj['id']}", cookies=auth_cookie(user_b))
    assert r.status_code == 404
    # meta update validated
    r = await client.patch(
        f"/projects/{proj['id']}/meta",
        json={"meta": MINI_DOC["meta"]},
        cookies=auth_cookie(user_a),
    )
    assert r.status_code == 200
    assert r.json()["meta"]["candidate"]["name"] == "Meena Krishnan"
    # invalid chapters rejected
    r = await client.patch(
        f"/projects/{proj['id']}/chapters",
        json={"chapters": [{"number": 1, "title": "X", "blocks": [{"type": "bogus"}]}]},
        cookies=auth_cookie(user_a),
    )
    assert r.status_code == 422
    # archive
    r = await client.delete(f"/projects/{proj['id']}", cookies=auth_cookie(user_a))
    assert r.status_code == 204
    r = await client.get(f"/projects/{proj['id']}", cookies=auth_cookie(user_a))
    assert r.status_code == 404


async def test_sources_and_quotes_flow(
    client: AsyncClient, user_a: User, user_b: User
) -> None:
    proj = await _mk_project(client, user_a)
    r = await client.post(
        f"/projects/{proj['id']}/sources",
        json={"kind": "book", "fields": MINI_SOURCES[_SRC_BOOK].fields},
        cookies=auth_cookie(user_a),
    )
    assert r.status_code == 201
    source = r.json()
    assert source["verified"] is False
    # quote must attach to an owned source
    r = await client.post(
        f"/projects/{proj['id']}/sources/{source['id']}/quotes",
        json={"text": "The condition of native is a nervous condition.",
              "page_or_loc": "20"},
        cookies=auth_cookie(user_a),
    )
    assert r.status_code == 201
    # cross-user cannot see or delete
    r = await client.get(f"/projects/{proj['id']}/sources", cookies=auth_cookie(user_b))
    assert r.status_code == 404
    r = await client.delete(
        f"/projects/{proj['id']}/sources/{source['id']}", cookies=auth_cookie(user_b)
    )
    assert r.status_code == 404


async def test_export_gates(
    client: AsyncClient, user_a: User, monkeypatch: pytest.MonkeyPatch
) -> None:
    proj = await _mk_project(client, user_a)
    cookies = auth_cookie(user_a)
    # G4 acknowledgment required
    r = await client.post(f"/projects/{proj['id']}/exports",
                          json={"formats": ["md"], "acknowledge": False}, cookies=cookies)
    assert r.status_code == 409 and "acknowledg" in r.json()["detail"].lower()
    # no chapters
    r = await client.post(f"/projects/{proj['id']}/exports",
                          json={"formats": ["md"], "acknowledge": True}, cookies=cookies)
    assert r.status_code == 409 and "no chapters" in r.json()["detail"]
    # happy path with stubbed background job
    await client.patch(f"/projects/{proj['id']}/chapters",
                       json={"chapters": MINI_DOC["chapters"]}, cookies=cookies)

    async def _noop(*args) -> None:
        return None

    monkeypatch.setattr("app.api.projects.run_export", _noop)
    r = await client.post(f"/projects/{proj['id']}/exports",
                          json={"formats": ["md", "txt"], "acknowledge": True},
                          cookies=cookies)
    assert r.status_code == 202
    rows = r.json()
    assert {x["format"] for x in rows} == {"md", "txt"}
    assert all(x["status"] == "running" for x in rows)
    # download before ready → 409
    r = await client.get(f"/exports/{rows[0]['id']}/download", cookies=cookies)
    assert r.status_code == 409


async def test_run_export_end_to_end(tmp_path, monkeypatch: pytest.MonkeyPatch) -> None:
    """Full background export of docx+md with local storage and real renderers.

    Seeds committed rows (run_export opens its own session) and cleans up.
    """
    from uuid import UUID as _UUID

    from app.core.config import get_settings
    from app.db.session import AsyncSessionLocal, engine
    from app.models.export import Export
    from app.models.institution import Institution
    from app.services.export_service import run_export
    from app.services.storage_service import LocalStorageService

    settings_obj = get_settings()
    monkeypatch.setattr(settings_obj, "LOCAL_STORAGE_DIR", str(tmp_path))
    storage = LocalStorageService()
    monkeypatch.setattr(
        "app.services.export_service.get_storage_service", lambda: storage
    )

    inst_id = user_id = project_id = None
    export_ids = []
    async with AsyncSessionLocal() as db:
        inst = Institution(
            name="Export College", short_name=f"EX{uuid4().hex[:6]}",
            email_domains="ex.edu", address="1 Export St",
            short_address="Exville", university_name="Test U",
            default_department="English", department_aided=False,
        )
        db.add(inst)
        await db.commit()
        await db.refresh(inst)
        inst_id = inst.id
        user = User(email=f"ex-{uuid4().hex[:8]}@ex.edu", full_name="Exa Porter",
                    institution_id=inst.id)
        db.add(user)
        await db.commit()
        await db.refresh(user)
        user_id = user.id
        project = Project(
            user_id=user.id, title="Export Project",
            meta=MINI_DOC["meta"], front_matter=MINI_DOC["front_matter"],
            chapters=MINI_DOC["chapters"], works_cited=MINI_DOC["works_cited"],
        )
        db.add(project)
        await db.commit()
        await db.refresh(project)
        project_id = project.id
        for k, src in MINI_SOURCES.items():
            db.add(Source(id=_UUID(k), project_id=project.id, user_id=user.id,
                          kind=src.kind, fields=src.fields, verified=True))
        for fmt in ("docx", "md"):
            row = Export(project_id=project.id, user_id=user.id, format=fmt,
                         status="running")
            db.add(row)
            await db.commit()
            await db.refresh(row)
            export_ids.append((row.id, fmt))

    try:
        for export_id, _fmt in export_ids:
            await run_export(export_id, project_id, user_id)
        async with AsyncSessionLocal() as db:
            for export_id, fmt in export_ids:
                row = (await db.execute(
                    select(Export).where(Export.id == export_id)
                )).scalar_one()
                assert row.status == "ready", f"{fmt}: {row.error_message}"
                assert row.checksum and row.size_bytes > 0
                path = await storage.open_local_path(row.storage_key)
                if fmt == "docx":
                    d = Document(path)
                    assert any("CHAPTER I" in p.text for p in d.paragraphs)
                else:
                    content = open(path, encoding="utf-8").read()
                    assert "## CHAPTER I" in content
    finally:
        async with AsyncSessionLocal() as db:
            await db.execute(delete(Export).where(Export.project_id == project_id))
            await db.execute(delete(Source).where(Source.project_id == project_id))
            await db.execute(delete(Project).where(Project.id == project_id))
            await db.execute(delete(User).where(User.id == user_id))
            await db.execute(delete(Institution).where(Institution.id == inst_id))
            await db.commit()
        await engine.dispose()
